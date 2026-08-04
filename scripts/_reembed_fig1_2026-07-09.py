# -*- coding: utf-8 -*-
"""Re-embed the rebuilt Figure 1 (workflow) into the docx.

Fig 1 was an ultra-wide 6.7:1 strip embedded at 7.27x1.08in (text microscopic,
sides cropped). The rebuild is 2.11:1; we keep the display WIDTH (7.27in) and
recompute a taller height so the boxes are legible. Field-safe (image part only);
byte-length + shape-index guard; backup; verify. Dry-run default.

  shape[0] Fig1 workflow <- fig1_open_workflow.png (3026x1436)  h 1.08 -> ~3.45in

Usage: python scripts/_reembed_fig1_2026-07-09.py [--apply]
"""
from __future__ import annotations

import re
import shutil
import struct
import sys
import zipfile
from pathlib import Path

import docx
from docx.shared import Emu

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs" / "paper1_open_source_alternative"
DOC = BASE / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_fig1_2026-07-09.docx")
SRC = BASE / "figures" / "fig1_open_workflow.png"

SHAPE_IDX = 0
OLD_BYTES_LEN = 38642   # current embedded Fig 1 (ultra-wide 2352x351)


def png_size(b: bytes) -> tuple[int, int]:
    assert b[12:16] == b"IHDR"
    return struct.unpack(">II", b[16:24])


def field_counts(p: Path) -> tuple[int, int]:
    fld = instr = 0
    with zipfile.ZipFile(p) as z:
        for n in z.namelist():
            if n.endswith(".xml"):
                x = z.read(n).decode("utf-8", "replace")
                fld += len(re.findall(r"<w:fldChar\b", x))
                instr += len(re.findall(r"<w:instrText\b", x))
    return fld, instr


def main(apply: bool) -> int:
    nb = SRC.read_bytes()
    nw, nh = png_size(nb)
    d = docx.Document(str(DOC))
    if len(d.inline_shapes) != 8:
        raise SystemExit(f"ABORT: expected 8 inline shapes, got {len(d.inline_shapes)}")
    sh = d.inline_shapes[SHAPE_IDX]
    part = d.part.related_parts[sh._inline.graphic.graphicData.pic.blipFill.blip.embed]
    cur = part._blob
    cw, ch = png_size(cur)
    if len(cur) != OLD_BYTES_LEN:
        raise SystemExit(f"ABORT: shape[0] bytes {len(cur)} != guard {OLD_BYTES_LEN}")

    w_emu = int(sh.width)
    old_h_emu = int(sh.height)
    new_h_emu = int(w_emu * nh / nw)
    print(f"Fig1: {cw}x{ch} -> {nw}x{nh}px   width {w_emu/914400:.2f}in kept, "
          f"height {old_h_emu/914400:.2f} -> {new_h_emu/914400:.2f}in, "
          f"eff DPI {nw/(w_emu/914400):.0f}")

    if not apply:
        print("\nDRY-RUN. Re-run with --apply.")
        return 0

    fb, ib = field_counts(DOC)
    shutil.copy2(DOC, BACKUP)
    print(f"backup -> {BACKUP.name}")
    part._blob = nb
    sh.height = Emu(new_h_emu)
    d.save(str(DOC))
    fa, ia = field_counts(DOC)
    d2 = docx.Document(str(DOC))
    b = d2.part.related_parts[d2.inline_shapes[0]._inline.graphic.graphicData.pic.blipFill.blip.embed]._blob
    print(f"SAVED. [verify] shapes={len(d2.inline_shapes)} paragraphs={len(d2.paragraphs)} "
          f"tables={len(d2.tables)} Fig1={png_size(b)} {len(b)}B  "
          f"fields fldChar {fb}->{fa} instrText {ib}->{ia} "
          f"[{'OK' if (fb, ib) == (fa, ia) else 'CHANGED!!'}]")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
