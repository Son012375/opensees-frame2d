"""Case 1 closed-form 3-way comparison Table 4a를 v2 docx의 §3.3.1 직후에 삽입.

`_patch_paper_draft.py` 실행 후, `_add_appendix_a2.py` 전에 실행해야 한다.

Para 83 (§3.3.1 body)을 수정하는 것은 _patch_paper_draft.py가 담당하고,
본 스크립트는 Table 4a (caption + 4×6 table)만 mid-document XML 삽입한다.

운영 순서:
    1. python scripts/_patch_paper_draft.py        # 원본 → v2 (모든 text rewrites)
    2. python scripts/_add_case1_table5a.py        # v2 → v2 (Table 4a 삽입)
    3. python scripts/_add_appendix_a2.py          # v2 → v2 (Appendix A 추가)

본 스크립트는 idempotent하지 않다 — 재실행 전 _patch_paper_draft.py로 v2를
원본 기반 재생성해야 한다.
"""
from __future__ import annotations

from pathlib import Path

import docx
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"

# Para 83은 §3.3.1 body. Table 4a는 이 직후에 삽입.
ANCHOR_PARA_INDEX = 83


# ============================================================
# Table data
# ============================================================

TABLE_5A_HEADERS = [
    "Quantity",
    "Closed-form",
    "OpenSees",
    "Midas Gen",
    "Δ OS−CF",
    "Δ MG−CF",
]

TABLE_5A_ROWS = [
    ["Midspan deflection |δ| (mm)", "5.4250", "5.4250", "5.4250", "<0.001%", "<0.001%"],
    ["Maximum bending moment |M_max| (kN·m)", "90.00", "90.00", "90.00", "0.0%", "0.0%"],
    ["Support reaction R_A = R_B (kN)", "30.00", "30.00", "30.00", "0.0%", "0.0%"],
]

TABLE_5A_COL_CM = [5.2, 2.4, 2.4, 2.4, 1.8, 1.8]


# ============================================================
# Helpers (subset of _add_appendix_a2.py)
# ============================================================

def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _populate_table(t, headers, rows, col_widths_cm, font_size=9) -> None:
    t.autofit = False
    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[ci].width = Cm(w)

    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
        _set_cell_borders(cell)

    for ri, row in enumerate(rows, start=1):
        for ci, v in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(v)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
            _set_cell_borders(cell)


# ============================================================
# Main
# ============================================================

def main() -> int:
    print(f"[*] Opening: {TARGET.relative_to(ROOT)}")
    d = docx.Document(str(TARGET))

    anchor_para = d.paragraphs[ANCHOR_PARA_INDEX]
    anchor_text_preview = anchor_para.text[:80]
    print(f"[*] Anchor para {ANCHOR_PARA_INDEX}: {anchor_text_preview!r}")

    # 1) Caption paragraph (created at end of document)
    caption_para = d.add_paragraph()
    caption_run = caption_para.add_run(
        "Table 4a. Case 1 three-way comparison: closed-form analytical "
        "solution, OpenSees, and Midas Gen."
    )
    caption_run.bold = True
    caption_run.font.size = Pt(10)

    # 2) Table (created at end of document)
    new_table = d.add_table(rows=1 + len(TABLE_5A_ROWS), cols=len(TABLE_5A_HEADERS))
    _populate_table(new_table, TABLE_5A_HEADERS, TABLE_5A_ROWS, TABLE_5A_COL_CM)

    # 3) Move caption and table to immediately after anchor paragraph.
    # addnext inserts as the *next* sibling. Order of calls is reverse of
    # final order: first move table (becomes anchor+1), then caption (anchor+1
    # pushes table to anchor+2). Result: anchor → caption → table → ...
    anchor_para._element.addnext(new_table._tbl)
    anchor_para._element.addnext(caption_para._element)

    d.save(str(TARGET))

    # Verify positioning
    d2 = docx.Document(str(TARGET))
    caption_at = ANCHOR_PARA_INDEX + 1
    print(f"[OK] Table 4a inserted after para {ANCHOR_PARA_INDEX}")
    print(f"     New caption (para {caption_at}): {d2.paragraphs[caption_at].text[:80]!r}")
    print(f"     Total paragraphs: {len(d2.paragraphs)}, total tables: {len(d2.tables)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
