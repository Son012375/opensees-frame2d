"""C1: re-embed the regenerated screening-summary figure (0.507 -> 0.467) into the
canonical v2 DOCX. This figure is the manuscript's "Figure 6" but corresponds to
inline shape index 5 (/word/media/image6.png) and to the generator's
fig5_screening_summary.png (final/fig5.png). The final/figN naming is offset from
the manuscript Figure numbering, so we target the shape by INDEX + an old-bytes
guard, not by the misleading final/fig6.png filename.

Steps: refresh final/fig5.png from the new render, back up the DOCX
(..._v2.pre_fig6_2026-07-01.docx), swap shape[5]'s image bytes (keep width,
recompute height from aspect), save, verify counts. Dry-run by default; --apply saves.
"""
from __future__ import annotations

import shutil
import struct
import sys
from pathlib import Path

import docx
from docx.shared import Emu

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs" / "paper1_open_source_alternative"
DOC = BASE / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_fig6_2026-07-01.docx")
NEW_PNG = BASE / "figures" / "fig5_screening_summary.png"
FINAL_PNG = BASE / "figures" / "final" / "fig5.png"

SHAPE_IDX = 5              # screening summary = manuscript Fig 6 = inline shape #6
OLD_BYTES_LEN = 41834     # known length of the pre-fix embedded screening image


def png_size(b: bytes) -> tuple[int, int]:
    assert b[12:16] == b"IHDR", "not a PNG / unexpected header"
    w, h = struct.unpack(">II", b[16:24])
    return w, h


def main(apply: bool) -> int:
    new_bytes = NEW_PNG.read_bytes()
    nw, nh = png_size(new_bytes)
    print(f"new screening png: {NEW_PNG.name} {nw}x{nh}px {len(new_bytes)} bytes")

    d = docx.Document(str(DOC))
    shapes = d.inline_shapes
    if len(shapes) != 8:
        raise SystemExit(f"ABORT: expected 8 inline shapes, got {len(shapes)}")
    sh = shapes[SHAPE_IDX]
    blip = sh._inline.graphic.graphicData.pic.blipFill.blip
    rId = blip.embed
    part = d.part.related_parts[rId]
    cur = part._blob
    cw, ch = png_size(cur)
    print(f"target shape[{SHAPE_IDX}]: rId={rId} partname={part.partname} "
          f"{cw}x{ch}px {len(cur)} bytes")
    if len(cur) != OLD_BYTES_LEN:
        raise SystemExit(
            f"ABORT: shape[{SHAPE_IDX}] current bytes {len(cur)} != expected old "
            f"{OLD_BYTES_LEN}; the screening image may have moved. Not touching.")

    old_w_emu = int(sh.width)
    old_h_emu = int(sh.height)
    new_h_emu = int(old_w_emu * nh / nw)
    print(f"display: width={old_w_emu/914400:.2f}in kept; "
          f"height {old_h_emu/914400:.2f}->{new_h_emu/914400:.2f}in")

    if not apply:
        print("\nDRY-RUN. Re-run with --apply to swap image + write final/fig5.png.")
        return 0

    # refresh the 'final' publication copy too
    shutil.copy2(NEW_PNG, FINAL_PNG)
    print(f"refreshed {FINAL_PNG.relative_to(ROOT)}")

    shutil.copy2(DOC, BACKUP)
    print(f"backup -> {BACKUP.name}")

    part._blob = new_bytes
    sh.height = Emu(new_h_emu)
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")

    d2 = docx.Document(str(DOC))
    media = [p.partname for p in d2.part.package.iter_parts() if "media" in str(p.partname)]
    swapped = d2.part.related_parts[d2.inline_shapes[SHAPE_IDX]._inline.graphic.graphicData.pic.blipFill.blip.embed]._blob
    print(f"[verify] inline_shapes={len(d2.inline_shapes)} media={len(media)} "
          f"paragraphs={len(d2.paragraphs)} tables={len(d2.tables)} "
          f"shape[{SHAPE_IDX}]_bytes={len(swapped)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
