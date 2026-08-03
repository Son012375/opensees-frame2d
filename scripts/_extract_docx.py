"""Dump the canonical v2 manuscript to a P-numbered text file.

Format:
  Pnnn [style]  paragraph text   (nnn = zero-padded paragraph index)
Tables are dumped after paragraphs with a global line counter so the checklist's
"dump lines 286-296" style references stay stable.

Run: opensees-mcp/Scripts/python.exe scratchpad/_extract_docx.py
Out:  d:/tmp/manuscript_v2_dump.txt
"""
from __future__ import annotations

from pathlib import Path

import docx

ROOT = Path(r"d:/son/opensees-MCP")
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
OUT = Path(r"d:/tmp/manuscript_v2_dump.txt")


def main() -> int:
    d = docx.Document(str(DOC))
    lines: list[str] = []
    lines.append(f"# dump of {DOC.name}")
    lines.append(f"# paragraphs={len(d.paragraphs)} tables={len(d.tables)} "
                 f"inline_shapes={len(d.inline_shapes)}")
    lines.append("")
    lines.append("=== PARAGRAPHS (Pnnn [style]) ===")
    for i, p in enumerate(d.paragraphs):
        txt = p.text.replace("\n", "\\n")
        style = p.style.name if p.style else "?"
        nruns = len(p.runs)
        lines.append(f"P{i:03d} [{style}] (runs={nruns}) {txt}")

    lines.append("")
    lines.append("=== TABLES ===")
    for ti, t in enumerate(d.tables):
        lines.append(f"--- Table T{ti} ({len(t.rows)}x{len(t.columns)}) ---")
        for ri, row in enumerate(t.rows):
            cells = [c.text.replace("\n", "\\n") for c in row.cells]
            lines.append(f"T{ti}R{ri}: " + " | ".join(cells))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"wrote {OUT} ({len(lines)} lines)")
    print(f"paragraphs={len(d.paragraphs)} tables={len(d.tables)} "
          f"inline_shapes={len(d.inline_shapes)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
