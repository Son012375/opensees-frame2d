"""Audit follow-up fixes (user-approved) — Tasks A, B, D.

TASK A — Appendix A reconciled to the actual 7-case / 36-combination analysis:
  P229  : snow removed from the "not generated" list (snow S=0.5 IS produced); note
          the KDS vertical-seismic DL factors (1.2924/0.8076) + orthogonal 100/30.
  Table A-7 (T17): expand to 36 (Gravity 4 incl. snow companions; Seismic X/Y 12 each
          incl. 100/30 orthogonal; Wind 4+4); Total 18->36; fix T17R6 "Section 4.1" ref.
  P231  : "eighteen-combination set" -> "thirty-six-combination set".

TASK B — §4 factual corrections:
  P114  : "support reactions agree within 0.01%" was false (4 CHECK metrics are reaction
          moments ~1.57%, vertical reactions 0.48-0.63%) -> rephrase to force-reaction sums.
  P115  : rigid-zone 0.51 reproduces drift+story-1 disp to 0.1% (roof ~2.8% off); the
          +6.6% column-stiffness variant reproduces drift AND roof to 0.1%.
  P116  : caption "A 0.5 rigid-zone factor" -> "A 0.51 rigid-zone factor".
  Table 7 (T7) R4: relabel "factor 0.50" -> "factor 0.51" with the exact-match values
          (0.000640 / 0.0% / 5.057 / 1.920) from case4_ablation.json.
  P167  : "Section 4.3.2" -> "Section 4.4" (rigid-zone attribution is in §4.4).

TASK D — comparison-metric convention (compare.py uses ref=max; the manuscript formula
          wrongly stated ref=Midas, so Table 6's 3.90% did not match its own formula):
  P093  : denominator |x_Midas| -> max(|x_OpenSees|, |x_Midas|) (symmetric ref, = the tool).
  P094  : prose updated to the symmetric-reference definition.
  P116  : append a note that Table 7's "vs Midas" column is the signed diff vs Midas
          (hence Case 4 +4.1% there vs the 3.90% symmetric max in Table 6).

Backup -> ..._v2.pre_auditfixes_2026-07-01.docx. Dry-run default; --apply saves.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_auditfixes_2026-07-01.docx")


def run_replace(par, old, new, tag, whole=False):
    for r in par.runs:
        if whole and r.text == old:
            r.text = new
            return
        if not whole and old in r.text:
            r.text = r.text.replace(old, new)
            return
    raise SystemExit(f"ABORT [{tag}]: not found: {old[:50]!r}")


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:40]!r}")
    return hits[0]


def cell_set(cell, old, new, tag):
    par = cell.paragraphs[0]
    if len(par.runs) != 1 or par.runs[0].text != old:
        raise SystemExit(f"ABORT [{tag}]: cell != {old!r} (got runs={len(par.runs)} {par.text!r})")
    par.runs[0].text = new


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs
    tables = d.tables

    # guards against double-apply
    if any("thirty-six-combination set" in p.text for p in paras):
        raise SystemExit("ABORT: already applied (thirty-six present).")

    # ---------- TASK A ----------
    p229 = find_para(paras, "are not generated for this example", "P229")
    run_replace(
        p229,
        " 3 combinations involving roof live load (L_r), snow (S), or rain (R) are not "
        "generated for this example because those load cases are not produced by the "
        "workflow for the Seoul site.",
        " 3 combinations involving roof live load (L_r) and rain (R) are not generated "
        "for this example because those load cases are not produced by the workflow for "
        "the Seoul site. The roof snow load (S = 0.5 kN/m2), which the workflow does "
        "generate for the Seoul site, is included as a companion action in the gravity "
        "and seismic combinations, and the seismic combinations apply the KDS "
        "vertical-seismic dead-load factors (1.2924 and 0.8076) together with the "
        "orthogonal 100%/30% directional combinations.",
        "P229")

    p231 = find_para(paras, "eighteen-combination set", "P231")
    run_replace(p231, "eighteen-combination set", "thirty-six-combination set", "P231")

    t17 = next(t for t in tables if [c.text for c in t.rows[0].cells][:2] == ["Group", "Count"])
    cell_set(t17.rows[1].cells[0], "Gravity only", "Gravity", "T17R1C0")
    cell_set(t17.rows[1].cells[1], "2", "4", "T17R1C1")
    cell_set(t17.rows[1].cells[2], "1.4 DL; 1.2 DL + 1.6 LL",
             "1.4 DL; 1.2 DL + 1.6 LL; 1.2 DL + 1.6 LL + 0.5 S; 1.2 DL + 1.0 LL + 1.6 S", "T17R1C2")
    cell_set(t17.rows[2].cells[0], "Seismic X (±)", "Seismic X (±, incl. 100%/30% orthogonal)", "T17R2C0")
    cell_set(t17.rows[2].cells[1], "4", "12", "T17R2C1")
    cell_set(t17.rows[2].cells[2], "1.2 DL + 1.0 LL ± 1.0 EQX; 0.9 DL ± 1.0 EQX",
             "1.2924 DL + 1.0 LL ± 1.0 EQX (± 0.3 EQY) + 0.2 S; 0.8076 DL ± 1.0 EQX (± 0.3 EQY)", "T17R2C2")
    cell_set(t17.rows[3].cells[0], "Seismic Y (±)", "Seismic Y (±, incl. 100%/30% orthogonal)", "T17R3C0")
    cell_set(t17.rows[3].cells[1], "4", "12", "T17R3C1")
    cell_set(t17.rows[3].cells[2], "1.2 DL + 1.0 LL ± 1.0 EQY; 0.9 DL ± 1.0 EQY",
             "1.2924 DL + 1.0 LL ± 1.0 EQY (± 0.3 EQX) + 0.2 S; 0.8076 DL ± 1.0 EQY (± 0.3 EQX)", "T17R3C2")
    cell_set(t17.rows[6].cells[1], "18", "36", "T17R6C1")
    cell_set(t17.rows[6].cells[2], "matches the count reported in Section 4.1",
             "matches the KDS 41 12 00 §1.7 requirement for the seven load cases (see Section 5.1)", "T17R6C2")

    # ---------- TASK B ----------
    p114 = find_para(paras, "support reactions agree within 0.01%", "P114")
    run_replace(
        p114,
        "The 12 CHECK metrics observed in Case 4 all involve lateral displacements, story "
        "drift ratios, or bending moments, whereas support reactions agree within 0.01%.",
        "The 12 CHECK metrics observed in Case 4 involve lateral displacements, story "
        "drift ratios, and moment quantities (element bending moments together with the "
        "four base reaction-moment components), whereas the base force reactions and "
        "their global sums agree to within ~0.1%.",
        "P114")

    p115 = find_para(paras, "reproduces the Midas Story 1 drift to within 0.1% and the story and roof", "P115")
    run_replace(
        p115,
        "a rigid-zone factor of 0.51, essentially the default panel-zone setting of Midas "
        "Gen, reproduces the Midas Story 1 drift to within 0.1% and the story and roof "
        "displacements to within 0.1%. Equivalently, the gap corresponds to a uniform 6.6% "
        "increase in effective column stiffness.",
        "a rigid-zone factor of 0.51, essentially the default panel-zone setting of Midas "
        "Gen, reproduces the Midas Story 1 drift and story-1 displacement to within 0.1% "
        "(the roof displacement remaining about 2.8% below the Midas value). Equivalently, "
        "the gap corresponds to a uniform 6.6% increase in effective column stiffness, "
        "which reproduces the Midas Story 1 drift and roof displacement to within 0.1%.",
        "P115")

    p116 = find_para(paras, "A 0.5 rigid-zone factor", "P116")
    run_replace(p116, "A 0.5 rigid-zone factor", "A 0.51 rigid-zone factor", "P116-factor")
    run_replace(
        p116,
        "whereas shear deformation increases the drift (wrong direction).",
        "whereas shear deformation increases the drift (wrong direction). The “vs Midas” "
        "column reports the signed difference relative to the Midas Gen value, so the "
        "Case 4 baseline reads +4.1% here versus the 3.90% symmetric maximum in Table 6.",
        "P116-note")

    t7 = next(t for t in tables if [c.text for c in t.rows[0].cells][:2] == ["Model variant", "Story Drift 1"])
    cell_set(t7.rows[4].cells[0], "+ rigid end zone, factor 0.50", "+ rigid end zone, factor 0.51", "T7R4C0")
    cell_set(t7.rows[4].cells[1], "0.000641", "0.000640", "T7R4C1")
    cell_set(t7.rows[4].cells[2], "+0.1%", "0.0%", "T7R4C2")
    cell_set(t7.rows[4].cells[3], "5.065", "5.057", "T7R4C3")
    cell_set(t7.rows[4].cells[4], "1.922", "1.920", "T7R4C4")

    p167 = find_para(paras, "rigid-zone effect identified for Case 4 in Section", "P167")
    run_replace(p167, ".3.2 could be offered", ".4 could be offered", "P167")

    # ---------- TASK D ----------
    p093 = find_para(paras, "x_OpenSees", "P093")
    run_replace(p093, "| / |", "| / max(|x_OpenSees|, |", "P093-den", whole=True)
    run_replace(p093, "| × 100 (%)", "|) × 100 (%)", "P093-close", whole=True)
    p094 = find_para(paras, "is used as the reference value in the benchmark comparison", "P094")
    run_replace(
        p094,
        " is used as the reference value in the benchmark comparison.",
        " appears in the numerator and the denominator is the larger of the two program "
        "magnitudes, so the reported relative difference is symmetric and does not "
        "privilege either program.",
        "P094")

    print("All audit fixes staged (Tasks A, B, D).")
    if not apply:
        print("\nDRY-RUN. P229 tail:", p229.text[-120:])
        print("T17 total row:", [t17.rows[6].cells[i].text for i in range(3)])
        print("P093:", p093.text)
        print("T7R4:", [t7.rows[4].cells[i].text for i in range(5)])
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"backup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
