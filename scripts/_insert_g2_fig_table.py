"""Insert Table 9 + Figure 8 into the v2 DOCX, right after the Case 4 attribution
paragraph (§3.3.2). Mid-document insertion via lxml addnext; styles copied from
the document's existing captions ('Normal') and tables ('Table Grid')."""
from __future__ import annotations

import shutil
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs" / "paper1_open_source_alternative"
DOC = BASE / "drafts" / "open_source_alternative_review_draft_v2.docx"
FIG8 = BASE / "figures" / "final" / "fig8.png"

ANCHOR = "do not indicate an error in the open-source analysis chain"

TABLE9 = [
    ("Model variant", "Story Drift 1", "vs Midas", "Roof disp. (mm)", "Story-1 disp. (mm)"),
    ("OpenSeesPy baseline (centerline; ETABS identical)", "0.000666", "+4.1%", "5.377", "1.998"),
    ("+ shear deformation (Timoshenko)", "0.000702", "+9.6%", "5.602", "2.105"),
    ("+ rigid end zone, factor 0.50", "0.000641", "+0.1%", "5.065", "1.922"),
    ("+ rigid end zone, factor 1.00 (full panel zone)", "0.000614", "-4.0%", "4.759", "1.843"),
    ("equivalent +6.6% column stiffness", "0.000640", "0.0%", "5.199", "1.921"),
    ("Midas Gen (reference)", "0.000640", "-", "5.201", "1.921"),
]

TBL_CAPTION = (
    "Table 9. Case 4 discrepancy ablation. Story Drift 1 (Midas Gen reference = "
    "0.000640); the ETABS centerline model returns the OpenSeesPy baseline values "
    "(0.00% difference). A 0.5 rigid-zone factor — Midas Gen's default panel-zone "
    "setting — reproduces Midas to within 0.1%, equivalent to a uniform +6.6% "
    "effective column stiffness, whereas shear deformation increases the drift "
    "(wrong direction)."
)
FIG_CAPTION = (
    "Figure 8. Case 4 discrepancy attribution: Story 1 drift versus the beam-column "
    "rigid-zone (end-offset) factor. The centerline model (factor 0; OpenSeesPy ≡ "
    "ETABS) lies +4.1% above the Midas Gen reference and the full geometric panel "
    "zone (factor 1) lies 4.0% below it; Midas is reproduced at a rigid-zone factor "
    "of 0.51."
)


def main() -> int:
    shutil.copy2(DOC, DOC.with_name("open_source_alternative_review_draft_v2.pre_g2insert_backup.docx"))
    d = docx.Document(str(DOC))

    anchor = next((p for p in d.paragraphs if ANCHOR in p.text), None)
    if anchor is None:
        print("ANCHOR NOT FOUND — aborting")
        return 1

    # --- build elements (appended at end, then relocated) ---
    cap_t = d.add_paragraph(TBL_CAPTION, style="Normal")

    tbl = d.add_table(rows=len(TABLE9), cols=5)
    tbl.style = "Table Grid"
    tbl.autofit = True
    for ri, row in enumerate(TABLE9):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.text = val
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)
                    if ri == 0 or ri == len(TABLE9) - 1:  # header + Midas row bold
                        run.font.bold = True

    p_img = d.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(FIG8), width=Inches(5.5))

    cap_f = d.add_paragraph(FIG_CAPTION, style="Normal")

    # --- relocate after the anchor, in order ---
    cur = anchor._p
    for elem in (cap_t._p, tbl._tbl, p_img._p, cap_f._p):
        cur.addnext(elem)
        cur = elem

    d.save(str(DOC))

    # verify
    d2 = docx.Document(str(DOC))
    media = [p.partname for p in d2.part.package.iter_parts() if "media" in p.partname]
    has_t9 = any("Table 9. Case 4 discrepancy ablation" in p.text for p in d2.paragraphs)
    has_f8 = any("Figure 8. Case 4 discrepancy attribution" in p.text for p in d2.paragraphs)
    print(f"[saved] inline_shapes={len(d2.inline_shapes)} media={len(media)} "
          f"tables={len(d2.tables)} paragraphs={len(d2.paragraphs)} "
          f"Table9_caption={has_t9} Fig8_caption={has_f8}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
