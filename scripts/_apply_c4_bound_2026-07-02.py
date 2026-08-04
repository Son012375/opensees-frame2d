"""Session 2026-07-02 (MAJOR REVISION): C4 — Case-4 panel-zone attribution -> non-unique bound.

Per case4_ablation.json: baseline centerline = +4.06% vs Midas (ETABS identical);
shear (Timoshenko) +9.6% (wrong direction, ruled out); rigid_offset 0.51 matches Story-1
drift & story-1 disp to 0.1% but leaves roof ~2.8% low (5.057 vs 5.201); colI +6.6% matches
Story-1 drift & roof to 0.1% but implies a different story-2 drift (0.001093 vs 0.001046).
=> two single-parameter mechanisms bracket the gap on DISJOINT metric subsets; neither
reproduces the full field. Downgrade "attribution"/"equivalently" -> a non-unique BOUND.

Edits (plain runs, field-safe):
  P113  rewrite (§4.4 ablation paragraph): attribute -> bound; state roof residual + the
        two mechanisms are non-unique; a true geometric panel zone would be needed.
  P114  Table 7 caption: "discrepancy ablation" -> "discrepancy bound"; drop "equivalent".
  P116  Figure 4 caption: "attribution" -> "bound"; note 0.51 does not reproduce the roof.
  T7R5  row label "equivalent +6.6% column stiffness" -> "alternative: +6.6% column stiffness".

Table 7 numeric rows unchanged (they are the honest ablation data). Backup ->
..._v2.pre_c4bound_2026-07-02.docx. Dry-run default; --apply.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_c4bound_2026-07-02.docx")

LOC_P113 = "the Case 4 frame was re-analyzed in OpenSeesPy under a controlled ablation"
NEW_P113 = (
    "To bound this discrepancy rather than leave it qualitative, the Case 4 frame was "
    "re-analyzed in OpenSeesPy under a controlled ablation of the candidate drivers (Table 7, "
    "Figure 4). Because the benchmark specification fixes shear deformation off and the section "
    "beta angle at zero in both programs, and because the OpenSeesPy and ETABS centerline models "
    "agree to 0.00% on the affected metrics, the residual is a commercial-side stiffening effect "
    "rather than a solver error. Enabling shear deformation increases the Story 1 drift by 5.4% "
    "in the wrong direction, which rules it out. Two single-parameter mechanisms then bracket the "
    "gap without uniquely identifying it. Introducing rigid end zones at the beam-column joints "
    "reduces the drift monotonically, from 4.1% above the Midas Gen value at the centerline limit "
    "(no rigid zone) to 4.0% below it at the full geometric panel-zone limit; a rigid-zone factor "
    "of 0.51 — consistent with, but not independently confirmed as, the Midas Gen panel-zone "
    "default — reproduces the Midas Story 1 drift and story-1 displacement to within 0.1%, but "
    "leaves the roof displacement about 2.8% below the Midas value. A uniform 6.6% increase in "
    "effective column stiffness instead reproduces the Story 1 drift and the roof displacement to "
    "within 0.1% but implies a different story-2 drift. Neither single-parameter surrogate "
    "reproduces the full displacement field — a true geometric panel zone would be required to "
    "match the story and roof responses simultaneously — but the two bracket the residual (all "
    "affected metrics remain below 4%) and are consistent in sign and magnitude with "
    "beam-column end-zone stiffening that the centerline benchmark models deliberately omit. "
    "Global equilibrium is preserved throughout (base reactions agree to about 0.1%), so the "
    "Case 4 differences reflect a commercial-side modeling convention rather than an error in "
    "the open-source analysis chain."
)

LOC_P114 = "Table 7. Case 4 discrepancy ablation."
NEW_P114 = (
    "Table 7. Case 4 discrepancy bound. Story Drift 1 (Midas Gen reference = 0.000640); the "
    "ETABS centerline model returns the OpenSeesPy baseline values (0.00% difference). A 0.51 "
    "rigid-zone factor (consistent with, but not independently confirmed as, the Midas Gen "
    "panel-zone default) reproduces the Story 1 drift to within 0.1% but leaves the roof about "
    "2.8% low, whereas a uniform +6.6% effective column stiffness matches the Story 1 drift and "
    "the roof but implies a different story-2 drift; the two bracket the gap non-uniquely and "
    "neither reproduces the full displacement field. Shear deformation increases the drift "
    "(wrong direction). The “vs Midas” column reports the signed difference relative to "
    "the Midas Gen value, so the Case 4 baseline reads +4.1% here versus the 3.90% symmetric "
    "maximum in Table 6."
)

LOC_P116 = "Figure 4. Case 4 discrepancy attribution:"
NEW_P116 = (
    "Figure 4. Case 4 discrepancy bound: Story 1 drift versus the beam-column rigid-zone "
    "(end-offset) factor. The centerline model (factor 0; OpenSeesPy ≡ ETABS) lies +4.1% "
    "above the Midas Gen reference and the full geometric panel zone (factor 1) lies 4.0% below "
    "it; the Story 1 drift is matched at a rigid-zone factor of 0.51, which does not "
    "simultaneously reproduce the roof displacement (see text)."
)

OLD_T7_CELL = "equivalent +6.6% column stiffness"
NEW_T7_CELL = "alternative: +6.6% column stiffness"

SENTINEL = "Two single-parameter mechanisms then bracket the gap without uniquely identifying it"


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:50]!r}")
    return hits[0]


def set_text(par, new, tag):
    xml = par._p.xml
    if "fldChar" in xml or "instrText" in xml:
        raise SystemExit(f"ABORT [{tag}]: paragraph carries a field; refusing rewrite")
    if not par.runs:
        par.add_run(new)
        return
    par.runs[0].text = new
    for r in par.runs[1:]:
        r.text = ""


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs
    if any(SENTINEL in p.text for p in paras):
        raise SystemExit("ABORT: already applied (P113 sentinel present).")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    print(f"BEFORE: paragraphs={np0} tables={nt0} inline_shapes={ni0} fldChar={fld0} instrText={ins0}")

    for needle, new, tag in [(LOC_P113, NEW_P113, "P113"), (LOC_P114, NEW_P114, "P114"),
                             (LOC_P116, NEW_P116, "P116")]:
        set_text(find_para(paras, needle, tag), new, tag)

    # T7R5 row-label cell
    hits = [(ri, ci, c) for ti, t in enumerate(d.tables) for ri, row in enumerate(t.rows)
            for ci, c in enumerate(row.cells) if OLD_T7_CELL in c.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [T7R5]: matched {len(hits)} cells (want 1)")
    ri, ci, cell = hits[0]
    if "fldChar" in cell._tc.xml or "instrText" in cell._tc.xml:
        raise SystemExit("ABORT [T7R5]: cell carries a field")
    done = False
    for cpar in cell.paragraphs:
        for r in cpar.runs:
            if OLD_T7_CELL in r.text:
                r.text = r.text.replace(OLD_T7_CELL, NEW_T7_CELL)
                done = True
    if not done:
        raise SystemExit("ABORT [T7R5]: label not found within a single run")

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    print(f"AFTER : paragraphs={np1} tables={nt1} inline_shapes={ni1} fldChar={fld1} instrText={ins1}")
    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0} -> {fld1}/{ins1}"
    assert (np1, nt1, ni1) == (np0, nt0, ni0), "structure count changed"

    print(f"\n--- P113 ({len(find_para(d.paragraphs, SENTINEL, 'x').text.split())} words) ---\n{find_para(d.paragraphs, SENTINEL, 'x').text}")
    print(f"\n--- P114 ---\n{find_para(d.paragraphs, 'Table 7. Case 4 discrepancy bound.', 'x').text}")
    print(f"\n--- P116 ---\n{find_para(d.paragraphs, 'Figure 4. Case 4 discrepancy bound:', 'x').text}")
    print(f"\n--- T7 row label now ---\n{cell.text}")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
