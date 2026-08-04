"""Session 2026-07-02 (MAJOR REVISION): C2 — rewrite P161 to match the shipped code.

P161 (§6.2 third limitation) currently states "compact sections with full lateral
support, K = 1.0, and no consideration of lateral-torsional buckling" — the OPPOSITE of
design_check.py. Verified in design_check.py:
  - K_MOMENT_FRAME = 1.2 (moment/unknown), K_BRACED = 1.0 (braced/shear wall)  [:176-191]
  - _classify_compactness (AISC B4) + _flb_capacity (F3.2) for I-sections        [:927-928]
  - _ltb_bending_capacity (AISC F2) applied when an unbraced length is supplied;
    default = slab-continuous lateral support -> Mp                              [:953-956]
  - column strong-axis bending -> Mp (assumes floor-level lateral support)       [:943-947]
And in example_section4_results.json the critical column carries "K": 1.2, compact,
interaction 0.4666 (~0.467) with slenderness 48.0 = 1.2*3000/75.1 -> the reported 0.467
is the shipped-default (K=1.2) value; no number changes needed elsewhere.

Rewrite keeps only the TRUE caveats: no IFC-driven automatic unbraced-length detection;
screening not certification; module not independently benchmarked. P161 is runs=3, no
field. Backup -> ..._v2.pre_c2p161_2026-07-02.docx. Dry-run default; --apply.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_c2p161_2026-07-02.docx")

LOC_P161 = "the preliminary design screening module adopts several simplifying assumptions"
NEW_P161 = (
    "Third, the preliminary design screening module is a screening tool rather than a "
    "certification engine, and its modeling assumptions bound the scope of its outcomes. Member "
    "strength follows the AISC 360-22 provisions: column compression uses an effective length "
    "factor of K = 1.2 for moment (and unclassified) frames and K = 1.0 for braced frames and "
    "shear walls — the conservative sway lower bound, adopted because K = 1.0 is unconservative "
    "for moment frames; flexural capacity applies compactness classification (B4) and "
    "flange-local-buckling reduction (F3.2) for doubly-symmetric I-sections, and "
    "lateral-torsional buckling (F2) is evaluated for beams whenever an unbraced length is "
    "supplied. By default, beam top flanges are treated as continuously braced by the floor "
    "slab, so the lateral-torsional-buckling reduction is triggered only when an unbraced length "
    "is provided explicitly; the workflow does not yet infer beam unbraced lengths automatically "
    "from the IFC model, and column strong-axis bending assumes lateral support at the floor "
    "levels. These assumptions are adopted for preliminary screening and do not replace a full "
    "code-compliance review, and the screening module has not been independently benchmarked "
    "against manual calculations or a separate reference program."
)

SENTINEL = "screening tool rather than a certification engine"


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:50]!r}")
    return hits[0]


def set_text(par, new, tag):
    xml = par._p.xml
    if "fldChar" in xml or "instrText" in xml:
        raise SystemExit(f"ABORT [{tag}]: paragraph carries a field; refusing rewrite")
    if not par.runs:
        par.add_run(new)
        return
    par.runs[0].text = new
    for r in par.runs[1:]:
        r.text = ""


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs
    if any(SENTINEL in p.text for p in paras):
        raise SystemExit("ABORT: already applied (P161 sentinel present).")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    print(f"BEFORE: paragraphs={np0} tables={nt0} inline_shapes={ni0} fldChar={fld0} instrText={ins0}")

    p161 = find_para(paras, LOC_P161, "P161")
    set_text(p161, NEW_P161, "P161")

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    print(f"AFTER : paragraphs={np1} tables={nt1} inline_shapes={ni1} fldChar={fld1} instrText={ins1}")
    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0} -> {fld1}/{ins1}"
    assert (np1, nt1, ni1) == (np0, nt0, ni0), "structure count changed"

    print(f"\n--- P161 now ({len(p161.text.split())} words) ---\n{p161.text}")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
