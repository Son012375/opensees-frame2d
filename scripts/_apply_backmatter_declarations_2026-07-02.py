"""Insert Elsevier mandatory back-matter declarations before REFERENCES (P175).

Blocks (Funding intentionally OMITTED per author instruction — still required before
submission): CRediT -> Declaration of competing interest -> Declaration of generative AI
-> Data availability. All plain-text paragraphs cloned from a Normal-style template
(no citation/SEQ fields touched). Titles set bold.

Author inputs (2026-07-02):
- CRediT: Son = comprehensive research; Baek = review, supervision, direction-setting.
- GenAI: Claude (Anthropic) used for broad manuscript editing -> disclosed.
- License: MIT (author: "anything is fine").
- Repo: https://github.com/Son012375/opensees-frame2d (PUBLIC). NOTE: repo must be
  updated (push feat/etabs-benchmark work + add LICENSE + tag + Zenodo DOI) before the
  statement is fully truthful; DOI left as a placeholder.

Backup -> ..._v2.pre_backmatter_2026-07-02.docx. Dry-run default; --apply saves.
"""
from __future__ import annotations

import copy
import shutil
import sys
from pathlib import Path

import docx
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_backmatter_2026-07-02.docx")

EN = "–"  # –
EM = "—"  # —

BLOCKS = [
    ("title", "CRediT authorship contribution statement"),
    ("body",  f"Myeong-Guk Son: Conceptualization, Methodology, Software, Validation, "
              f"Formal analysis, Investigation, Data curation, Writing {EN} original draft, "
              f"Visualization."),
    ("body",  f"Jang-Woon Baek: Conceptualization, Supervision, Writing {EN} review & editing, "
              f"Project administration."),
    ("title", "Declaration of competing interest"),
    ("body",  "The authors declare that they have no known competing financial interests or "
              "personal relationships that could have appeared to influence the work reported "
              "in this paper."),
    ("title", "Declaration of generative AI and AI-assisted technologies in the writing process"),
    ("body",  "During the preparation of this work the authors used Claude (Anthropic) to "
              "assist with editing, language refinement, and formatting revisions of the "
              "manuscript. After using this tool, the authors reviewed and edited the content "
              "as needed and take full responsibility for the content of the publication."),
    ("title", "Data availability"),
    ("body",  f"The source code of the open-source pipeline {EM} the node-element IFC parser, "
              f"the KDS load-generation module, and the OpenSeesPy three-dimensional analysis "
              f"driver {EM} together with the IFC application-example model, the five benchmark "
              f"case definitions, and the OpenSeesPy{EN}Midas Gen{EN}ETABS comparison data, are "
              f"openly available at https://github.com/Son012375/opensees-frame2d, released "
              f"under the MIT license. A versioned snapshot archived at Zenodo (DOI: to be "
              f"inserted upon acceptance) reproduces the results reported in this paper."),
]

SENTINEL = "Declaration of competing interest"


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def has_field(p):
    from docx.oxml.ns import qn
    return any(el.tag in (qn("w:fldChar"), qn("w:instrText")) for el in p._p.iter())


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    if any(SENTINEL in p.text for p in paras):
        raise SystemExit("ABORT: already applied (competing-interest present).")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    print(f"BEFORE: paragraphs={np0} tables={nt0} images={ni0} fldChar={fld0} instrText={ins0}")

    # anchor: REFERENCES heading (Heading 1, exact text 'REFERENCES')
    refs = [p for p in paras if p.text.strip() == "REFERENCES"]
    if len(refs) != 1:
        raise SystemExit(f"ABORT: REFERENCES anchor matched {len(refs)}")
    p_refs = refs[0]

    # template: a plain Normal-style APPENDIX body paragraph (NOT the title P000).
    # Anchor on the Appendix A.1 scope paragraph, which is regular body prose in Normal style.
    tmpl = None
    for p in paras:
        if ("This appendix records the clause-by-clause" in p.text
                and p.style and p.style.name == "Normal" and len(p.runs) >= 1 and not has_field(p)):
            tmpl = p
            break
    if tmpl is None:
        raise SystemExit("ABORT: appendix Normal template paragraph not found")
    print(f"template = Normal body para: {tmpl.text[:40]!r} (runs={len(tmpl.runs)}, "
          f"bold0={tmpl.runs[0].font.bold}, size0={tmpl.runs[0].font.size})")

    added = 0
    for kind, text in BLOCKS:
        new_el = copy.deepcopy(tmpl._p)
        p_refs._p.addprevious(new_el)
        para = Paragraph(new_el, p_refs._parent)
        # collapse to a single run carrying our text
        runs = para.runs
        runs[0].text = text
        runs[0].font.bold = True if kind == "title" else False
        runs[0].font.italic = False
        for r in runs[1:]:
            r.text = ""
        added += 1

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    print(f"AFTER : paragraphs={np1} tables={nt1} images={ni1} fldChar={fld1} instrText={ins1} (+{added} paras)")

    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0} -> {fld1}/{ins1}"
    assert np1 == np0 + len(BLOCKS), f"para delta {np1-np0} != {len(BLOCKS)}"
    assert (nt1, ni1) == (nt0, ni0), "table/image count changed"

    # show inserted block in order
    print("\n--- inserted back-matter (in document order) ---")
    started = False
    for p in d.paragraphs:
        if p.text.strip() == "CRediT authorship contribution statement":
            started = True
        if started:
            if p.text.strip() == "REFERENCES":
                break
            b = "[B]" if any(r.font.bold for r in p.runs if r.text) else "   "
            print(f"  {b} {p.text[:110]}")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
