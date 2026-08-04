"""Session 2026-07-02 (MAJOR REVISION): C1 body corrections.

Three edits that make the body say what the validation files say (C1 decisive):

  P085 (§4.1)  APPEND one principled definition sentence: benchmark = independently
               modeled matched-input case vs a commercial program; cross-check =
               identical generated loads replayed to an independently assembled model,
               isolating assembly/recovery from load generation.
  P132 (§5.2)  REWRITE the ETABS paragraph: it is a solver-equivalence check on the
               assembled node-element model under identical loads — NOT a test of the
               IFC parse or the KDS load generation (those are evidenced separately:
               parse by counts/diagnostics §3.3, loads by hand-check Appendix A). Fixes
               "cross-check of the IFC-to-analysis pipeline" / "pipeline itself" overclaim.
  T0R9 (Table1) last cell: "Midas Gen (112 metrics) + ETABS" ->
               "Midas Gen benchmark (112 metrics) + ETABS cross-check (31)".

Grounded in validation/ifc_example_etabs_compare.json (_meta.model = node-element export;
identical OpenSees-generated loads; 31/31 OK <1%) and e2_ifc_etabs_scoping.md (physical
.ifc not in repo -> fallback node-element path). No field/SEQ/citation touched.
Backup -> ..._v2.pre_c1body_2026-07-02.docx. Dry-run default; --apply.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_c1body_2026-07-02.docx")

# ---- P085 append (principled definition) ----
LOC_P085 = "This isolates the comparison to the numerical modeling and response-recovery behavior"
P085_APPEND = (
    " Throughout this paper the term benchmark denotes such an independently modeled, "
    "matched-input comparison against a commercial program, whereas cross-check (as in the "
    "ETABS comparison of Section 5.2) denotes a comparison in which identical generated loads "
    "are replayed to an independently assembled model, isolating model assembly and response "
    "recovery from load generation."
)

# ---- P132 wholesale rewrite (§5.2 ETABS scope) ----
LOC_P132 = "the same 87-member example was rebuilt in ETABS"
NEW_P132 = (
    "As an independent cross-check of model assembly and response recovery, the same "
    "87-member example was rebuilt in ETABS and driven with the identical KDS-generated member "
    "and nodal loads. Because both models receive the same generated loads on the same geometry, "
    "the comparison isolates model assembly and linear response recovery from the IFC parse and "
    "the load generation — a solver-equivalence check rather than a test of those two stages, "
    "which are evidenced separately: the parse by the extracted node, element, and member counts "
    "and validation diagnostics of Section 3.3, and the generated load values by the "
    "clause-by-clause hand-check of Appendix A. All 31 compared global-response metrics agree to "
    "within 1%: the dead-load corner reactions match to 0.94%, the equivalent-static story drifts "
    "and base shears to within 0.01%, and the governing displacement envelope to within 0.04% "
    "(both models using E = 205,000 MPa for SS275). This confirms that the assembled node-element "
    "model and its recovered global response reproduce a second commercial reference, "
    "complementing the controlled five-case Midas Gen benchmark of Section 4, which remains the "
    "primary numerical validation; end-to-end validation of the IFC parse against a commercial "
    "model is identified as future work in Section 6.3."
)

# ---- Table 1 T0R9 last cell ----
OLD_CELL = "Midas Gen (112 metrics) + ETABS"
NEW_CELL = "Midas Gen benchmark (112 metrics) + ETABS cross-check (31)"

SENTINEL = "a solver-equivalence check rather than a test of those two stages"


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:50]!r}")
    return hits[0]


def set_text(par, new, tag):
    xml = par._p.xml
    if "fldChar" in xml or "instrText" in xml:
        raise SystemExit(f"ABORT [{tag}]: paragraph carries a field; refusing rewrite")
    if not par.runs:
        par.add_run(new)
        return
    par.runs[0].text = new
    for r in par.runs[1:]:
        r.text = ""


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    if any(SENTINEL in p.text for p in paras):
        raise SystemExit("ABORT: already applied (P132 sentinel present).")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    print(f"BEFORE: paragraphs={np0} tables={nt0} inline_shapes={ni0} fldChar={fld0} instrText={ins0}")

    # P085 append
    p085 = find_para(paras, LOC_P085, "P085")
    p085.add_run(P085_APPEND)

    # P132 rewrite
    p132 = find_para(paras, LOC_P132, "P132")
    set_text(p132, NEW_P132, "P132")

    # T0R9 cell
    t1 = d.tables[0]
    hits = [(ri, ci, c) for ri, row in enumerate(t1.rows)
            for ci, c in enumerate(row.cells) if OLD_CELL in c.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [T0R9]: matched {len(hits)} cells (want 1)")
    ri, ci, cell = hits[0]
    cxml = cell._tc.xml
    if "fldChar" in cxml or "instrText" in cxml:
        raise SystemExit("ABORT [T0R9]: cell carries a field")
    cpar = cell.paragraphs[0]
    if not cpar.runs:
        cpar.add_run(NEW_CELL)
    else:
        cpar.runs[0].text = NEW_CELL
        for r in cpar.runs[1:]:
            r.text = ""

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    print(f"AFTER : paragraphs={np1} tables={nt1} inline_shapes={ni1} fldChar={fld1} instrText={ins1}")
    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0} -> {fld1}/{ins1}"
    assert (np1, nt1, ni1) == (np0, nt0, ni0), "structure count changed"

    print(f"\n--- P085 tail ---\n{p085.text[-360:]}")
    print(f"\n--- P132 now ({len(p132.text.split())} words) ---\n{p132.text}")
    print(f"\n--- T0R9 cell[{ri},{ci}] now ---\n{cell.text}")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
