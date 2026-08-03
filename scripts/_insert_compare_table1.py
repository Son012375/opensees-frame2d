"""Part A step 2 - insert the Related-Work comparison table as Table 1 (plain Word
grid, no fields) immediately after the §2.4 prose paragraph ('Table 1 positions ...').
The table sequence was already opened to start at Table 1 by _apply_table_renumber.py.

Backup -> ..._v2.pre_table1_backup.docx.  Run with --apply (default dry-run).
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_table1_backup.docx")

CAPTION = "Table 1. Positioning of the present work against representative recent systems."

TABLE1 = [
    ("System", "Primary input", "Representation", "Connectivity repair / diagnostics", "Solver", "Code system", "Commercial cross-validation"),
    ("Ramaji & Memari (2018)", "IFC (coordination view)", "analytical model", "not surfaced", "external", "—", "none"),
    ("Hasan et al. (2019)", "BIM", "geometrically accurate FE model", "partial", "BIM-centered", "—", "qualitative"),
    ("Leonardi et al. (2024)", "openBIM / IFC", "FE (masonry aggregates)", "not surfaced", "OpenSees-based", "Eurocode context", "feasibility-level"),
    ("Singh et al. (2024)", "architectural IFC", "structural analytical model", "not surfaced", "export to FE", "—", "none"),
    ("Rudenko & Petryna (2025)", "BIM", "multi-complexity FE", "not surfaced", "generic FE", "—", "none"),
    ("Llanos & Delgadillo (2025)", "parametric", "OpenSeesPy RC frame", "not surfaced", "OpenSeesPy", "—", "none"),
    ("Liang et al. (2025)", "natural language", "LLM-generated script", "not surfaced", "OpenSeesPy", "—", "20 analytical problems"),
    ("MCP approach (Buildings, 2025)", "natural language", "LLM + MCP commands", "not surfaced", "OpenSeesPy / ETABS", "—", "ETABS (~1%)"),
    ("This work", "IFC + NL + manual", "node-element analysis graph", "yes (merge / snap / split + diagnostics)", "OpenSeesPy", "KDS (clause-traced)", "Midas Gen + ETABS, 112 metrics"),
]

ANCHOR = "Table 1 positions the present work against representative recent systems"


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))

    if any("Positioning of the present work against representative recent" in p.text
           for p in d.paragraphs):
        print("ABORT: comparison-table caption already present.")
        return 1

    anchors = [p for p in d.paragraphs if ANCHOR in p.text]
    if len(anchors) != 1:
        print(f"ABORT: anchor matched {len(anchors)} (want 1).")
        return 1
    anchor = anchors[0]

    cap = d.add_paragraph(style="Caption")
    cap.add_run(CAPTION)

    tbl = d.add_table(rows=len(TABLE1), cols=7)
    tbl.style = "Table Grid"
    for ri, row in enumerate(TABLE1):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.text = val
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(8)
                    if ri == 0 or ri == len(TABLE1) - 1:
                        run.font.bold = True

    # relocate caption then table to just after the §2.4 prose paragraph
    anchor._p.addnext(cap._p)
    cap._p.addnext(tbl._tbl)

    if not apply:
        print(f"DRY-RUN: would insert caption + {len(TABLE1)}x7 table after {anchor.text[:40]!r}")
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"backup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
