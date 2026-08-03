"""Polish edits A+B+C for DBE submission (plain runs / table cells only; no fields touched).

A. M4 Korean glyphs -> English + keyword trim (8 -> 6):
   '표 ' -> 'Table ' (16x), '식 ' -> 'Eq. ' (6x) across body paragraphs and table cells;
   'retail (판매장)'->'retail', 'office (사무실)'->'office',
   'S (단주기)'->'S (short-period)', 'S₁ (1초 주기)'->'S₁ (1-second period)'.
   Keywords: drop title-duplicative 'Midas Gen', 'node-element model', 'steel frame buildings';
   keep 6 incl. new 'benchmark validation'.
   (P176/'Academic References' orphan deletion DEFERRED to Word: doc has a live ZOTERO_BIBL
    bibliography field; deleting adjacent EndNote-styled paragraphs risks corrupting it.)

B. M2 Abstract: merge 3 paragraphs (P007-P009) into ONE <=250-word paragraph, preserving the
   ETABS 31/31 cross-check and 112-metric (100<1%/12<4%) figures.

C. M3 + eq wording:
   - P115: soften "essentially the default panel-zone setting of Midas Gen" (unsourced/curve-fit).
   - Table 7 caption: same softening ("Midas Gen's default panel-zone setting").
   - P094: fix the relative-difference definition ("x_Midas appears in the numerator" is wrong;
     numerator is the absolute difference) and define both operands.

Backup -> ..._v2.pre_polish_2026-07-02.docx. Dry-run default; --apply saves.
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_polish_2026-07-02.docx")

KW_OLD = ("IFC; OpenSeesPy; Midas Gen; Korean Design Standards; node-element model; "
          "open-source structural analysis; BIM-to-analysis workflow; steel frame buildings")
KW_NEW = ("IFC; OpenSeesPy; Korean Design Standards; BIM-to-analysis workflow; "
          "open-source structural analysis; benchmark validation")

TERMS = [
    ("retail (판매장)", "retail"),
    ("office (사무실)", "office"),
    ("S (단주기)", "S (short-period)"),
    ("S₁ (1초 주기)", "S₁ (1-second period)"),
]

ABSTRACT = (
    "Commercial structural-analysis programs remain central to building design practice, yet "
    "their closed preprocessing and result-recovery pipelines can limit independent inspection, "
    "reproduction, and customization of the analysis path. This paper presents an open-source "
    "pipeline linking node-element IFC parsing, Korean Design Standard (KDS) load automation, and "
    "OpenSeesPy three-dimensional frame analysis for steel frame buildings, supporting regular "
    "orthogonal grids and orthogonal irregular plans such as L-shape and setback through "
    "zone-based decomposition. IFC parsing produces an inspectable analysis graph of nodes, "
    "elements, supports, sections, materials, and validation diagnostics, and KDS automation "
    "generates dead, live, seismic, and wind cases and assembles code combinations, with every "
    "value traced to its originating clause in a hand-check appendix and four workflow-introduced "
    "simplifications disclosed. The pipeline is benchmarked against Midas Gen across five "
    "controlled cases comparing 112 response metrics: 100 agree within 1% and the remaining 12 "
    "within 4% once sign and local-axis conventions are aligned, with the larger discrepancies "
    "confined to element-formulation-sensitive quantities of a single three-dimensional case "
    "rather than global equilibrium. An IFC-derived three-story application example and a "
    "node-element L-shaped demonstration exercise the full pipeline end-to-end under "
    "KDS-generated loading; the application example's global response is independently "
    "cross-checked against ETABS, agreeing on all 31 compared metrics within 1%. Within the "
    "benchmark-validated scope — regular steel frames under elastic three-dimensional analysis — "
    "the pipeline serves as a credible, transparent, KDS-traced alternative computational path "
    "alongside the commercial baseline; orthogonal irregular plans are supported and demonstrated "
    "as workflow-execution cases rather than benchmark-validated cases."
)

P115_OLD = "a rigid-zone factor of 0.51, essentially the default panel-zone setting of Midas Gen, reproduces the Midas Story 1 drift and story-1 displacement to within 0.1%"
P115_NEW = "a rigid-zone factor of 0.51, consistent with but not independently confirmed as the panel-zone default of Midas Gen, reproduces the Midas Story 1 drift and story-1 displacement to within 0.1%"

T7_OLD = "A 0.51 rigid-zone factor — Midas Gen's default panel-zone setting — reproduces Midas to within 0.1%"
T7_NEW = "A 0.51 rigid-zone factor (consistent with, but not independently confirmed as, the Midas Gen panel-zone default) reproduces Midas to within 0.1%"

P094_OLD = "where x_Midas appears in the numerator and the denominator is the larger of the two program magnitudes, so the reported relative difference is symmetric and does not privilege either program."
P094_NEW = "where x_OpenSees and x_Midas are the corresponding response magnitudes from the OpenSeesPy and Midas Gen models; the numerator is their absolute difference and the denominator is the larger of the two magnitudes, so the reported relative difference is symmetric and does not privilege either program."

HANGUL = re.compile(r'[가-힣]')


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:45]!r}")
    return hits[0]


def repl_in_para(par, old, new, tag):
    for r in par.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return
    raise SystemExit(f"ABORT [{tag}]: not found in single run: {old[:45]!r}")


def iter_all_paragraphs(d):
    yield from d.paragraphs
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                yield from c.paragraphs


def glyph_replace_runs(d):
    """Hangul '표'/'식' sit in isolated runs (surrounded by ' '); substring-replace per run.
    '표'->'Table', '식'->'Eq.' — adjacent runs already carry the surrounding spaces."""
    n_pyo = n_sik = 0
    for p in iter_all_paragraphs(d):
        for r in p.runs:
            t = r.text
            if "표" in t:
                n_pyo += t.count("표"); t = t.replace("표", "Table")
            if "식" in t:
                n_sik += t.count("식"); t = t.replace("식", "Eq.")
            if t != r.text:
                r.text = t
    return n_pyo, n_sik


# occupancy / period cells: whole-cell rewrite (Korean tokens live in their own runs w/ split parens)
CELL_REWRITES = [
    ("LL, retail (판매장)", "LL, retail"),
    ("LL, office (사무실)", "LL, office"),
    ("S (단주기)", "S (short-period)"),
    ("S₁ (1초 주기)", "S₁ (1-second period)"),
]


def term_replace(d):
    n = 0
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                for para in c.paragraphs:
                    cur = para.text
                    for old, new in CELL_REWRITES:
                        if cur.strip() == old:
                            para.runs[0].text = new
                            for r in para.runs[1:]:
                                r.text = ""
                            n += 1
    return n


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    # guard: not already applied (old abstract P008 still present)
    if not any("Node-element IFC parsing follows the methodology established in prior BIM-to-FEM work and produces an inspectable analysis graph containing" in p.text for p in paras):
        raise SystemExit("ABORT: abstract already merged (guard).")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    h0 = sum(len(HANGUL.findall(p.text)) for p in iter_all_paragraphs(d))
    print(f"BEFORE: paragraphs={np0} tables={nt0} images={ni0} fldChar={fld0} instrText={ins0} hangul={h0}")

    # ---- A: keywords ----
    kw = find_para(paras, "open-source structural analysis; BIM-to-analysis workflow", "keywords")
    repl_in_para(kw, KW_OLD, KW_NEW, "keywords")

    # ---- A: Korean glyphs + terms ----
    npyo, nsik = glyph_replace_runs(d)
    nterm = term_replace(d)
    print(f"  glyphs: 표->Table x{npyo}, 식->Eq. x{nsik}, terms x{nterm}")

    # ---- C: panel-zone softening + equation wording ----
    p115 = find_para(paras, "To attribute this discrepancy rather than leave it qualitative", "P115")
    repl_in_para(p115, P115_OLD, P115_NEW, "P115")
    t7cap = find_para(paras, "Table 7. Case 4 discrepancy ablation", "T7cap")
    repl_in_para(t7cap, T7_OLD, T7_NEW, "T7cap")
    # P094 spans 3 runs: r0='where ', r1='x_Midas', r2=' appears...program. Each metric...'
    p094 = find_para(paras, "so the reported relative difference is symmetric", "P094")
    r = p094.runs
    assert r[0].text == "where " and r[1].text == "x_Midas", f"P094 runs unexpected: {[x.text for x in r[:3]]}"
    assert "program. Each metric is then assigned" in r[2].text, "P094 r2 tail unexpected"
    tail = r[2].text.split("program.", 1)[1]   # ' Each metric is then assigned...'
    r[0].text = P094_NEW                         # ends with '...either program.'
    r[1].text = ""
    r[2].text = tail

    # ---- B: abstract merge (do last; deletes 2 paragraphs) ----
    p007 = find_para(paras, "Commercial structural-analysis programs remain central", "P007")
    p008 = find_para(paras, "produces an inspectable analysis graph containing nodes, line elements", "P008")
    p009 = find_para(paras, "Of the 112 metrics, 100 agree with Midas Gen within 1%", "P009")
    # set P007 to full abstract (single run), blank extras
    p007.runs[0].text = ABSTRACT
    for r in p007.runs[1:]:
        r.text = ""
    # delete P008, P009 elements
    for p in (p008, p009):
        p._p.getparent().remove(p._p)

    wc = len(ABSTRACT.split())
    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    h1 = sum(len(HANGUL.findall(p.text)) for p in iter_all_paragraphs(d))
    print(f"AFTER : paragraphs={np1} tables={nt1} images={ni1} fldChar={fld1} instrText={ins1} hangul={h1}")
    print(f"  abstract word count = {wc}")

    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0} -> {fld1}/{ins1}"
    assert np1 == np0 - 2, f"para delta {np1-np0} != -2"
    assert (nt1, ni1) == (nt0, ni0), "table/image count changed"
    assert h1 == 0, f"Hangul remaining: {h1}"
    assert wc <= 252, f"abstract too long: {wc} words"

    print("\n--- new keywords ---\n ", kw.text)
    print("\n--- new abstract ---\n ", p007.text)
    print("\n--- P115 tail ---\n ", p115.text[-300:])
    print("\n--- T7 caption ---\n ", t7cap.text)
    print("\n--- P094 ---\n ", p094.text)

    if not apply:
        print("\nDRY-RUN (no save).")
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
