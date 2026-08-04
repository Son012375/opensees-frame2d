"""C2: re-embed the refreshed 3D deformed-shape figure (Fig 5) into the canonical v2
DOCX with the current post-gravity-fix values (peak nodal displacement ~12.0 mm).

Target = inline shape index 4 (/word/media/image5.png), the deformed shape shown as
manuscript Figure 5. Regenerate the PNG first with:
    python scripts/_render_fig4_deformed.py
then run this. Keeps display width, recomputes height from aspect. Refreshes
final/fig4.png too. Backup -> ..._v2.pre_fig5_2026-07-01.docx. --apply to save.
"""
from __future__ import annotations

import shutil
import struct
import sys
from pathlib import Path

import docx
from docx.shared import Emu

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs" / "paper1_open_source_alternative"
DOC = BASE / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_fig5_2026-07-01.docx")
NEW_PNG = BASE / "figures" / "fig4_ifc_deformed.png"
FINAL_PNG = BASE / "figures" / "final" / "fig4.png"

SHAPE_IDX = 4              # deformed shape = manuscript Fig 5 = inline shape #5
OLD_BYTES_LEN = 378266    # known length of the pre-refresh embedded deformed image


def png_size(b: bytes) -> tuple[int, int]:
    assert b[12:16] == b"IHDR", "not a PNG / unexpected header"
    w, h = struct.unpack(">II", b[16:24])
    return w, h


def main(apply: bool) -> int:
    new_bytes = NEW_PNG.read_bytes()
    nw, nh = png_size(new_bytes)
    print(f"new deformed png: {NEW_PNG.name} {nw}x{nh}px {len(new_bytes)} bytes")

    d = docx.Document(str(DOC))
    shapes = d.inline_shapes
    if len(shapes) != 8:
        raise SystemExit(f"ABORT: expected 8 inline shapes, got {len(shapes)}")
    sh = shapes[SHAPE_IDX]
    rId = sh._inline.graphic.graphicData.pic.blipFill.blip.embed
    part = d.part.related_parts[rId]
    cur = part._blob
    cw, ch = png_size(cur)
    print(f"target shape[{SHAPE_IDX}]: rId={rId} partname={part.partname} {cw}x{ch}px {len(cur)} bytes")
    if len(cur) != OLD_BYTES_LEN:
        raise SystemExit(
            f"ABORT: shape[{SHAPE_IDX}] current bytes {len(cur)} != expected {OLD_BYTES_LEN}; not touching.")

    old_w_emu, old_h_emu = int(sh.width), int(sh.height)
    new_h_emu = int(old_w_emu * nh / nw)
    print(f"display: width={old_w_emu/914400:.2f}in kept; height {old_h_emu/914400:.2f}->{new_h_emu/914400:.2f}in")

    if not apply:
        print("\nDRY-RUN. Re-run with --apply.")
        return 0

    shutil.copy2(NEW_PNG, FINAL_PNG)
    print(f"refreshed {FINAL_PNG.relative_to(ROOT)}")
    shutil.copy2(DOC, BACKUP)
    print(f"backup -> {BACKUP.name}")

    part._blob = new_bytes
    sh.height = Emu(new_h_emu)
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")

    d2 = docx.Document(str(DOC))
    media = [p.partname for p in d2.part.package.iter_parts() if "media" in str(p.partname)]
    swapped = d2.part.related_parts[d2.inline_shapes[SHAPE_IDX]._inline.graphic.graphicData.pic.blipFill.blip.embed]._blob
    print(f"[verify] inline_shapes={len(d2.inline_shapes)} media={len(media)} "
          f"paragraphs={len(d2.paragraphs)} tables={len(d2.tables)} shape[{SHAPE_IDX}]_bytes={len(swapped)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
