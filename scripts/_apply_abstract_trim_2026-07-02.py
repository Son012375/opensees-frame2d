"""Session 2026-07-02: trim the abstract to <=250 words (DBE guideline) — wording only,
no number or claim changed. Wholesale field-safe rewrite of P007.
Backup -> ..._v2.pre_abstrim_2026-07-02.docx. Dry-run default; --apply.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_abstrim_2026-07-02.docx")

LOC = "Two evidence standards are distinguished: a benchmark (an independently modeled"
NEW_ABSTRACT = (
    "Commercial structural-analysis programs remain central to building design, yet their "
    "closed preprocessing and result recovery limit independent inspection, reproduction, and "
    "customization. This paper presents an open-source alternative for regular steel frame "
    "buildings that couples clause-traced Korean Design Standard (KDS) load automation with "
    "OpenSeesPy three-dimensional analysis, evaluated against commercial software at the "
    "response-quantity level. A node-element layer converts BIM/IFC input into an inspectable "
    "analysis graph; KDS automation then generates dead, live, seismic, wind, and snow cases "
    "and assembles code combinations, every value traced to its originating clause in a "
    "hand-check appendix and four workflow-introduced simplifications disclosed. Two evidence "
    "standards are distinguished: a benchmark (an independently modeled, matched-input case "
    "compared against Midas Gen) and a cross-check (identical generated loads replayed to an "
    "independently assembled model, isolating assembly and recovery). Across five Midas Gen "
    "benchmark cases, 100 of 112 response metrics agree within 1% and the remaining 12 within "
    "4% once sign and local-axis conventions are aligned, the larger differences confined to "
    "element-formulation-sensitive quantities of a single symmetric three-dimensional case, not "
    "global equilibrium. An IFC-derived example is cross-checked against ETABS — a "
    "solver-equivalence check on the assembled model under identical loads, not a test of the "
    "parse or load generation — agreeing on 31 global-response metrics within 1%. Orthogonal "
    "irregular plans are supported through zone-based decomposition and demonstrated end-to-end, "
    "with commercial benchmarking reserved for future work. Within this scope the pipeline is a "
    "credible, transparent, KDS-traced alternative path, and its clause-trace, disclosure, and "
    "hand-check protocol offers a transferable, code-agnostic evidence standard."
)


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs
    hits = [p for p in paras if LOC in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT: matched {len(hits)} abstract paras (want 1)")
    p = hits[0]
    if "originating clause in a hand-check appendix" not in p.text:
        raise SystemExit("ABORT: located paragraph is not the abstract")
    if "closed preprocessing and result recovery limit" in p.text:
        raise SystemExit("ABORT: already trimmed")
    if "fldChar" in p._p.xml or "instrText" in p._p.xml:
        raise SystemExit("ABORT: abstract carries a field")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    wc_old = len(p.text.split())
    p.runs[0].text = NEW_ABSTRACT
    for r in p.runs[1:]:
        r.text = ""
    wc_new = len(p.text.split())

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    print(f"BEFORE: paragraphs={np0} fldChar={fld0} instrText={ins0} abstract_words={wc_old}")
    print(f"AFTER : paragraphs={np1} fldChar={fld1} instrText={ins1} abstract_words={wc_new}")
    assert (fld1, ins1) == (fld0, ins0), "FIELD COUNT CHANGED"
    assert (np1, nt1, ni1) == (np0, nt0, ni0), "structure count changed"
    print(f"\n--- abstract now ({wc_new} words) ---\n{p.text}")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
