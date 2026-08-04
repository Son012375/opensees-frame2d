"""D (author-approved: prose-only + light scoping reconcile) — close reviewer gap G3.

Adds a prose ETABS cross-check paragraph to §5.2 (mirroring the §5.3 P144 cross-check,
no new table / no SEQ caption fields) and lightly reconciles the two scoping sentences
so the section stays internally consistent:

  new  after P133  : "As an independent cross-check of the IFC-to-analysis pipeline ...
                      All 31 compared global-response metrics agree to within 1% ..."
  P134 (§5.2)      : "... rather than additional commercial-baseline benchmark metrics."
                      -> "... independently cross-checked against ETABS as reported above,
                          with the controlled commercial benchmark campaign remaining the
                          five cases of Section 4."
  P153 (§6.1)      : append ", with the global response independently cross-checked
                      against ETABS to within 1%." to the "consistent with engineering
                      expectations ..." sentence.

Numbers from validation/ifc_example_etabs_compare.json + ifc_example_etabs_writeup.md
(31/31 OK, E=205 GPa both, DL corner 0.94%, drifts+base shear <=0.01%, envelope <=0.04%).

Backup -> ..._v2.pre_g3_2026-07-01.docx. Dry-run by default; --apply saves.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_g3_2026-07-01.docx")

NEW_PARA = (
    "As an independent cross-check of the IFC-to-analysis pipeline, the same 87-member "
    "example was rebuilt in ETABS and driven with the identical KDS-generated member and "
    "nodal loads, isolating model assembly and analysis from load generation. All 31 "
    "compared global-response metrics agree to within 1%: the dead-load corner reactions "
    "match to 0.94%, the equivalent-static story drifts and base shears to within 0.01%, "
    "and the governing displacement envelope to within 0.04% (both models using E = "
    "205,000 MPa for SS275). This cross-check confirms the global response of the "
    "IFC-derived node-element pipeline itself against a commercial reference, "
    "complementing the controlled five-case benchmark of Section 4, which remains the "
    "primary numerical validation."
)

P134_OLD = (
    "Accordingly, the reported response quantities are interpreted as workflow "
    "demonstration results rather than additional commercial-baseline benchmark metrics."
)
P134_NEW = (
    "Accordingly, the reported response quantities are interpreted as workflow "
    "demonstration results, independently cross-checked against ETABS as reported above, "
    "with the controlled commercial benchmark campaign remaining the five cases of "
    "Section 4."
)

P153_OLD = (
    "returns drift, displacement, and member-strength results consistent with "
    "engineering expectations under KDS-generated loading."
)
P153_NEW = (
    "returns drift, displacement, and member-strength results consistent with "
    "engineering expectations under KDS-generated loading, with the global response "
    "independently cross-checked against ETABS to within 1%."
)


def _find_one(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} paragraphs (want 1) for {needle[:40]!r}")
    return hits[0]


def _edit_run(par, old, new, tag):
    for r in par.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return
    raise SystemExit(f"ABORT [{tag}]: old substring not in a single run: {old[:40]!r}")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    if any("same 87-member example was rebuilt in ETABS" in p.text for p in paras):
        raise SystemExit("ABORT: G3 cross-check paragraph already present.")

    p133 = _find_one(paras, "screening results across the evaluated load combinations is presented", "P133")
    p134 = _find_one(paras, "rather than additional commercial-baseline benchmark metrics", "P134")
    # note: the same "... consistent with engineering expectations ..." tail also
    # appears in the abstract (P9) and the conclusion (P170, past-tense "returned");
    # anchor on the full present-tense §6.1 sentence so only P153 matches.
    p153 = _find_one(paras, P153_OLD, "P153")

    # reconcile P134 / P153 (plain-run edits)
    _edit_run(p134, P134_OLD, P134_NEW, "P134")
    _edit_run(p153, P153_OLD, P153_NEW, "P153")

    # insert new cross-check paragraph immediately after P133
    new_p = d.add_paragraph(style="Body Text Indent")
    new_p.add_run(NEW_PARA)
    p133._p.addnext(new_p._p)

    print("Planned edits:")
    print("  [D-new ] inserted §5.2 ETABS cross-check paragraph after P133")
    print("  [P134  ] reconciled 'rather than benchmark' -> 'cross-checked ... benchmark = §4'")
    print("  [P153  ] appended 'global response cross-checked against ETABS to within 1%'")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply.")
        print("P134 now:", p134.text[-160:])
        print("P153 tail:", p153.text[-160:])
        print("new para :", new_p.text[:120])
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
