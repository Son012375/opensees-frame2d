"""Task 3 - insert the new Section 2 'Related Work' PROSE only (no table, no fields).

The user already renumbered the Heading-1 sections (3..7) and left an empty Section 2
slot; Task 1 added the roadmap clause and fixed cross-refs. This script inserts the
2.1-2.4 prose from related_work_section.md as plain text, immediately before the
'3. PROPOSED OPEN-SOURCE WORKFLOW AND METHODOLOGY' heading, matching peer paragraph
styles (Heading 1 / Body Text Indent + bold / Body Text Indent 2).

NOT inserted (field / Word work, handed to the user):
  * the Table 1 positioning comparison table + its caption
  * EndNote citation fields (the parenthetical author-year strings remain plain text)

Run with --apply to write; default dry-run. Aborts if the section already exists.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_relwork_backup.docx")

H1 = "2. RELATED WORK"

P_INTRO = (
    "Three streams of prior work bear on the proposed pipeline: extraction of "
    "analysis-ready structural models from BIM/IFC data, open-source BIM-to-finite-"
    "element workflows built on OpenSees, and the emerging use of generative AI and "
    "large language models for structural analysis. The present work is positioned "
    "within the first two streams and deliberately defers the third."
)
S21 = "2.1 IFC/BIM-to-Analytical-Model Extraction"
P_21 = (
    "Converting a coordination-oriented BIM into an analysis-ready model has been "
    "studied for over a decade. Ramaji and Memari (2018) interpret structural "
    "analytical models from the coordination view of building information models, and "
    "Hasan et al. (2019) generate geometrically accurate structural-analysis models "
    "within BIM-centered software. Fernández-Mora et al. (2022) review the broader "
    "integration of the structural project into the BIM paradigm and document the "
    "persistent gap between authoring models and analysis-ready models. More recently, "
    "Singh et al. (2024) automate the generation of a structural analytical model "
    "directly from an architectural IFC model to improve open-BIM interoperability, "
    "and Rudenko and Petryna (2025) generate finite-element models of varying "
    "complexity and dimensionality from BIM. These works establish the extraction "
    "methodology that the present pipeline adopts. The distinguishing feature here is "
    "an explicit, inspectable node-element analysis graph coupled with an automated "
    "connectivity-repair and validation-diagnostics layer (Section 3.3.2), rather than "
    "a single-pass translation whose intermediate representation is not surfaced for "
    "inspection."
)
S22 = "2.2 Open-Source BIM-to-FEM and OpenSeesPy Workflows"
P_22 = (
    "OpenSees (McKenna, 2011) and its Python interface OpenSeesPy (Zhu et al., 2018) "
    "provide an open computational substrate that a growing number of workflows "
    "target. Leonardi et al. (2024) present a scalable, open BIM-based workflow for "
    "the structural analysis of masonry building aggregates, and Llanos and "
    "Delgadillo (2025) perform linear seismic analysis and optimization of "
    "reinforced-concrete frames in OpenSeesPy. An IFC-based framework has also been "
    "proposed to integrate structural-analysis results back into BIM for "
    "code-compliance support (Buildings, 2026, 16(4):746). These confirm OpenSeesPy as "
    "a credible analysis engine for BIM-driven pipelines. The present work adds Korean "
    "Design Standard (KDS) clause-traced load automation and, distinctively, a "
    "metric-level cross-validation against two independent commercial programs (Midas "
    "Gen and ETABS); most BIM-to-FEM studies report agreement at the feasibility or "
    "visualization level rather than at the response-quantity level."
)
S23 = "2.3 Generative AI and Large Language Models for Structural Analysis (Deferred)"
P_23 = (
    "A rapidly growing 2024–2026 body of work applies generative AI and large language "
    "models (LLMs) to structural modeling. Liao et al. (2024) and Xie et al. (2025) "
    "survey generative-AI and AI-driven structural design automation, and Wang et al. "
    "(2025) apply BIM to the seismic reliability analysis of reinforced-concrete "
    "structures. Closest to an analysis pipeline, Liang et al. (2025) use an LLM to "
    "translate natural-language “structural analysis word problems” into executable "
    "OpenSeesPy scripts, reporting full accuracy with GPT-4o on a 20-problem benchmark; "
    "and a Model Context Protocol approach (Buildings, 2025, 15(17):3190) couples "
    "GPT-4o to OpenSeesPy and ETABS, reporting agreement within roughly one percent on "
    "storey drift, nodal displacement, and vibrational period for reinforced-concrete "
    "frames. This cluster is the most active and the most architecturally adjacent to "
    "a natural-language front-end. The present paper deliberately excludes the LLM, "
    "retrieval-augmented, and chatbot layers from its contributions and validates only "
    "the underlying BIM-to-analysis substrate, so that the quantitative commercial-"
    "baseline evidence is not entangled with the reliability of LLM-generated output; "
    "the natural-language and KDS-retrieval-augmented layers are reserved for a "
    "separate study."
)
S24 = "2.4 Positioning"
P_24 = (
    "Table 1 positions the present work against representative recent systems across "
    "input modality, intermediate representation, connectivity repair and diagnostics, "
    "solver, code system, and the depth of commercial cross-validation. None of the "
    "individual components — IFC parsing, OpenSeesPy analysis, or KDS load lookup — is "
    "claimed as novel in isolation. The contribution is their joint realization as a "
    "transparent, inspectable, clause-traced toolchain for the Korean regulatory "
    "context, with a surfaced node-element representation and connectivity diagnostics, "
    "validated at the response-quantity level against two commercial references."
)

# (style, text, bold_run)
BLOCKS = [
    ("Heading 1", H1, False),
    ("Body Text Indent 2", P_INTRO, False),
    ("Body Text Indent", S21, True),
    ("Body Text Indent 2", P_21, False),
    ("Body Text Indent", S22, True),
    ("Body Text Indent 2", P_22, False),
    ("Body Text Indent", S23, True),
    ("Body Text Indent 2", P_23, False),
    ("Body Text Indent", S24, True),
    ("Body Text Indent 2", P_24, False),
]


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))

    # guard: Related Work must not already exist
    if any(p.text.strip().upper().startswith("2. RELATED WORK") for p in d.paragraphs):
        print("ABORT: a '2. RELATED WORK' heading already exists.")
        return 1

    anchors = [p for p in d.paragraphs
               if p.style.name == "Heading 1" and p.text.strip().startswith("3.")
               and "PROPOSED" in p.text.upper()]
    if len(anchors) != 1:
        print(f"ABORT: found {len(anchors)} anchor headings (want 1).")
        return 1
    anchor = anchors[0]

    new_paras = []
    for style, text, bold in BLOCKS:
        p = d.add_paragraph(style=style)
        r = p.add_run(text)
        if bold:
            r.bold = True
        new_paras.append(p)

    # relocate the appended paragraphs to just before the anchor, in order
    cur = anchor._p
    for p in new_paras:
        cur.addprevious(p._p)
    # (addprevious before the same anchor preserves insertion order)

    if not apply:
        print("DRY-RUN: would insert 10 paragraphs before", repr(anchor.text.strip()[:45]))
        for style, text, _ in BLOCKS:
            print(f"   ({style}) {text[:60]!r}")
        print("\npass --apply to write.")
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"backup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
