"""Session 2026-07-02 (peer-review MAJOR REVISION): C1 + C6 framing inversion.

Highest-leverage revision per the editorial report (peer_review_2026-07-02.md):
invert the four high-visibility surfaces so they say exactly what the honest methods
body and validation files already say.

Author decisions (confirmed this session):
  - C6 = FULL framing inversion; title rebalanced away from IFC parsing (author: IFC
    foregrounding not important; keep a scoped "alternative" positioning reviewers
    will not challenge).
  - C3 = expand irregular-plan validation; irregular positioned as the distinguishing
    class, benchmark reserved (Midas handoff prepared separately). Interim framing here
    keeps irregular as demonstrated + ETABS-cross-checked, NOT yet benchmark-validated.

Five plain-run wholesale rewrites (field-safe: refuse if paragraph carries a field):
  P000  title            -> transparency/clause-trace/benchmark foregrounded; IFC dropped
  P007  abstract         -> substrate demotion + benchmark/cross-check definition + ETABS
                            scoped as solver-equivalence + transferable evidence standard
  P018  contributions    -> lead with the transferable code-agnostic evidence standard
  P169  conclusion (res) -> ETABS = solver-equivalence check (isolates parse+loads); fixes
                            the 2026-07-02 overclaim "extending ... to the IFC-to-analysis
                            pipeline itself" (C1 decisive)
  P170  conclusion (pos) -> contribution = transferable evidence standard; irregular =
                            distinguishing class with Midas benchmark as planned extension

Numbers unchanged & verified against validation/ifc_example_etabs_compare.json (31/31 OK
within 1%) and Table 6 (100/112 <1%, 12 <4%). No field/SEQ/citation touched (targets are
plain runs=1..3). Backup -> ..._v2.pre_reframe_2026-07-02.docx. Dry-run default; --apply.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_reframe_2026-07-02.docx")

# ------------------------------------------------------------------ new text
NEW_TITLE = (
    "A Transparent, Clause-Traced Open-Source Alternative for Steel-Frame Structural "
    "Analysis: Response-Quantity Benchmarking of a KDS-Automated BIM-to-OpenSeesPy "
    "Workflow against Midas Gen and ETABS"
)

NEW_ABSTRACT = (
    "Commercial structural-analysis programs remain central to building design, yet their "
    "closed preprocessing and result-recovery pipelines can limit independent inspection, "
    "reproduction, and customization of the analysis path. This paper presents a transparent "
    "open-source alternative for regular steel frame buildings that couples clause-traced "
    "Korean Design Standard (KDS) load automation with OpenSeesPy three-dimensional frame "
    "analysis and evaluates it against commercial software at the response-quantity level. A "
    "node-element parsing layer converts BIM/IFC input into an inspectable analysis graph as "
    "enabling substrate; KDS automation then generates dead, live, seismic, wind, and snow "
    "cases and assembles code combinations, with every value traced to its originating clause "
    "in a hand-check appendix and four workflow-introduced simplifications disclosed. Two "
    "evidence standards are distinguished: a benchmark (an independently modeled, matched-input "
    "case compared against Midas Gen) and a cross-check (identical generated loads replayed to "
    "an independently assembled model, isolating assembly and recovery). Across five controlled "
    "Midas Gen benchmark cases, 100 of 112 response metrics agree within 1% and the remaining "
    "12 within 4% once sign and local-axis conventions are aligned, the larger differences "
    "confined to element-formulation-sensitive quantities of a single symmetric "
    "three-dimensional case, not global equilibrium. An IFC-derived example is also "
    "cross-checked against ETABS — a solver-equivalence check on the assembled model under "
    "identical loads, not a test of the parse or load generation — agreeing on 31 "
    "global-response metrics within 1%. Orthogonal irregular plans are supported through "
    "zone-based decomposition and demonstrated end-to-end, with commercial benchmarking "
    "reserved as ongoing work. Within this scope the pipeline is a credible, transparent, "
    "KDS-traced alternative computational path, and its clause-trace, disclosure, and "
    "hand-check protocol is offered as a transferable, code-agnostic evidence standard."
)

NEW_CONTRIB = (
    "The contributions of the paper are threefold. First, and most centrally, the workflow "
    "establishes and applies a transferable, code-agnostic evidence standard for open-source "
    "BIM-to-analysis pipelines: every automated load value is traced to its originating clause "
    "and hand-checked, every workflow-introduced simplification is disclosed, and the resulting "
    "model is compared against a commercial baseline at the level of individual response "
    "quantities rather than at the feasibility or visualization level typical of prior "
    "BIM-to-FEM work. Second, this standard is instantiated for the Korean Design Standards "
    "(KDS 41 12 00, KDS 41 17 00, and KDS 17 10 00) with a clause-by-clause hand-check of the "
    "generated dead, live, seismic, wind, and snow values and the four workflow-introduced "
    "simplifications (MEP allowance, wind topographic factor, gust factor, and pressure "
    "coefficient) disclosed, and quantified against Midas Gen across five controlled cases "
    "(100 of 112 metrics within 1%, the remaining 12 within 4%) and against ETABS on an "
    "IFC-derived example (31 global-response metrics within 1%). Third, these results are made "
    "reproducible and auditable by a transparent, inspectable open-source pipeline that surfaces "
    "the node-element analysis graph — following established BIM-to-FEM methodology rather than "
    "claiming parsing novelty — and supports orthogonal irregular plans (L-shape, setback) "
    "through zone-based decomposition, demonstrated end-to-end in Section 5.3. The "
    "evidence-standard protocol, not any single component, is the portable contribution: a team "
    "adopting Eurocode or ASCE 7 would replace the clause database and load logic while retaining "
    "the trace, disclose, and benchmark discipline."
)

NEW_P169 = (
    "Across five controlled benchmark cases spanning two-dimensional frames, a symmetric "
    "three-dimensional moment frame, and a supplementary geometrically nonlinear case, 100 of "
    "112 response metrics agreed with Midas Gen within 1% and the remaining 12 within 4%, with "
    "the larger discrepancies concentrated in element-formulation-sensitive quantities of the "
    "single three-dimensional case rather than in global equilibrium. An IFC-derived three-story "
    "application example exercised the full pipeline end-to-end and returned drift, displacement, "
    "and member-strength results consistent with engineering expectations under KDS-generated "
    "loading; its global response was independently cross-checked against ETABS on the assembled "
    "model under identical applied loads — a solver-equivalence check that isolates model "
    "assembly and response recovery from the IFC parse and the load generation — with all 31 "
    "compared metrics agreeing within 1%."
)

NEW_P170 = (
    "Within the benchmark-validated scope — regular orthogonal steel frame buildings under "
    "elastic three-dimensional analysis — the results support positioning the pipeline as a "
    "credible, transparent, KDS-traced alternative computational path alongside the commercial "
    "baseline. The contribution lies primarily in the transferable evidence standard — "
    "clause-tracing, simplification disclosure, and response-quantity benchmarking — that the "
    "inspectable workflow makes reproducible and auditable, rather than in methodological novelty "
    "for any single component. The same pipeline supports orthogonal irregular plan "
    "configurations (L-shape, setback) through zone-based decomposition, demonstrated in "
    "Section 5.3 and cross-checked against ETABS; their independent commercial benchmarking "
    "against Midas Gen is the primary planned extension of this evidence to a distinguishing "
    "structural class."
)

EDITS = [
    ("Toward an Open-Source BIM-to-Analysis Workflow for Steel Frame Buildings", NEW_TITLE, "P000"),
    ("closed preprocessing and result-recovery pipelines can limit independent inspection", NEW_ABSTRACT, "P007"),
    ("The contributions of the paper are threefold.", NEW_CONTRIB, "P018"),
    ("Across five controlled benchmark cases spanning two-dimensional frames", NEW_P169, "P169"),
    ("the results support the position that the pipeline serves as a credible", NEW_P170, "P170"),
]

SENTINEL = "A Transparent, Clause-Traced Open-Source Alternative"


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:50]!r}")
    return hits[0]


def set_text(par, new, tag):
    """Field-safe wholesale rewrite: refuse if the paragraph carries any field."""
    xml = par._p.xml
    if "fldChar" in xml or "instrText" in xml:
        raise SystemExit(f"ABORT [{tag}]: paragraph carries a field; refusing wholesale rewrite")
    if not par.runs:
        par.add_run(new)
        return
    par.runs[0].text = new
    for r in par.runs[1:]:
        r.text = ""


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    if any(SENTINEL in p.text for p in paras):
        raise SystemExit("ABORT: already applied (title sentinel present).")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    print(f"BEFORE: paragraphs={np0} tables={nt0} inline_shapes={ni0} fldChar={fld0} instrText={ins0}")

    targets = []
    for needle, new, tag in EDITS:
        par = find_para(paras, needle, tag)
        set_text(par, new, tag)
        targets.append((tag, par, new))

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    print(f"AFTER : paragraphs={np1} tables={nt1} inline_shapes={ni1} fldChar={fld1} instrText={ins1}")

    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0} -> {fld1}/{ins1}"
    assert (np1, nt1, ni1) == (np0, nt0, ni0), "structure count changed"

    for tag, par, _ in targets:
        wc = len(par.text.split())
        print(f"\n--- {tag} ({wc} words) ---\n{par.text}")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
