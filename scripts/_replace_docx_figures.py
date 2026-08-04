"""Replace the 7 embedded figures in the canonical v2 DOCX with the new
publication-quality set (figures/final/fig1..fig7.png), in document order.

Keeps each figure's existing display WIDTH (so nothing overflows the page) and
recomputes HEIGHT from the new image's aspect ratio (so nothing is distorted).
A backup (..._v2.pre_figs_backup.docx) is written first.
"""
from __future__ import annotations

import shutil
import struct
from pathlib import Path

import docx
from docx.shared import Emu

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs" / "paper1_open_source_alternative"
DOC = BASE / "drafts" / "open_source_alternative_review_draft_v2.docx"
FINAL = BASE / "figures" / "final"


def png_size(b: bytes) -> tuple[int, int]:
    assert b[12:16] == b"IHDR", "not a PNG / unexpected header"
    w, h = struct.unpack(">II", b[16:24])
    return w, h


def main() -> int:
    backup = DOC.with_name("open_source_alternative_review_draft_v2.pre_figs_backup.docx")
    shutil.copy2(DOC, backup)
    print(f"[backup] {backup.name}")

    d = docx.Document(str(DOC))
    shapes = d.inline_shapes
    print(f"[doc] {len(shapes)} inline shapes")
    if len(shapes) != 7:
        print("[!] expected 7 inline shapes; aborting to be safe")
        return 1

    for idx, shape in enumerate(shapes, start=1):
        new_path = FINAL / f"fig{idx}.png"
        new_bytes = new_path.read_bytes()
        nw, nh = png_size(new_bytes)

        blip = shape._inline.graphic.graphicData.pic.blipFill.blip
        rId = blip.embed
        part = d.part.related_parts[rId]
        old_w_emu = int(shape.width)
        old_h_emu = int(shape.height)

        part._blob = new_bytes  # swap image bytes in place
        shape.height = Emu(int(old_w_emu * nh / nw))  # keep width, fix aspect

        print(f"  Fig {idx}: {new_path.name} {nw}x{nh}px  "
              f"width={old_w_emu/914400:.2f}in  "
              f"h {old_h_emu/914400:.2f}->{int(shape.height)/914400:.2f}in")

    d.save(str(DOC))
    print(f"[saved] {DOC.name}")

    # verify
    d2 = docx.Document(str(DOC))
    media = [p.partname for p in d2.part.package.iter_parts() if "media" in p.partname]
    print(f"[verify] inline_shapes={len(d2.inline_shapes)}, media={len(media)}, "
          f"paragraphs={len(d2.paragraphs)}, tables={len(d2.tables)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
