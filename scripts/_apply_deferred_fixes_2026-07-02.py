"""Handle the three deferred items from the full-manuscript audit.

ITEM 1 (P091 / §3.4.2): P091 says sign conventions are aligned "as described in
  Section 3.4.2", but §3.4.2 did not describe it. Append a sign/local-axis convention
  sentence to §3.4.2's post-processing paragraph (P068) so the cross-reference is valid.

ITEM 2 (Appendix A snow hand-check): snow (S = 0.5 kN/m2) is now generated but was not
  hand-checked. Insert a new subsection "A.6 Snow Load (KDS 41 12 00 §6)" with the
  clause computation (S_g = 0.5 Seoul; C_b0.7·C_e1.0·C_t1.0·I_s1.0 -> S_flat 0.35;
  minimum S_m = I_s·S_g = 0.5 governs -> S = 0.5 kN/m2 = 60 kN / 120 m2 roof, matching
  the auto value). Renumber the following subsections: A.6 Load Combinations -> A.7,
  A.7 Summary -> A.8. (No table added, so Table A-7 (combos) keeps its number; verified
  there are no in-text cross-references to A.6/A.7.) Update the P204 category list +
  count (Forty-nine -> Fifty) and P231 count (49 -> 50).

ITEM 3 (Table 9 ratio rounding): 0.21%/0.107 and 0.46%/0.232 are internally inconsistent
  (0.21/2.0 = 0.105). Make the ratios match the displayed drifts: 0.105 and 0.230.

Snow numbers verified by running load_generator.generate_snow_loads on example_input.json.
Backup -> ..._v2.pre_deferred_2026-07-02.docx. Dry-run default; --apply saves.
"""
from __future__ import annotations

import copy
import shutil
import sys
from pathlib import Path

import docx
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_deferred_2026-07-02.docx")

SIGN_SENTENCE = (
    " Prior to comparison, all response quantities are transformed to a common textbook "
    "sign convention — sagging-positive bending, with the two programs' local-axis shear "
    "and moment signs reconciled — so that the OpenSeesPy and Midas Gen quantities are "
    "compared on identical conventions."
)

SNOW_HEADING = "A.6 Snow Load (KDS 41 12 00 §6)"
SNOW_BODY = (
    "The flat-roof design snow load is S = C_b · C_e · C_t · I_s · S_g, with C_b = 0.7, "
    "C_e = 1.0 (partial exposure), C_t = 1.0 (heated building), and I_s = 1.0 (Importance "
    "Class II). For the Seoul site the basic ground snow load retrieved from the hazard "
    "database is S_g = 0.5 kN/m2, giving a formula value of 0.7 × 0.5 = 0.35 kN/m2. "
    "Because the code minimum roof snow load S_m = I_s · S_g = 0.5 kN/m2 (for S_g ≤ 1.0) "
    "governs, the design roof snow load is 0.5 kN/m2, applied to the 120 m2 roof as "
    "60.0 kN — matching the auto-generated value."
)


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:40]!r}")
    return hits[0]


def repl(par, old, new, tag):
    for r in par.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return
    raise SystemExit(f"ABORT [{tag}]: not found {old[:40]!r}")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    if any("A.6 Snow Load" in p.text for p in paras):
        raise SystemExit("ABORT: already applied (A.6 Snow present).")

    # ---- ITEM 1: sign convention into §3.4.2 (P068) ----
    p068 = find_para(paras, "Post-processing extracts nodal displacements, support reactions", "P068")
    p068.add_run(SIGN_SENTENCE)

    # ---- ITEM 3: Table 9 ratios ----
    t9 = next(t for t in d.tables
              if [c.text for c in t.rows[0].cells][:3] == ["Category", "Parameter", "Value"]
              and len(t.rows) == 11)
    repl(t9.rows[5].cells[2].paragraphs[0], "ratio 0.107", "ratio 0.105", "T9R5")
    repl(t9.rows[6].cells[2].paragraphs[0], "ratio 0.232", "ratio 0.230", "T9R6")

    # ---- ITEM 2: appendix snow subsection + renumber ----
    p226 = find_para(paras, "A.6 Load Combinations", "P226-combos")
    p230 = find_para(paras, "A.7 Summary", "P230-summary")
    p204 = find_para(paras, "quantities are tabulated across dead, live", "P204")
    p231 = find_para(paras, "Across all 49 individually checked quantities", "P231")

    # new snow heading + body cloned from the A.6-combos heading and a body para for style fidelity
    heading_src = p226._p            # subsection-heading style
    body_src = find_para(paras, "The MEP allowance", "P210-body")._p
    new_heading_p = copy.deepcopy(heading_src)
    new_body_p = copy.deepcopy(body_src)
    p226._p.addprevious(new_heading_p)
    p226._p.addprevious(new_body_p)
    nh = Paragraph(new_heading_p, p226._parent)
    nb = Paragraph(new_body_p, p226._parent)
    for r in nh.runs[1:]:
        r.text = ""
    nh.runs[0].text = SNOW_HEADING
    for r in nb.runs[1:]:
        r.text = ""
    nb.runs[0].text = SNOW_BODY

    # renumber following subsections
    repl(p226, "A.6 Load Combinations", "A.7 Load Combinations", "P226-renum")
    repl(p230, "A.7 Summary", "A.8 Summary", "P230-renum")
    # update category list + counts
    repl(p204, "quantities are tabulated across dead, live, equivalent lateral seismic, wind, and combination categories",
         "quantities are checked across dead, live, equivalent lateral seismic, wind, snow, and combination categories", "P204-cat")
    repl(p204, "Forty-nine", "Fifty", "P204-count")
    repl(p231, "Across all 49 individually checked quantities",
         "Across all 50 individually checked quantities", "P231-count")

    print("Staged: Item1 (sign conv §3.4.2), Item2 (A.6 Snow + renumber A.7/A.8), Item3 (Table 9 ratios).")
    if not apply:
        print("DRY-RUN. new heading:", nh.text)
        print("new body[:80]:", nb.text[:80])
        print("P226 now:", p226.text[:40], "| P230 now:", p230.text)
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"backup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
