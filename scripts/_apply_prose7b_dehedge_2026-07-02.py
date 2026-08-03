"""Session 2026-07-02 (MAJOR REVISION): prose 7b — de-hedge (trim redundant refrains).

Reviewer flow item: the "demonstration vs benchmark / benchmark not a replacement" refrain
recurs 15+ times. Keep the authoritative statements (P104 summary; P154/P170 alternative-not-
replacement; P155 demonstration-scope; P073/P161 screening-not-certification; P121 §5 intro)
and remove three PURE duplicates that add nothing:

  P081  drop "is not framed as a complete replacement test; rather," (dup of P104/P017).
  P133  drop trailing ", with the controlled commercial benchmark campaign remaining the five
        cases of Section 4" (dup of P121/P155).
  P149  drop opening "Distinct in scope from the five-case benchmark of Section 4, " (dup of
        P140/P155); recapitalize "This".

Run-level edits only (no field touched). Backup -> ..._v2.pre_dehedge_2026-07-02.docx.
Dry-run default; --apply.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_dehedge_2026-07-02.docx")

EDITS = [
    ("The benchmark is not framed as a complete replacement test",
     "The benchmark is not framed as a complete replacement test; rather, it examines",
     "The benchmark examines", "P081"),
    ("Distinct in scope from the five-case benchmark of Section 4, this example illustrates",
     "Distinct in scope from the five-case benchmark of Section 4, this example illustrates",
     "This example illustrates", "P149"),
]

# P133 trailing hedge clause is split across run[0]/run[1]:
#   run[0] "...as reported above, with the controlled commercial benchmark "
#   run[1] "campaign remaining the five cases of Section 4."
P133_LOC = "the reported response quantities are interpreted as workflow demonstration results"
P133_R0_OLD = ", with the controlled commercial benchmark "
P133_R0_NEW = "."
P133_R1_OLD = "campaign remaining the five cases of Section 4."
P133_R1_NEW = ""

SENTINEL = "The benchmark examines whether"


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:50]!r}")
    return hits[0]


def run_replace(par, old, new, tag):
    for r in par.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return
    raise SystemExit(f"ABORT [{tag}]: substring not in a single run: {old[:50]!r}")


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs
    if any(SENTINEL in p.text for p in paras):
        raise SystemExit("ABORT: already applied (P081 sentinel present).")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    print(f"BEFORE: paragraphs={np0} tables={nt0} inline_shapes={ni0} fldChar={fld0} instrText={ins0}")

    for loc, old, new, tag in EDITS:
        par = find_para(paras, loc, tag)
        run_replace(par, old, new, tag)

    # P133 two-run trim
    p133 = find_para(paras, P133_LOC, "P133")
    r0ok = r1ok = False
    for r in p133.runs:
        if P133_R0_OLD in r.text and not r0ok:
            r.text = r.text.replace(P133_R0_OLD, P133_R0_NEW); r0ok = True
        if P133_R1_OLD in r.text and not r1ok:
            r.text = r.text.replace(P133_R1_OLD, P133_R1_NEW); r1ok = True
    if not (r0ok and r1ok):
        raise SystemExit(f"ABORT [P133]: r0ok={r0ok} r1ok={r1ok}")

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    print(f"AFTER : paragraphs={np1} tables={nt1} inline_shapes={ni1} fldChar={fld1} instrText={ins1}")
    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0} -> {fld1}/{ins1}"
    assert (np1, nt1, ni1) == (np0, nt0, ni0), "structure count changed"

    for needle, tag in [("The benchmark examines whether", "P081"),
                        ("This example illustrates the workflow's irregular-plan", "P149"),
                        (P133_LOC, "P133")]:
        p = find_para(d.paragraphs, needle, tag)
        print(f"\n--- {tag} ---\n{p.text[:360]}")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
