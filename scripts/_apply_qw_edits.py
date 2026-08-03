"""Apply the safe quick-win redlines to the canonical manuscript (v2 DOCX).

Edits applied:
  QW-1  Add a "Young's modulus E" row to the OpenSeesPy config table (Table 2).
  QW-2  Section 5.3 precision: note the P-Delta/drift/torsion analysis primitives
        are already implemented, only their certified-compliance integration is future.
  QW-3  Abstract: state that the cross-program comparison aligns sign/local-axis
        conventions (pre-empts the "calibrated reconciliation" objection, neutrally).

NOT applied here: QW-4 (Figure 5 reference already present in para 118),
QW-5 (author email — needs author confirmation).

A backup (..._v2.pre_qw_backup.docx) was made before running.
"""
from __future__ import annotations

from pathlib import Path
import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"


def replace_in_para(p, old, new) -> bool:
    if old not in p.text:
        return False
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    # substring spans multiple runs -> collapse into first run (paragraph is plain text)
    full = p.text.replace(old, new)
    for i, r in enumerate(p.runs):
        r.text = full if i == 0 else ""
    return True


def set_cell(cell, text, template_cell):
    cell.text = text
    tpl_runs = template_cell.paragraphs[0].runs
    if tpl_runs and cell.paragraphs[0].runs:
        t = tpl_runs[0].font
        r = cell.paragraphs[0].runs[0].font
        r.name = t.name
        r.size = t.size
        r.bold = t.bold


def main() -> int:
    d = docx.Document(str(DOC))
    log = []

    # ---- QW-3: abstract (para 9) ----
    QW3_OLD = "the remaining 12 within 4%, with the larger discrepancies"
    QW3_NEW = ("the remaining 12 within 4% once the two programs' sign and "
               "local-axis conventions are aligned, with the larger discrepancies")
    ok3 = any(replace_in_para(p, QW3_OLD, QW3_NEW) for p in d.paragraphs)
    log.append(f"QW-3 abstract sign/axis note: {'APPLIED' if ok3 else 'NOT FOUND'}")

    # ---- QW-2: section 5.3 (para 150) ----
    QW2_OLD = ("and torsional irregularity assessment, would convert the "
               "preliminary screening into a complete code-compliance pathway.")
    QW2_NEW = ("and torsional irregularity assessment — whose underlying "
               "analysis primitives (P-Delta amplification, drift amplification, "
               "and torsional-irregularity evaluation) are already implemented and "
               "were exercised in the verification suite — would convert the "
               "preliminary screening into a complete, certified code-compliance pathway.")
    ok2 = any(replace_in_para(p, QW2_OLD, QW2_NEW) for p in d.paragraphs)
    log.append(f"QW-2 sec5.3 precision: {'APPLIED' if ok2 else 'NOT FOUND'}")

    # ---- QW-1: add E row to config table (table with elasticBeamColumn) ----
    ok1 = False
    for t in d.tables:
        cells = [c.text for row in t.rows for c in row.cells]
        if any("elasticBeamColumn" in c for c in cells):
            # avoid duplicate if rerun
            if any("Young's modulus" in c for c in cells):
                log.append("QW-1 E row: ALREADY PRESENT")
                ok1 = True
                break
            tpl = t.rows[1].cells  # an existing data row, for font matching
            new_row = t.add_row().cells
            set_cell(new_row[0], "Young’s modulus E", tpl[0])
            set_cell(new_row[1], "210,000 MPa", tpl[1])
            set_cell(new_row[2],
                     "Structural steel; identical value in OpenSeesPy and the "
                     "commercial reference for all benchmark cases", tpl[2])
            ok1 = True
            log.append("QW-1 E row: APPLIED (Table with elasticBeamColumn)")
            break
    if not ok1:
        log.append("QW-1 E row: CONFIG TABLE NOT FOUND")

    if ok1 and ok2 and ok3:
        d.save(str(DOC))
        log.append(f"SAVED -> {DOC.name}")
    else:
        log.append("NOT SAVED (one or more edits failed; investigate)")

    print("\n".join(log))
    return 0 if (ok1 and ok2 and ok3) else 1


if __name__ == "__main__":
    raise SystemExit(main())
