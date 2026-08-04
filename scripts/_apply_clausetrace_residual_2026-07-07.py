# -*- coding: utf-8 -*-
"""Apply the residual clause-trace / label fixes from the 2026-07-07 DBE re-audit.

All low/cosmetic clause-tracing corrections; no reported analysis value changes.
Field-safe (Zotero fldChar/instrText untouched): every edit is a plain-text-run
replacement located by content anchor. Pre-scan proves each target string occurs
exactly where expected; ABORT on any ambiguity or if already applied. Dry-run by
default; dated backup on --apply.

  W1   Wind topographic factor Kzt clause  §5.2.5 -> §5.5.5
         (KDS Kzt_scope DB record clause_id=5.5.5, table 5.5-4; Kzr is §5.5.4)
         - P231 body: "computed per KDS §5.2.5 from the site topography"
         - Table A-5 (wind param) row D2 clause cell "KDS 41 12 00 §5.2.5 (flat terrain)"
  Ce   Snow exposure-coeff label  "(partial exposure)" -> "(normal exposure)"
         (Ce=1.0 is the 'normal' category; partial=0.9 in KDS 표 4.3-1 / snow DB)
         - P233 body: "C_e = 1.0 (partial exposure)"
  NC-1 Table 9 design-drift demand/capacity ratio  0.105/0.230 -> 0.107/0.231
         (example_section4_results.json design_drift_x/y_ratio = 0.1074 / 0.2315)
         - Table 9 R5 "0.21% (ratio 0.105)", R6 "0.46% (ratio 0.230)"

Usage: python scripts/_apply_clausetrace_residual_2026-07-07.py [--apply]
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_clausetrace_2026-07-07.docx")

SECT = "§"  # §


def _iter_cells(d):
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                yield c


def _replace_in_paragraph(p, old, new, tag):
    """Replace `old`->`new` in the single run of paragraph `p` that contains it.

    Aborts if `old` does not occur exactly once, or spans multiple runs.
    """
    n = p.text.count(old)
    if n != 1:
        raise SystemExit(f"ABORT [{tag}]: {old!r} occurs {n}x in paragraph (want 1). text={p.text!r}")
    hits = [i for i, r in enumerate(p.runs) if old in r.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: {old!r} not confined to one run (run-hits={len(hits)}). text={p.text!r}")
    r = p.runs[hits[0]]
    r.text = r.text.replace(old, new)


def _replace_in_cell(cell, old, new, tag):
    for p in cell.paragraphs:
        if old in p.text:
            _replace_in_paragraph(p, old, new, tag)
            return
    raise SystemExit(f"ABORT [{tag}]: {old!r} not in cell. text={cell.text!r}")


def _find_para(paras, substr, tag):
    hits = [p for p in paras if substr in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: para anchor {substr!r} matched {len(hits)} (want 1).")
    return hits[0]


def _prescan(d):
    """Prove the document is in the expected pre-edit state before touching anything."""
    paras = d.paragraphs
    cells = list(_iter_cells(d))
    all_texts = [p.text for p in paras] + [c.text for c in cells]

    def total(sub):
        return sum(t.count(sub) for t in all_texts)

    print("Pre-scan (occurrence counts across all paragraphs + table cells):")
    checks = {
        SECT + "5.2.5": 2,          # P231 + Table A-5 D2
        "(partial exposure)": 1,    # P233
        "(ratio 0.105)": 1,         # Table 9 R5
        "(ratio 0.230)": 1,         # Table 9 R6
    }
    ok = True
    for sub, want in checks.items():
        got = total(sub)
        flag = "OK" if got == want else "MISMATCH"
        if got != want:
            ok = False
        print(f"  {sub!r:24s} found {got} (want {want}) [{flag}]")

    # idempotency: none of the *target* strings may already be present
    already = {
        SECT + "5.5.5": total(SECT + "5.5.5"),
        "(normal exposure)": total("(normal exposure)"),
        "(ratio 0.107)": total("(ratio 0.107)"),
        "(ratio 0.231)": total("(ratio 0.231)"),
    }
    for sub, got in already.items():
        if got:
            ok = False
            print(f"  ALREADY-APPLIED: {sub!r} present {got}x -> ABORT")
    if not ok:
        raise SystemExit("ABORT: pre-scan state does not match expectation (see above).")
    print("  pre-scan OK\n")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    _prescan(d)

    plan = []

    # -- W1: P231 body §5.2.5 -> §5.5.5 --------------------------------------
    p231 = _find_para(paras, "would otherwise be computed per KDS", "W1/P231")
    _replace_in_paragraph(p231, SECT + "5.2.5", SECT + "5.5.5", "W1/P231")
    plan.append("[W1]   P231 body    KDS %s5.2.5 -> %s5.5.5" % (SECT, SECT))

    # -- W1: Table A-5 row D2 clause cell §5.2.5 -> §5.5.5 -------------------
    t15 = d2cell = None
    for t in d.tables:
        for r in t.rows:
            if r.cells[0].text.strip() == "D2" and "Topographic factor" in r.cells[1].text:
                t15, d2cell = t, r.cells[2]
                break
        if t15 is not None:
            break
    if d2cell is None:
        raise SystemExit("ABORT [W1/T15]: Table A-5 D2 row not found.")
    _replace_in_cell(d2cell, SECT + "5.2.5", SECT + "5.5.5", "W1/T15-D2")
    plan.append("[W1]   Table A-5 D2  KDS %s5.2.5 -> %s5.5.5" % (SECT, SECT))

    # -- Ce: P233 snow exposure label ----------------------------------------
    p233 = _find_para(paras, "flat-roof design snow load is S", "Ce/P233")
    _replace_in_paragraph(p233, "(partial exposure)", "(normal exposure)", "Ce/P233")
    plan.append("[Ce]   P233 body    C_e (partial exposure) -> (normal exposure)")

    # -- NC-1: Table 9 design-drift ratios -----------------------------------
    t9 = None
    for t in d.tables:
        if any("(ratio 0.105)" in c.text for r in t.rows for c in r.cells):
            t9 = t
            break
    if t9 is None:
        raise SystemExit("ABORT [NC-1/T9]: Table 9 (ratio 0.105) not found.")
    c5 = next(c for r in t9.rows for c in r.cells if "(ratio 0.105)" in c.text)
    c6 = next(c for r in t9.rows for c in r.cells if "(ratio 0.230)" in c.text)
    _replace_in_cell(c5, "(ratio 0.105)", "(ratio 0.107)", "NC-1/T9R5")
    _replace_in_cell(c6, "(ratio 0.230)", "(ratio 0.231)", "NC-1/T9R6")
    plan.append("[NC-1] Table 9 R5   (ratio 0.105) -> (ratio 0.107)")
    plan.append("[NC-1] Table 9 R6   (ratio 0.230) -> (ratio 0.231)")

    print("Planned edits:")
    for line in plan:
        print("  " + line)

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply.")
        print("  P231 (§5.5.5?):", (SECT + "5.5.5") in p231.text)
        print("  T15 D2 cell  :", d2cell.text)
        print("  P233 tail    :", p233.text[:90], "...")
        print("  T9 R5 cell   :", c5.text)
        print("  T9 R6 cell   :", c6.text)
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED  -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
