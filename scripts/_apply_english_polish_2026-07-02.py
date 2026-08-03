"""Session 2026-07-02: apply vetted native-English/flow edits from the multi-lens review
workflow (wksgmmrrf). 49 vetted old->new pairs; 3 skipped by author judgment (see SKIP).

Each edit is applied ONLY if its exact "old" string occurs in exactly ONE paragraph AND
within a single run (no citation/field split). Non-conforming edits are skipped and reported
(not fatal), so the run applies what is safe and leaves the rest for manual handling. No
number/claim/scoping/citation is touched (enforced by the workflow + these guards).

Backup -> ..._v2.pre_engpolish_2026-07-02.docx. Dry-run default; --apply.
"""
from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_engpolish_2026-07-02.docx")
EDITS_JSON = Path(r"C:/Users/youm/AppData/Local/Temp/claude/d--son-opensees-MCP/10bb2238-4ffa-4a07-bb0b-524f28b056a2/scratchpad/vetted_edits.json")

# Author-skipped (word budget / meaning-preservation):
SKIP_OLD = {
    "Orthogonal irregular plans are supported through zone-based decomposition and demonstrated end-to-end, with commercial benchmarking reserved for future work.",
    "The contribution lies primarily in the transferable evidence standard",
    "is the primary planned extension of this evidence to a distinguishing structural class",
}


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def paras_with(paras, needle):
    return [p for p in paras if needle in p.text]


def apply_in_single_run(par, old, new):
    """Replace old->new inside the one run that fully contains old. Return True if done."""
    if "fldChar" in par._p.xml or "instrText" in par._p.xml:
        # paragraph has a field: only allow if old sits fully inside a plain run
        pass
    for r in par.runs:
        if old in r.text:
            r.text = r.text.replace(old, new, 1)
            return True
    return False


def main(apply: bool) -> int:
    edits = json.loads(EDITS_JSON.read_text(encoding="utf-8"))
    d = docx.Document(str(DOC))
    paras = d.paragraphs
    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)

    applied, skipped = [], []
    for e in edits:
        old, new = e["old"], e["new"]
        if old in SKIP_OLD:
            skipped.append((e["section"], "author-skip", old[:60]))
            continue
        hits = paras_with(paras, old)
        if len(hits) == 0:
            skipped.append((e["section"], "not found", old[:60]))
            continue
        if len(hits) > 1:
            skipped.append((e["section"], f"{len(hits)} matches", old[:60]))
            continue
        if apply_in_single_run(hits[0], old, new):
            applied.append((e["section"], old[:55]))
        else:
            skipped.append((e["section"], "spans runs", old[:60]))

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0}->{fld1}/{ins1}"
    assert (np1, nt1, ni1) == (np0, nt0, ni0), "structure count changed"

    # abstract word count check
    ab = [p for p in d.paragraphs if "closed preprocessing and result recovery limit" in p.text]
    abw = len(ab[0].text.split()) if ab else -1

    print(f"BEFORE fld={fld0}/{ins0} paras={np0} tables={nt0} imgs={ni0}")
    print(f"AFTER  fld={fld1}/{ins1} paras={np1} tables={nt1} imgs={ni1}")
    print(f"APPLIED {len(applied)} / {len(edits)}  |  abstract words now = {abw}")
    for s, o in applied:
        print(f"  [{s}] {o}")
    if skipped:
        print(f"\nSKIPPED {len(skipped)}:")
        for s, why, o in skipped:
            print(f"  [{s}] ({why}) {o}")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
