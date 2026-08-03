# -*- coding: utf-8 -*-
"""Apply Tier-1 fixes from the 2026-07-07 adversarial pipeline re-audit.

Field-safe (Zotero fldChar/instrText untouched): every target paragraph/cell is a
single plain-text run (verified by probe). Content-anchored, ABORT on ambiguity,
dry-run by default, dated backup.

  N1          Table 8 wind Fx/Fy  100.0 / 120.0 -> 93.1 / 111.7  (stale; body/App.A = 93.1/111.7)
  SNOW-CLAUSE P232  "A.6 Snow Load (KDS 41 12 00 §6)" -> "... §4)"  (snow is §4; DB tbl 표 4.3-1)
  G2          P217 + Table A-2 (T12) "Table 3.1-1" -> "Table 3.2-1"  (actual KDS 표 3.2-1)
  G1          Table A-1 (T11) R1 clause "...Table 4.2-1" -> "...Table 4.2-1 (building-adjusted)"
              + P216 disclosure sentence (table prints 24.5; 24.0 = building value)
  SNOW-01     Table A-7 (T17) Wind X/Y combos add "+ 0.5 S"; P237 "gravity and seismic"
              -> "gravity, wind, and seismic"  (code now generates 1.2D+1.0L+1.0W+0.5S per §1.7-4)

Usage: python scripts/_apply_tier1_reaudit_2026-07-07.py [--apply]
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_tier1_2026-07-07.docx")

SECT = "§"   # §
PM = "±"     # ±

P216_PREPEND = (
    "The reinforced-concrete unit weight is adopted as the building-structure value of "
    "24.0 kN/m³; KDS 24 12 21 Table 4.2-1 lists 24.5 kN/m³ for reinforced concrete, "
    "the higher value reflecting the reinforcement ratio assumed in bridge design. "
)


def _repl_run(container, run_idx, old, new, tag):
    r = container.runs[run_idx]
    if old not in r.text:
        raise SystemExit(f"ABORT [{tag}]: {old!r} not in run {run_idx}: {r.text!r}")
    r.text = r.text.replace(old, new)
    return r.text


def _find_para(paras, substr):
    hits = [p for p in paras if substr in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT: para anchor {substr!r} matched {len(hits)} (want 1).")
    return hits[0]


def _find_table(d, pred, tag):
    hits = [t for t in d.tables if _safe(pred, t)]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: table anchor matched {len(hits)} (want 1).")
    return hits[0]


def _safe(pred, t):
    try:
        return pred(t)
    except Exception:
        return False


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    if any("Table 3.2-1" in p.text for p in paras):
        raise SystemExit("ABORT: Tier-1 fix appears already applied (Table 3.2-1 present).")

    plan = []

    # -- G2: P217 heading Table 3.1-1 -> 3.2-1 --------------------------------
    p217 = _find_para(paras, "A.3 Live Load (KDS 41 12 00")
    _repl_run(p217, 0, "Table 3.1-1", "Table 3.2-1", "G2/P217")
    plan.append("[G2]  P217  Table 3.1-1 -> 3.2-1")

    # -- SNOW-CLAUSE: P232 §6 -> §4 -------------------------------------------
    p232 = _find_para(paras, "A.6 Snow Load (KDS 41 12 00")
    _repl_run(p232, 0, SECT + "6)", SECT + "4)", "SNOW-CLAUSE/P232")
    plan.append("[SNOW-CL] P232  (KDS 41 12 00 %s6) -> %s4)" % (SECT, SECT))

    # -- SNOW-01: P237 companion combos ---------------------------------------
    p237 = _find_para(paras, "included as a companion action in the gravity and seismic")
    _repl_run(p237, 0,
              "in the gravity and seismic combinations",
              "in the gravity, wind, and seismic combinations", "SNOW-01/P237")
    plan.append("[SNOW-01] P237  'gravity and seismic' -> 'gravity, wind, and seismic'")

    # -- G1: P216 disclosure prepend ------------------------------------------
    p216 = _find_para(paras, "The MEP allowance (A5) is a workflow constant")
    r216 = p216.runs[0]
    if r216.text.startswith("The reinforced-concrete unit weight is adopted"):
        raise SystemExit("ABORT: P216 already carries RC disclosure.")
    r216.text = P216_PREPEND + r216.text
    plan.append("[G1]  P216  prepend RC 24.5->24.0 building-value disclosure")

    # -- N1: Table 8 wind value 100.0/120.0 -> 93.1/111.7 ---------------------
    t8 = _find_table(d, lambda t: [c.text for c in t.rows[0].cells][:3] == ["Category", "Parameter", "Value"]
                     and any(r.cells[0].text == "Wind" for r in t.rows), "N1/T8")
    wind_row = next(r for r in t8.rows if r.cells[0].text == "Wind")
    vcell = wind_row.cells[2]
    _repl_run(vcell.paragraphs[0], 0, "100.0 / 120.0 ", "93.1 / 111.7 ", "N1/T8-wind")
    plan.append("[N1]  Table 8  Wind 100.0/120.0 -> 93.1/111.7 kN")

    # -- G1: Table A-1 (T11) R1 clause cell -----------------------------------
    t11 = _find_table(d, lambda t: len(t.rows) >= 2 and t.rows[1].cells[0].text == "A1"
                      and "RC unit weight" in t.rows[1].cells[1].text, "G1/T11")
    clause_cell = t11.rows[1].cells[2]
    _repl_run(clause_cell.paragraphs[0], 0,
              "KDS 24 12 21 Table 4.2-1", "KDS 24 12 21 Table 4.2-1 (building-adjusted)", "G1/T11R1")
    plan.append("[G1]  Table A-1 R1  clause -> '...Table 4.2-1 (building-adjusted)'")

    # -- G2: Table A-2 (T12) B1/B2 clause cells Table 3.1-1 -> 3.2-1 ----------
    t12 = _find_table(d, lambda t: len(t.rows) >= 3 and t.rows[1].cells[0].text == "B1"
                      and "LL, retail" in t.rows[1].cells[1].text, "G2/T12")
    for ridx in (1, 2):
        cc = t12.rows[ridx].cells[2]
        _repl_run(cc.paragraphs[0], 0, "Table 3.1-1", "Table 3.2-1", f"G2/T12R{ridx}")
    plan.append("[G2]  Table A-2 B1/B2  Table 3.1-1 -> 3.2-1")

    # -- SNOW-01: Table A-7 (T17) Wind X/Y combos add +0.5S -------------------
    t17 = _find_table(d, lambda t: [c.text for c in t.rows[0].cells] == ["Group", "Count", "Combinations"], "SNOW-01/T17")
    for w in ("WX", "WY"):
        wr = next(r for r in t17.rows if r.cells[0].text.startswith("Wind") and w in r.cells[2].text)
        _repl_run(wr.cells[2].paragraphs[0], 0,
                  f"{w}; 0.9 DL", f"{w} + 0.5 S; 0.9 DL", f"SNOW-01/T17-{w}")
    plan.append("[SNOW-01] Table A-7 Wind X/Y  '± 1.0 W;' -> '± 1.0 W + 0.5 S;'")

    print("Planned Tier-1 edits:")
    for line in plan:
        print("  " + line)

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply.")
        print("  P216 head:", p216.runs[0].text[:120])
        print("  T8 wind  :", [wind_row.cells[i].text for i in range(3)])
        print("  T11R1    :", [t11.rows[1].cells[i].text for i in range(3)])
        print("  T12R1    :", [t12.rows[1].cells[i].text for i in range(4)])
        print("  T17 WindX:", t17.rows[4].cells[2].text)
        print("  T17 WindY:", t17.rows[5].cells[2].text)
        print("  P217     :", p217.text)
        print("  P232     :", p232.text)
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
