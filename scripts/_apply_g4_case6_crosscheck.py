"""G4 - insert the Case 6 (L-shape) OpenSeesPy-vs-ETABS cross-check into §5.3, and
reconcile the section's closing sentence (which currently says the example is shown
'rather than through numerical comparison against a commercial reference').

Author-approved scoped framing: claim global-response + total-reaction agreement,
attribute the re-entrant corner/boundary gravity-reaction gap to a tributary-area
decomposition convention, and reserve per-member forces (sign conventions + the
setback-column end moment) for the dedicated validation campaign.

Backup -> ..._v2.pre_g4_backup.docx. Run with --apply (default dry-run).
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_g4_backup.docx")

CROSSCHECK = (
    "As an independent cross-check of the irregular-plan model, the same L-shaped "
    "configuration was also analyzed in ETABS under the four load cases, with both "
    "programs assigned identical member sections, supports, and the same "
    "tributary-distributed gravity loads. The global response agrees closely: the "
    "Zone A roof displacements, the maximum inter-story drifts, and the inter-zone "
    "torsional displacement difference match to within 0.1% under both the X- and "
    "Y-direction equivalent-static cases (for example, 36.5 mm versus 36.5 mm for the "
    "Zone A fifth-floor X-displacement and 14.24 mm versus 14.24 mm for the X-direction "
    "torsional difference), and the total vertical base reaction matches to 0.0% "
    "(3,458 kN dead load in both programs). The differences are confined to the "
    "per-support gravity-reaction allocation at the re-entrant boundary: the Zone A "
    "corner support and the shared-boundary support carry 20.6% and 8.5% more "
    "dead-load reaction in the open-source model (195.5 versus 155.2 kN and 269.7 "
    "versus 246.8 kN), while the far Zone B corner agrees to 0.2% (90.4 versus "
    "90.3 kN). The same ratios hold for live load, and because the total reaction and "
    "the full lateral response are preserved, these differences reflect the "
    "tributary-area decomposition convention at the L-shape re-entrant corner rather "
    "than the underlying analysis. A full per-member force comparison — including the "
    "column local-axis moment-sign conventions and the setback-column end moments — is "
    "reserved for the dedicated irregular-plan validation campaign (Section 6.3)."
)

ANCHOR = "the total vertical base reaction is 3,458"

REPLACE_148 = [
    ("from the benchmark validation reported in Section 4, this example",
     "from the five-case benchmark of Section 4, this example"),
    ("through end-to-end execution and gravity-equilibrium verification rather than "
     "through numerical comparison against a commercial reference.",
     "through end-to-end execution, gravity-equilibrium verification, and the ETABS "
     "cross-check of global response reported above."),
]


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    if any("also analyzed in ETABS under the four load cases" in p.text for p in paras):
        print("ABORT: cross-check paragraph already present.")
        return 1

    anchors = [p for p in paras if ANCHOR in p.text]
    if len(anchors) != 1:
        print(f"ABORT: gravity anchor matched {len(anchors)} paras (want 1).")
        return 1
    gravity_p = anchors[0]

    c148 = [p for p in paras if "Distinct in scope from the benchmark validation reported" in p.text]
    if len(c148) != 1:
        print(f"ABORT: closing para matched {len(c148)} paras (want 1).")
        return 1
    closing = c148[0]
    # verify both old substrings present
    for old, _ in REPLACE_148:
        if old not in closing.text:
            print(f"ABORT: closing-para substring not found: {old[:50]!r}")
            return 1

    # build new paragraph matching the §5.3 body style
    new_p = d.add_paragraph(style="본문_1")
    new_p.add_run(CROSSCHECK)
    # move it directly after the gravity paragraph
    gravity_p._p.addnext(new_p._p)

    # reconcile closing sentence (single run)
    for old, new in REPLACE_148:
        for r in closing.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
                break

    if not apply:
        print("DRY-RUN: would insert cross-check para after gravity para, and edit closing para.")
        print("  closing now starts:", closing.text[:130])
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"backup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
