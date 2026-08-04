# -*- coding: utf-8 -*-
"""Re-embed Figs 2/3/5/6 at full source resolution (fix docx downscaling).

The manuscript embeds these four matplotlib figures downscaled to 130-220
effective DPI (pixelated, esp. Fig 3 distribution @130 and Fig 6 screening @167).
The hi-res generator outputs already exist; this swaps the embedded image blob for
the full-res source, keeps the display WIDTH, recomputes height from the (preserved)
aspect. Field-safe: only image parts are replaced — no paragraph/run/field XML is
touched. Per-target byte-length guard + aspect guard; backup; verify. Dry-run default.

  shape[1] Fig2 parity        <- fig6_parity_plot.png          (1690x1774)  ~412 DPI
  shape[2] Fig3 distribution  <- fig7_ok_check_distribution.png(3256x1417)  ~473 DPI
  shape[4] Fig5 deformed      <- fig4_ifc_deformed.png         (2773x1949)  ~462 DPI
  shape[5] Fig6 screening     <- fig5_screening_summary.png    (3270x1428)  ~630 DPI

Usage: python scripts/_reembed_hires_figs_2026-07-09.py [--apply]
"""
from __future__ import annotations

import re
import shutil
import struct
import sys
import zipfile
from pathlib import Path

import docx
from docx.shared import Emu

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs" / "paper1_open_source_alternative"
DOC = BASE / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_reembed_2026-07-09.docx")
FIGS = BASE / "figures"

# (shape_idx, manuscript label, source png, expected current embedded byte-length)
TARGETS = [
    (1, "Fig2 parity",       FIGS / "fig6_parity_plot.png",           188220),
    (2, "Fig3 distribution", FIGS / "fig7_ok_check_distribution.png",  42375),
    (4, "Fig5 deformed",     FIGS / "fig4_ifc_deformed.png",          290293),
    (5, "Fig6 screening",    FIGS / "fig5_screening_summary.png",      41252),
]


def png_size(b: bytes) -> tuple[int, int]:
    assert b[12:16] == b"IHDR", "not a PNG"
    return struct.unpack(">II", b[16:24])


def field_counts(p: Path) -> tuple[int, int]:
    fld = instr = 0
    with zipfile.ZipFile(p) as z:
        for n in z.namelist():
            if n.endswith(".xml"):
                x = z.read(n).decode("utf-8", "replace")
                fld += len(re.findall(r"<w:fldChar\b", x))
                instr += len(re.findall(r"<w:instrText\b", x))
    return fld, instr


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    shapes = d.inline_shapes
    if len(shapes) != 8:
        raise SystemExit(f"ABORT: expected 8 inline shapes, got {len(shapes)}")

    plan = []
    for idx, label, src, old_len in TARGETS:
        if not src.exists():
            raise SystemExit(f"ABORT [{label}]: source missing {src}")
        sh = shapes[idx]
        part = d.part.related_parts[sh._inline.graphic.graphicData.pic.blipFill.blip.embed]
        cur = part._blob
        cw, ch = png_size(cur)
        nb = src.read_bytes()
        nw, nh = png_size(nb)
        if len(cur) != old_len:
            raise SystemExit(
                f"ABORT [{label}]: shape[{idx}] current bytes {len(cur)} != guard {old_len}")
        # aspect guard: new source must be the same figure (aspect within 1.5%)
        if abs((cw / ch) - (nw / nh)) / (cw / ch) > 0.015:
            raise SystemExit(
                f"ABORT [{label}]: aspect mismatch embedded {cw}x{ch} vs source {nw}x{nh}")
        w_emu = int(sh.width)
        new_h_emu = int(w_emu * nh / nw)
        eff_dpi = nw / (w_emu / 914400)
        plan.append((idx, label, part, nb, sh, new_h_emu,
                     f"{cw}x{ch}->{nw}x{nh}px  ~{eff_dpi:.0f} DPI  "
                     f"h {int(sh.height)/914400:.2f}->{new_h_emu/914400:.2f}in"))

    print("Planned re-embeds:")
    for idx, label, *_rest, info in plan:
        print(f"  shape[{idx}] {label:18s} {info}")

    if not apply:
        print("\nDRY-RUN. Re-run with --apply.")
        return 0

    fb, ib = field_counts(DOC)
    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    for idx, label, part, nb, sh, new_h_emu, _info in plan:
        part._blob = nb
        sh.height = Emu(new_h_emu)
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")

    # verify
    fa, ia = field_counts(DOC)
    d2 = docx.Document(str(DOC))
    print(f"[verify] inline_shapes={len(d2.inline_shapes)} paragraphs={len(d2.paragraphs)} "
          f"tables={len(d2.tables)}  fields fldChar {fb}->{fa} instrText {ib}->{ia} "
          f"[{'OK' if (fb, ib) == (fa, ia) else 'CHANGED!!'}]")
    for idx, label, *_ in plan:
        sh = d2.inline_shapes[idx]
        b = d2.part.related_parts[sh._inline.graphic.graphicData.pic.blipFill.blip.embed]._blob
        print(f"  shape[{idx}] {label:18s} now {png_size(b)} px, {len(b)} bytes")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
