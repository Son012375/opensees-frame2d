"""논문 초안 docx 정정 패치 (V2 누적: §4 정합 + §5.1 재작성).

원본(open_source_alternative_review_draft.docx)에서 다음을 일괄 수정해
_v2.docx로 저장한다.

수정 근거: 본 세션의 코드 정정 + framing 재설정 결과를 반영
  1. Fa/Fv 선형보간 누락 (KDS 17 10 00 §4.2.1 ②) 수정
  2. Cu 선형보간 누락 (KDS 41 17 00 §7.2.3 표 7.2-1) 수정
  3. Vu 추출 버그 (eleForce → eleResponse 'localForce') 수정 [별도 세션]
  → V_base 213 → 239.8 kN, max member ratio 0.193 (beam) → 0.507 (column)
  4. §5.1 framing 재작성: feasibility hedging → defensible position with
     prior-literature foundation + Appendix A hand-check + workflow
     simplification disclosure

원본은 보존하고 _v2.docx로 저장. 재실행하면 _v2.docx를 덮어쓴다.

Usage:
    python scripts/_patch_paper_draft.py
"""
from __future__ import annotations

import shutil
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft.docx"
DST = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"


def _replace_run(run, old: str, new: str) -> None:
    """Run 내 정확한 부분문자열을 교체. 없으면 AssertionError."""
    assert old in run.text, f"expected substring {old!r} in run, got {run.text!r}"
    run.text = run.text.replace(old, new)


def main() -> int:
    shutil.copy2(SRC, DST)
    d = docx.Document(str(DST))

    edits: list[dict] = []

    # ─── Title (para 0) — capability scope 확장 (regular → regular + orthogonal irregular)
    # 직전 "Regular Steel Frame Analysis" 좁힘 결정을 capability scope 확장에 맞춰 재조정.
    # Validation against Midas Gen → Benchmark against Midas Gen (validation은 regular cases sub-scope임을 암시)
    p0 = d.paragraphs[0]
    NEW_TITLE = (
        "Toward an Open-Source BIM-to-Analysis Workflow for Steel Frame "
        "Buildings: Node-Element IFC Parsing, KDS-Based Load Automation, "
        "and OpenSeesPy Benchmark against Midas Gen"
    )
    p0.runs[0].text = NEW_TITLE
    # Clear any additional runs that might carry stale title fragments
    for r in p0.runs[1:]:
        r.text = ""
    edits.append({"loc": "Title (para 0)",
                  "what": "Regular Steel Frame Analysis → BIM-to-Analysis Workflow for Steel Frame Buildings (capability scope 확장); Validation against → Benchmark against (validation 어휘 약화)"})

    # ─── Abstract rewrite (paras 7, 8, 9) ───────────────────
    ABS_P1 = (
        "Commercial structural-analysis programs remain central to building design "
        "practice, yet their closed preprocessing and result-recovery pipelines can "
        "limit independent inspection, reproduction, and customization of the "
        "analysis path. This paper presents an open-source pipeline that links "
        "node-element IFC parsing, Korean Design Standard (KDS) load automation, "
        "and OpenSeesPy three-dimensional frame analysis for steel frame buildings, "
        "supporting regular orthogonal grids and orthogonal irregular plan "
        "configurations (such as L-shape and setback) through zone-based "
        "decomposition, and evaluates it against Midas Gen as a commercial "
        "reference baseline for the regular cases reported here."
    )
    d.paragraphs[7].runs[0].text = ABS_P1

    ABS_P2 = (
        "Node-element IFC parsing follows the methodology established in prior "
        "BIM-to-FEM work and produces an inspectable analysis graph containing "
        "nodes, line elements, supports, sections, materials, and validation "
        "diagnostics. KDS-based load automation generates dead, live, equivalent "
        "lateral seismic, and wind cases and assembles KDS load combinations; every "
        "generated value is traced to its originating regulation clause through a "
        "hand-check appendix, with the small number of workflow-introduced "
        "simplifications disclosed explicitly. The pipeline's three-dimensional "
        "results are benchmarked against Midas Gen across five controlled cases "
        "spanning two-dimensional frames, a three-dimensional moment frame, and a "
        "supplementary geometrically nonlinear case, comparing 112 response metrics "
        "covering displacements, story drifts, reactions, and member forces."
    )
    d.paragraphs[8].runs[0].text = ABS_P2

    ABS_P3 = (
        "Of the 112 metrics, 100 agree with Midas Gen within 1% and the remaining "
        "12 within 4%, with the larger discrepancies concentrated in "
        "element-formulation-sensitive quantities of a single three-dimensional "
        "case rather than in global equilibrium. An IFC-derived three-story "
        "regular application example, together with a node-element L-shaped "
        "irregular-plan demonstration, exercises the full pipeline end-to-end "
        "and returns results consistent with engineering expectations under "
        "KDS-generated loading. Within the benchmark-validated scope — regular "
        "steel frames under elastic three-dimensional analysis — the pipeline "
        "serves as a credible, transparent, KDS-traced alternative computational "
        "path alongside the commercial baseline; orthogonal irregular plans are "
        "supported and demonstrated as workflow-execution cases rather than "
        "benchmark-validated cases."
    )
    d.paragraphs[9].runs[0].text = ABS_P3

    edits.append({"loc": "Abstract (paras 7-9)",
                  "what": "3 paragraphs rewritten (drop self-novelty claim, surface 3 integrated capabilities + hand-check appendix, close consistent with §5.1)"})

    # ─── §1 Introduction rewrite (paras 14–21) ──────────────
    # Phase A rewrites + B6 global trim pass (trimmed where length/density allowed).

    # §1 ¶1 motivation (B6 trim: 692c → ~540c)
    SEC1_P1 = (
        "Building Information Modeling (BIM) is now routine in architectural, "
        "engineering, and construction practice, yet the transfer from a "
        "design-oriented BIM model to an analysis-ready structural model remains "
        "a persistent bottleneck. IFC is a vendor-neutral exchange format for "
        "building geometry, materials, spatial hierarchy, and structural "
        "components, but IFC objects are authored for coordination rather than "
        "finite-element analysis. Structural engineers must still interpret "
        "member centerlines, connectivity, boundary conditions, section "
        "assignments, and analysis loads before a reliable computational model "
        "can be assembled."
    )
    d.paragraphs[14].runs[0].text = SEC1_P1

    # §1 ¶2 transition (B6 trim: 777c → ~660c)
    SEC1_P2 = (
        "This interpretation step is especially important for frame buildings, "
        "where a coordination model stores beams, columns, and slabs as building "
        "objects while a structural solver requires nodes, elements, degrees of "
        "freedom, loads, constraints, and sign conventions. When the conversion "
        "is performed manually or through proprietary preprocessing tools, the "
        "resulting workflow can be difficult to inspect, reproduce, or extend "
        "outside the originating environment. Commercial structural-analysis "
        "platforms such as Midas Gen provide capable, well-validated "
        "environments; positioning an open-source pathway against such a "
        "platform therefore requires evidence at the level of individual "
        "response quantities rather than capability claims."
    )
    d.paragraphs[15].runs[0].text = SEC1_P2

    SEC1_P3 = (
        "Prior BIM-to-analysis studies have established methodology for IFC-based "
        "structural model interpretation, finite-element model generation, and "
        "OpenBIM workflows (Ramaji & Memari, 2018; Hasan et al., 2019; Leonardi "
        "et al., 2024; Rudenko & Petryna, 2025). The present work builds on this "
        "established methodology rather than reformulating it. Three capabilities, "
        "however, are not yet jointly available in a single transparent toolchain "
        "for the Korean regulatory context. First, the IFC-derived structural "
        "graph is rarely paired with regional code-based load automation that "
        "traces each generated value to its originating clause. Second, "
        "integration with the Korean Design Standards (KDS) in an open-source "
        "frame-analysis pipeline remains limited. Third, metric-level comparison "
        "against a widely used commercial reference baseline is rarely reported "
        "at the response-quantity level for the resulting workflow as a whole. "
        "Their joint availability supports the paper's main evidentiary "
        "contribution: response-quantity-level comparison with a commercial "
        "baseline for a transparent BIM-to-analysis workflow; none of the "
        "three is claimed as novel in isolation."
    )
    d.paragraphs[16].runs[0].text = SEC1_P3

    SEC1_P4 = (
        "OpenSees and OpenSeesPy provide a strong computational foundation for "
        "such workflows. OpenSees is an open-source structural and "
        "earthquake-engineering simulation framework (McKenna, 2011), and "
        "OpenSeesPy exposes its modeling and analysis capabilities in Python "
        "(Zhu et al., 2018). This makes it possible to script model generation, "
        "parameter studies, and post-processing in a fully inspectable "
        "environment. OpenSeesPy does not, however, by itself solve the "
        "BIM-to-analysis problem: nodes, elements, transformations, loads, "
        "constraints, and result extraction must still be generated by a reliable "
        "preprocessing pipeline."
    )
    d.paragraphs[17].runs[0].text = SEC1_P4

    SEC1_P5 = (
        "This paper presents an open-source pipeline that integrates the three "
        "capabilities identified above for the case of regular steel frame "
        "buildings. IFC structural members are extracted into a node-element "
        "analysis graph using established methodology, KDS-based load automation "
        "generates dead, live, equivalent lateral seismic, and wind cases "
        "together with KDS load combinations, and the resulting OpenSeesPy "
        "three-dimensional model is benchmarked against Midas Gen as a "
        "commercial reference. A hand-check appendix records every generated KDS "
        "value clause-by-clause and explicitly discloses the small number of "
        "workflow-introduced simplifications, so that the benchmark agreement "
        "reported in Section 3 is not the product of hidden assumptions."
    )
    d.paragraphs[18].runs[0].text = SEC1_P5

    SEC1_P6 = (
        "The benchmark-validated scope is intentionally bounded. Numerical "
        "benchmarking against Midas Gen is restricted to regular orthogonal "
        "steel frame buildings under elastic three-dimensional analysis, with "
        "auxiliary geometric-nonlinear results from a supplementary five-story "
        "case. The implemented workflow capability extends to orthogonal "
        "irregular plans (L-shape, T-shape, setback) through the zone-based "
        "decomposition described in Section 2.3.3; non-orthogonal grids, "
        "curved members, rotated zones, reinforced concrete or composite "
        "sections, dynamic response-spectrum or time-history analysis, "
        "nonlinear material constitutive behavior, and full code-compliance "
        "certification fall outside the implemented scope. Within the "
        "benchmark-validated scope, the pipeline is positioned not as a "
        "replacement for commercial design environments but as a transparent, "
        "KDS-traced computational alternative, benchmarked in this study "
        "against Midas Gen and inspectable enough for practitioners and "
        "researchers to verify, extend, or contest."
    )
    d.paragraphs[19].runs[0].text = SEC1_P6

    # §1 ¶7 contributions — framing shift: evidence first, integration as framework
    SEC1_P7 = (
        "The contributions of the paper are threefold. First, the workflow is "
        "benchmarked against Midas Gen across five controlled cases and 112 "
        "response metrics, providing quantitative evidence at the "
        "response-quantity level — 100 within 1% and the remaining 12 within 4% — that an "
        "open-source pipeline can reproduce commercial-baseline responses for "
        "the validated structural class. Second, the KDS load automation is "
        "verified clause-by-clause against KDS 41 12 00, KDS 41 17 00, and "
        "KDS 17 10 00 in a hand-check appendix, with the four "
        "workflow-introduced simplifications (MEP allowance, wind topographic "
        "factor, gust factor, and pressure coefficient) disclosed, so the "
        "benchmark agreement is not the product of hidden input assumptions. "
        "Third, the workflow itself is a transparent, inspectable open-source "
        "pipeline integrating node-element IFC parsing — following "
        "established BIM-to-FEM methodology — with KDS load automation and "
        "OpenSeesPy three-dimensional frame analysis; this integration is "
        "what makes the quantitative evidence and the clause-traced load "
        "automation reproducible and auditable, rather than a methodological "
        "contribution in isolation. The same pipeline supports orthogonal "
        "irregular plan configurations through zone-based decomposition, "
        "illustrated by the L-shaped node-element demonstration in Section 4.3."
    )
    d.paragraphs[20].runs[0].text = SEC1_P7

    SEC1_P8 = (
        "The remainder of this paper is organized as follows. Section 2 presents "
        "the proposed open-source workflow and node-element methodology, "
        "including zone-based decomposition for orthogonal irregular plans and "
        "the disclosure of workflow-introduced simplifications. Section 3 "
        "reports the benchmark validation against Midas Gen. Section 4 "
        "presents an IFC-derived regular application example (Sections 4.1–4.2) "
        "and a node-element L-shaped irregular-plan demonstration (Section 4.3). "
        "Section 5 discusses the "
        "position of the workflow within its benchmark-validated scope and its "
        "limitations, and Section 6 concludes the paper. Appendix A provides "
        "the clause-by-clause hand-check of the KDS load generation."
    )
    d.paragraphs[21].runs[0].text = SEC1_P8

    edits.append({"loc": "§1 Introduction (paras 15-21)",
                  "what": "7 paragraphs rewritten: tone-shift on Midas, explicit prior-lit citations, integration framing, 4→3 contributions, Appendix A in outline. ¶1 (para 14) unchanged."})

    # ─── §2.4.1 workflow simplifications disclosure (para 48) ───
    # §5.1과 Abstract가 인용하는 §2.4 disclosure를 본문에 실체화.
    # 기존 2-sentence 위치 단락에 4가지 simplification (MEP, Kzt, Gf, Cp) + Appendix A pointer 추가.
    SEC2_4_1_DISCLOSURE = (
        "In this study, KDS automation is treated as reproducible analysis input "
        "generation, not as final code-compliance certification. This distinction "
        "avoids overstating the design authority of the workflow while still "
        "demonstrating that regional code-based loading can be incorporated into an "
        "open-source analysis pipeline. Four simplifications introduced by the "
        "workflow itself are disclosed here so that they are not absorbed silently "
        "into the generated values: (i) a mechanical/electrical/plumbing (MEP) "
        "allowance of 0.5 kN/m² applied to all floors as a workflow default in "
        "lieu of fixture-specific computation; (ii) a wind topographic factor Kzt "
        "of 1.0 corresponding to flat-terrain conditions; (iii) a wind gust factor "
        "Gf taken as a fixed value per exposure category rather than computed "
        "through the full KDS 41 12 00 §5 expression that depends on building "
        "geometry, period, and damping; and (iv) a total wind pressure coefficient "
        "Cp of 1.3 representing 0.8 (windward) + 0.5 (leeward) without aspect-ratio "
        "refinement from the KDS pressure-coefficient table. A complete "
        "clause-by-clause hand-check of the generated values against KDS 41 12 00, "
        "KDS 41 17 00, and KDS 17 10 00 — including each of these "
        "simplifications — is provided in Appendix A."
    )
    d.paragraphs[48].runs[0].text = SEC2_4_1_DISCLOSURE
    edits.append({"loc": "§2.4.1 disclosure (para 48)",
                  "what": "4 workflow simplifications (MEP/Kzt/Gf/Cp) explicit + Appendix A pointer"})

    # ─── §2.4.2 body (para 52) — append compat sentence for irregular plans ──
    SEC2_4_2_BODY = (
        "The structural analysis module assembles a three-dimensional "
        "OpenSeesPy frame model using the nodes and elements stored in the "
        "StructuralModel. Nodes are modeled with six degrees of freedom, and "
        "frame members are represented using elastic beam-column elements "
        "with appropriate geometric transformations, support constraints, "
        "and optional member releases. The same model-assembly procedure "
        "applies to both regular grids and the zone-decomposed irregular "
        "plans of Section 2.3.3; the OpenSeesPy formulation does not require "
        "special handling because zone boundary node merging produces a "
        "single connected node-element graph."
    )
    d.paragraphs[52].runs[0].text = SEC2_4_2_BODY
    edits.append({"loc": "§2.4.2 body (para 52)",
                  "what": "Append compat sentence: same model-assembly for regular + zone-decomposed irregular plans (links to new §2.3.3)"})

    # ─── §2.5.2 close (para 64) — remove NL/RAG/chatbot leftover sentence ──
    # 본문 어디서도 NL/RAG/chatbot을 다루지 않으므로 이를 "excluded ... reserved for future work"
    # 로 §2 본문에서 언급하는 것은 §5.2 ¶6 제거와 동일한 scope-confusion 패턴. 마지막 문장만 제거.
    SEC2_5_2_CLOSE = (
        "The model-generation and response-recovery portions of the workflow "
        "are validated in Section 3 through a Midas Gen commercial-baseline "
        "comparison; the KDS load automation is hand-checked in Appendix A."
    )
    d.paragraphs[64].runs[0].text = SEC2_5_2_CLOSE
    edits.append({"loc": "§2.5.2 close (para 64)",
                  "what": "NL/KDS-RAG/chatbot 'reserved for future work' 문장 제거 (§5.2 ¶6 제거와 동일 사유 — 본문이 다루지 않는 기능을 'excluded'로 언급하면 scope confusion)"})

    # ─── §3.3.1 Case 1 closed-form 3-way rewrite (para 83) ──
    # A-3 처방: Midas 순환논리 해소. 단순보 PL³/(48EI), PL/4, P/2 해석해 검증 추가.
    # Para 83은 2 runs였으나, 신규 text는 runs[0]에 consolidate, runs[1]은 비움.
    SEC3_3_1 = (
        "Across the three two-dimensional cases, all 37 metrics fall within the "
        "OK category, with maximum relative differences below 0.1%. Case 1, a "
        "simply supported beam of length L = 6.0 m under a midspan point load "
        "P = 60 kN, also admits an exact closed-form solution from elementary "
        "beam theory — midspan deflection δ = PL³/(48EI), maximum bending moment "
        "M_max = PL/4, and support reactions R_A = R_B = P/2 — and is therefore "
        "independently verified against the textbook values in addition to the "
        "Midas Gen comparison. The three-way comparison is presented in "
        "Table 4a, with relative differences below 0.001% on all three "
        "quantities; OpenSees reproduces the closed-form solution, and Midas "
        "Gen does so independently, removing the circularity of using a "
        "commercial program as the sole reference for a problem with a known "
        "analytical answer. Case 2, which introduces lateral loading and "
        "moment-frame interaction at beam-column joints, produces displacement "
        "and member-force results that remain within 0.1% of the reference "
        "values. Case 3 extends the comparison to a three-story frame under "
        "story-wise lateral forces; the resulting inter-story drift ratios and "
        "column base moments again remain within 0.1%. Taken together, these "
        "results show close numerical agreement for the tested two-dimensional "
        "linear cases."
    )
    p83 = d.paragraphs[83]
    p83.runs[0].text = SEC3_3_1
    if len(p83.runs) > 1:
        p83.runs[1].text = ""
    edits.append({"loc": "§3.3.1 (para 83)",
                  "what": "Case 1 closed-form 3-way verification added (independent analytical baseline removes Midas-circular-logic critique); reference to Table 4a (inserted by _add_case1_table5a.py)"})

    # ─── §3.3.4 Summary (para 89) — Phase A framing alignment ──
    # "replacement potential" → §5.1 verbatim "alternative computational path alongside the commercial baseline"
    SEC3_3_4 = (
        "Overall, the benchmark results show that the proposed open-source "
        "workflow can reproduce Midas Gen reference responses with close "
        "agreement across the tested cases. The agreement supports positioning "
        "the workflow as a defensible computational alternative for the tested "
        "regular steel-frame analyses, rather than as evidence of universal parity across "
        "structural classes and analysis types beyond that scope."
    )
    d.paragraphs[89].runs[0].text = SEC3_3_4
    edits.append({"loc": "§3.3.4 Summary (para 89)",
                  "what": "Phase A framing: 'replacement potential' → 'credible alternative computational path alongside the commercial baseline'"})

    # ─── §3.4 Discussion close (para 100) — §5.1 cross-reference ──
    SEC3_4_CLOSE = (
        "For the structural configurations and loading conditions tested in "
        "this study, all benchmark results fall within the OK or CHECK "
        "categories, and no metric reaches the FAIL threshold. These "
        "observations support the numerical consistency of the present "
        "workflow for the analyses tested, and underpin the discussion in "
        "Section 5.1 of how the workflow should be positioned relative to "
        "Midas Gen."
    )
    d.paragraphs[100].runs[0].text = SEC3_4_CLOSE
    edits.append({"loc": "§3.4 close (para 100)",
                  "what": "'open-source analysis pathway' → '§5.1 alternative computational path' verbatim cross-reference"})

    # ─── §4 도입부 (para 103) — claim-level 분리 명시 ────────
    # SCIE reviewer가 §4를 §3 benchmark의 연장으로 오독하지 않도록 첫 두 문장으로
    # claim level 분리. 기존 정보(3층 IFC Revit 2x3)는 세 번째 문장으로 흡수.
    SEC4_OPEN = (
        "This section presents an end-to-end execution demonstration of the "
        "open-source workflow on IFC-derived input. It is distinct in scope "
        "from the benchmark validation reported in Section 3. The application "
        "example uses a regular three-story steel frame building represented "
        "by an IFC model exported from Autodesk Revit in the IFC 2x3 schema "
        "(ISO 16739-1:2018)."
    )
    d.paragraphs[103].runs[0].text = SEC4_OPEN
    edits.append({"loc": "§4 opening (para 103)",
                  "what": "claim-level separation added at §4 opening (workflow demonstration vs §3 benchmark)"})

    # ─── §4.1 본문 (para 106 run 8) ─────────────────────────
    # "Cs = 0.118" → "Cs = 0.126"  AND  "V = 213.0 " → "V = 239.8 "
    p106 = d.paragraphs[106].runs[8]
    _replace_run(p106, "Cs = 0.118", "Cs = 0.126")
    _replace_run(p106, "V = 213.0 ", "V = 239.8 ")
    edits.append({"loc": "§4.1 body (para 106)", "what": "Cs 0.118→0.126; V 213.0→239.8 kN"})

    # ─── §4.2 결과 본문 (para 111 run 0) ─────────────────────
    p111 = d.paragraphs[111].runs[0]
    _replace_run(p111, "4.9 mm", "5.4 mm")
    _replace_run(p111, "10.8 mm", "12.0 mm")
    _replace_run(p111, "0.20%", "0.22%")
    _replace_run(p111, "0.42%", "0.47%")
    edits.append({"loc": "§4.2 results (para 111)", "what": "disp 4.9/10.8→5.4/12.0; drift 0.20/0.42→0.22/0.47"})

    # ─── §4.2 부재강도 본문 (para 112 run 1) ─────────────────
    # Vu 버그 수정 후 최종 narrative: beam-governed → column-governed
    p112 = d.paragraphs[112].runs[1]
    _replace_run(
        p112,
        "0.204 at a second-story X-direction beam under the H1-1b equation",
        "0.507 at a second-story column under the H1-1b equation (governed by 1.2DL + 1.0LL + 1.0EQY)"
    )
    _replace_run(
        p112,
        "The maximum column interaction ratio is 0.109, occurring at the third story",
        "The maximum beam interaction ratio is 0.322, occurring at the first story"
    )
    _replace_run(p112, "remain low", "remain within the allowable bound")
    edits.append({"loc": "§4.2 member ratio (para 112)",
                  "what": "0.204 (beam-gov) → 0.507 (column-gov, H1-1b, EQY); column 0.109 → beam 0.322"})

    # ─── Table 6 (table index 5): Seismic Base shear cell ───
    # row 9 col 2: "213.0 " → "239.8 " (Cs도 같이)
    t5_r9_c2 = d.tables[5].rows[9].cells[2].paragraphs[0]
    # runs[0]='213.0 ', runs[1]='kN', runs[2]=' (Cs = 0.118, Ta = 0.42 ', runs[3]='s', runs[4]=')'
    _replace_run(t5_r9_c2.runs[0], "213.0 ", "239.8 ")
    _replace_run(t5_r9_c2.runs[2], "Cs = 0.118", "Cs = 0.126")
    edits.append({"loc": "Table 6 (Seismic row)", "what": "213.0→239.8 kN; Cs 0.118→0.126"})

    # ─── Table 7 (table index 6): result rows ───────────────
    t6 = d.tables[6]
    _replace_run(t6.rows[4].cells[2].paragraphs[0].runs[0], "4.9 / 10.8 mm", "5.4 / 12.0 mm")
    _replace_run(t6.rows[5].cells[2].paragraphs[0].runs[0], "0.20% (ratio 0.098)", "0.22% (ratio 0.108)")
    _replace_run(t6.rows[6].cells[2].paragraphs[0].runs[0], "0.42% (ratio 0.212)", "0.47% (ratio 0.233)")
    # ─── §4 close (para 114) — drop Midas-IFC commitment, reaffirm scope separation ──
    SEC4_CLOSE = (
        "This application example demonstrates the full open-source execution "
        "path from IFC-derived node-element model construction to KDS load "
        "generation, OpenSeesPy analysis, and preliminary review output. "
        "Accordingly, the reported response quantities are interpreted as "
        "workflow demonstration results rather than additional "
        "commercial-baseline benchmark metrics."
    )
    d.paragraphs[114].runs[0].text = SEC4_CLOSE
    edits.append({"loc": "§4 close (para 114)",
                  "what": "Midas-IFC 'will be added' 약속 제거; '§3과 다른 claim level (workflow demonstration vs benchmark)' 재확인"})

    _replace_run(t6.rows[8].cells[2].paragraphs[0].runs[0], "0.204 (beam, Story 2, H1-1b)", "0.507 (column, Story 2, H1-1b)")
    _replace_run(t6.rows[9].cells[1].paragraphs[0].runs[0], "Max column ratio", "Max beam ratio")
    _replace_run(t6.rows[9].cells[2].paragraphs[0].runs[0], "0.109 (Story3)", "0.322 (Story 1)")
    edits.append({"loc": "Table 7 (rows 4-9)",
                  "what": "disp/drift/member aligned; row 8 column-governed; row 9 renamed Max column → Max beam"})

    # ─── §5.1 rewrite (paras 121-125) ───────────────────────
    # §5.1 header: keep numbering runs 0-2 ("5", ".", "1"), rewrite run 3 title
    p121 = d.paragraphs[121]
    _replace_run(p121.runs[3], " Interpretation of the Validated Scope",
                 " Position within the Validated Scope and Practical Significance")

    # §5.1 ¶1 (para 122)
    SEC5_1_P1 = (
        "For the structural class examined — regular orthogonal steel frame buildings "
        "analyzed elastically in three dimensions — the evidence in this paper supports "
        "a defensible position: the proposed open-source pipeline reproduces "
        "commercial-baseline response quantities closely enough to serve as a credible "
        "alternative computational pathway alongside Midas Gen for the analyses it "
        "covers. The benchmark comparison in Section 3 agrees with Midas Gen within 1% "
        "for 100 of 112 response metrics and within 4% for the remaining 12, with the "
        "differences concentrated in element-formulation-sensitive quantities of a "
        "single three-dimensional case. The IFC application example in Section 4 "
        "exercises the full chain end-to-end on a 348-element model assembled from a "
        "Revit IFC 2x3 export and returns drift, displacement, and member-strength "
        "results consistent with engineering expectations under KDS-generated loading."
    )
    p122 = d.paragraphs[122]
    p122.runs[0].text = SEC5_1_P1

    # §5.1 ¶2 (para 123) — integration as contribution
    SEC5_1_P2 = (
        "The practical significance is the close numerical match — "
        "quantified at the response-quantity level against Midas Gen as a "
        "commercial reference — together with the transparency that lets "
        "these results be inspected and reproduced. Node-element IFC parsing "
        "follows the methodology established in prior BIM-to-FEM work "
        "(Ramaji & Memari, 2018; Hasan et al., 2019; Rudenko & Petryna, 2025) and "
        "produces an inspectable analysis graph rather than a black-box model. "
        "KDS-based load automation traces every generated quantity to the originating "
        "regulation clause, hand-checked against the Korean Design Standards in "
        "Appendix A; the four workflow-introduced simplifications (MEP allowance, "
        "wind topographic factor, gust factor, leeward-windward pressure coefficient) "
        "are disclosed in Section 2.4 rather than absorbed silently. "
        "Commercial-baseline benchmarking "
        "quantifies agreement across 112 response metrics, replacing visual or "
        "qualitative comparison common in similar BIM-to-analysis studies. None of "
        "these three components is novel in isolation; their joint availability "
        "is what supports interpreting these results as response-quantity-level "
        "evidence rather than as an isolated numerical comparison."
    )
    p123 = d.paragraphs[123]
    p123.runs[0].text = SEC5_1_P2

    # §5.1 ¶3 (para 124) — scope boundary
    # B6 trim: 784c → ~620c (combine framing sentences, drop redundant pipeline-recovery clause)
    SEC5_1_P3 = (
        "The pipeline is intentionally scoped. The benchmark-validated "
        "structural class is regular orthogonal steel frame buildings under "
        "elastic three-dimensional analysis, with auxiliary geometric-nonlinear "
        "results from a supplementary five-story case. Orthogonal irregular "
        "plan configurations (L-shape, T-shape, setback) are supported through "
        "zone-based decomposition and demonstrated in Section 4.3; their "
        "numerical benchmarking against a commercial reference is left to "
        "future work. Non-orthogonal grids and rotated zones, reinforced "
        "concrete or composite sections, dynamic analysis, nonlinear material "
        "behavior, and full code-compliance certification fall outside the "
        "implemented scope. Within the benchmark-validated scope, the workflow "
        "is positioned not as a replacement for commercial design environments "
        "but as a transparent, reproducible, KDS-traced alternative path — "
        "inspectable enough that practitioners and researchers can verify, "
        "extend, or contest its outputs."
    )
    p124 = d.paragraphs[124]
    p124.runs[0].text = SEC5_1_P3

    # §5.1 ¶4 (para 125) — application examples (dual: regular IFC + L-shape node-element)
    SEC5_1_P4 = (
        "The application examples in Section 4 are presented as end-to-end "
        "execution demonstrations under the same workflow as the benchmarked "
        "cases — an IFC-derived regular example (Sections 4.1–4.2) and a "
        "node-element L-shaped irregular-plan example (Section 4.3). The "
        "numerical validation against the commercial reference baseline remains "
        "scoped to the five controlled cases reported in Section 3; the "
        "demonstrations exercise the pipeline end-to-end on real BIM-derived "
        "and node-element data respectively but are not themselves additional "
        "benchmarks."
    )
    p125 = d.paragraphs[125]
    p125.runs[0].text = SEC5_1_P4

    edits.append({"loc": "§5.1 (paras 121-125)",
                  "what": "header retitled + 4 paragraphs rewritten (framing reset, prior-lit citations, Appendix A surfaced, replacement claim bounded)"})

    # ─── §5.2 ¶2 (para 129) — mandatory rewrite: §2.3와 모순 해소 + 비정형 capability 명시 ──
    # 기존 "grid-based model assembly" → §2.3의 node-element graph 와 직접 모순
    # 기존 "Irregular geometry (setbacks, ...) falls outside" → 사실 아님 (zone-based로 지원)
    SEC5_2_P2 = (
        "First, the benchmark-validated structural scope is restricted to "
        "regular orthogonal steel frames with elastic material behavior. The "
        "workflow implementation extends to orthogonal irregular plans "
        "(L-shape, T-shape, setback) through zone-based decomposition "
        "(Section 2.3.3), illustrated by the application example in "
        "Section 4.3, but quantitative numerical benchmarking against Midas "
        "Gen for such configurations is reserved for future work. "
        "Non-orthogonal grids, curved members, rotated zones, internal "
        "column removal within a rectangular zone, reinforced concrete or "
        "composite members, and nonlinear material constitutive models fall "
        "outside the implemented scope. These boundaries reflect both the "
        "controlled benchmark-validation scope adopted in this study and the "
        "current implementation choices, including the zone-based orthogonal "
        "model assembly and elastic beam-column element formulation."
    )
    d.paragraphs[129].runs[0].text = SEC5_2_P2
    edits.append({"loc": "§5.2 ¶2 (para 129)",
                  "what": "MANDATORY rewrite — §2.3 node-element graph와의 모순 해소 ('grid-based' 제거) + orthogonal irregular plans capability 명시 + benchmark vs implementation scope 분리"})

    # ─── §5.2 ¶4 (para 131) terminal period fix ─────────────
    p131 = d.paragraphs[131].runs[0]
    if not p131.text.endswith("."):
        p131.text = p131.text + "."
    edits.append({"loc": "§5.2 ¶4 (para 131)",
                  "what": "added missing terminal period"})

    # ─── §5.3 ¶1 (para 136) full rewrite ────────────────────
    # 5 future items → 4 (IFC-Midas commitment, hand-check done, chatbot dropped);
    # add RC/composite, dynamic, full code-compliance to match §6 ¶4 extensions.
    # B6 trim: 1,809c → ~1,500c (combine intro sentences, tighten parser/dynamic descriptions)
    SEC5_3_P1 = (
        "Several directions for future work follow from the limitations "
        "identified in Section 5.2 and the extensions noted in Section 6. "
        "First, parser robustness should be evaluated "
        "across multiple BIM authoring tools — Tekla, ArchiCAD, and "
        "additional Revit export profiles — and on larger regular steel frame "
        "models. Second, the irregular-plan capability demonstrated in "
        "Section 4.3 should be benchmarked against a commercial reference "
        "baseline through additional controlled cases covering L-shape, "
        "T-shape, and setback configurations, with attention to "
        "zone-boundary node merging, rigid-diaphragm constraint extension "
        "across zones, and tributary-area calculation in "
        "inclusion–exclusion form. Third, the element-formulation-sensitive response metrics "
        "observed in the primary three-dimensional benchmark case (Case 4) "
        "should be studied through a targeted ablation covering local-axis "
        "convention, beta angle, shear deformation (Euler-Bernoulli versus "
        "Timoshenko), rigid-offset treatment, and the rigid-diaphragm "
        "constraint. Fourth, extension to reinforced concrete and composite "
        "member generation requires section-property derivation from rebar "
        "and composite cross-sections, KDS reinforced-concrete design checks, "
        "and independent benchmark cases. Fifth, dynamic analysis support — "
        "response-spectrum and time-history analysis with KDS-conforming "
        "inputs — would extend the validated analysis class beyond elastic "
        "static cases. Finally, full KDS code-compliance procedures, "
        "including drift amplification with member-specific C_d, story "
        "stability checks, P-delta amplification, and torsional irregularity "
        "assessment, would convert the preliminary screening into a complete "
        "code-compliance pathway."
    )
    d.paragraphs[136].runs[0].text = SEC5_3_P1
    edits.append({"loc": "§5.3 ¶1 (para 136)",
                  "what": "5 items → 4 items (IFC-Midas commitment + hand-check done + chatbot dropped; add RC/composite + dynamic + full KDS code-compliance to match §6)"})

    # ─── §6 Conclusion rewrite (paras 139–142) ──────────────
    # 4 단락 모두 재작성: integration framing, Abstract verbatim tolerance,
    # §5.1 verbatim position word, chatbot extension 제거, Midas-IFC commitment.

    SEC6_P1 = (
        "This paper presented an open-source pipeline that integrates "
        "node-element IFC parsing — following the methodology established in "
        "prior BIM-to-FEM work — with KDS-based load automation and OpenSeesPy "
        "three-dimensional frame analysis for regular steel frame buildings, "
        "accompanied by a clause-by-clause hand-check of the generated KDS "
        "values against the originating regulation clauses."
    )
    d.paragraphs[139].runs[0].text = SEC6_P1

    SEC6_P2 = (
        "Across five controlled benchmark cases spanning two-dimensional "
        "frames, a three-dimensional moment frame, and a supplementary "
        "geometrically nonlinear case, 100 of 112 response metrics agreed with "
        "Midas Gen within 1% and the remaining 12 within 4%, with the larger "
        "discrepancies concentrated in element-formulation-sensitive "
        "quantities of a single three-dimensional case rather than in global "
        "equilibrium. An IFC-derived three-story application example exercised "
        "the full pipeline end-to-end and returned drift, displacement, and "
        "member-strength results consistent with engineering expectations "
        "under KDS-generated loading."
    )
    d.paragraphs[140].runs[0].text = SEC6_P2

    SEC6_P3 = (
        "Within the benchmark-validated scope — regular orthogonal steel frame "
        "buildings under elastic three-dimensional analysis — the results "
        "support the position that the pipeline serves as a credible, "
        "transparent, KDS-traced alternative computational path alongside "
        "the commercial baseline. The contribution lies primarily in the "
        "quantitative evidence supporting this position; the inspectable "
        "workflow makes that evidence reproducible and auditable rather than "
        "claiming methodological novelty for any single component. The same "
        "pipeline supports orthogonal irregular plan configurations through "
        "zone-based decomposition (illustrated in Section 4.3), with their "
        "numerical benchmarking against Midas Gen reserved for future work."
    )
    d.paragraphs[141].runs[0].text = SEC6_P3

    SEC6_P4 = (
        "Extensions beyond the validated scope — reinforced concrete and "
        "composite sections, dynamic response-spectrum or time-history "
        "analysis, nonlinear material constitutive behavior, broader IFC "
        "authoring-tool coverage, and full KDS code-compliance procedures — "
        "require additional development and independent validation."
    )
    d.paragraphs[142].runs[0].text = SEC6_P4

    edits.append({"loc": "§6 Conclusion (paras 139-142)",
                  "what": "4 paragraphs rewritten: integration framing, Abstract verbatim tolerances, §5.1 verbatim position word, chatbot dropped, Midas-IFC commitment added"})

    # ─── §5.2 ¶6 (para 133) deletion — MUST BE LAST ─────────
    # NL/RAG/chatbot disclosure removed (out of paper scope).
    # XML element removal shifts all subsequent paragraph indices by -1.
    p133_orig_text = d.paragraphs[133].text[:80]
    p133_element = d.paragraphs[133]._element
    p133_element.getparent().remove(p133_element)
    edits.append({"loc": "§5.2 ¶6 (para 133)",
                  "what": f"REMOVED — '{p133_orig_text}...' (out of paper scope; matches §6 chatbot drop)"})

    d.save(str(DST))

    print(f"[OK] Patched draft: {DST.relative_to(ROOT)}")
    print(f"     Original preserved: {SRC.relative_to(ROOT)}")
    print()
    print("Edits applied:")
    for i, e in enumerate(edits, 1):
        print(f"  {i}. {e['loc']:35s}  {e['what']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
