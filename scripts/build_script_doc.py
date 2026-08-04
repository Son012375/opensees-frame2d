# -*- coding: utf-8 -*-
"""LLM-AE-AI 중간발표 대본 (5분) 문서 생성"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(r"C:\Users\youm\Desktop\백장운_수업")
OUT_PATH = OUT_DIR / "LLM-AE-AI_중간발표_대본_손명국.docx"
FONT = "맑은 고딕"

NAVY = RGBColor(0x11, 0x2D, 0x4E)
BLUE = RGBColor(0x3F, 0x72, 0xAF)
GRAY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB0, 0x2A, 0x2A)


def set_font(run, size=10, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for k in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(k), FONT)


def add_para(doc, text="", size=10, bold=False, color=None, space_after=4, line_spacing=1.3, indent=0, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if indent:
        pf.left_indent = Cm(indent)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, color=color)
    return p


def add_slide_header(doc, slide_no, title, time_str, extra_note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "3F72AF")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(f"SLIDE {slide_no} — ")
    set_font(run, size=13, bold=True, color=BLUE)
    run = p.add_run(title)
    set_font(run, size=13, bold=True, color=NAVY)
    run = p.add_run(f"   · {time_str}")
    set_font(run, size=10, color=GRAY)
    if extra_note:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        run = p2.add_run(extra_note)
        set_font(run, size=9, color=RED, bold=True)


def add_script(doc, text):
    """대본 본문 — 인용 스타일."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "3F72AF")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_font(run, size=11)


def add_cue(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.6)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.25
    run = p.add_run("▶ " + text)
    set_font(run, size=9, color=GRAY, bold=False)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=15, bold=True, color=NAVY)


def add_table(doc, headers, rows, widths_cm=None, highlight_row=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=9, bold=True)
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "DCE6F1")
        tc_pr.append(shd)
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.2
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            bold = (highlight_row is not None and r - 1 in highlight_row)
            set_font(run, size=9, bold=bold)
    if widths_cm:
        for row in table.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
    return table


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # ── 표제 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LLM-AE-AI 중간발표 — 발표 대본 (5분)")
    set_font(run, size=17, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("3D Editor + KDS-RAG 설계 에이전트")
    set_font(run, size=11, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("손명국 · 2026310208 · 경희대학교 건축공학과 · 2026.05.12")
    set_font(run, size=10, color=GRAY)

    # ── 사용 안내 ──
    add_para(doc, "[ 사용 안내 ]", size=10, bold=True, color=BLUE, space_after=2)
    add_para(doc, "• 파란 좌측 라인이 그어진 인용 블록이 실제 발표 대본입니다.", size=9, color=GRAY, space_after=2)
    add_para(doc, "• ▶ 회색 표시는 행동 메모(강조 · 멈춤 · 시선 처리)로 발화하지 않습니다.", size=9, color=GRAY, space_after=2)
    add_para(doc, "• 빨간 메모는 해당 슬라이드의 발표 포인트입니다.", size=9, color=GRAY, space_after=6)

    # ──────────────── SLIDE 1 ────────────────
    add_slide_header(doc, 1, "표지", "약 10초", "인사 후 1초 멈춤, 청중과 시선 맞추기")
    add_script(doc, "안녕하십니까. 건축공학과 손명국입니다. 발표 주제는 3D Editor + KDS-RAG 설계 에이전트, IFC 기반 모델링과 KDS 기반 자동 설계 보조 시스템 개발 계획입니다.")

    # ──────────────── SLIDE 2 ────────────────
    add_slide_header(doc, 2, "CONTENTS", "약 5초", "빠르게 넘김")
    add_script(doc, "발표는 Background, Proposal, Method, Plan 네 부분으로 진행하겠습니다.")

    # ──────────────── SLIDE 3 ────────────────
    add_slide_header(doc, 3, "01 Background · 1.1 V2 현황과 한계", "약 35초")
    add_script(doc, "먼저 배경입니다. 현재 IFC 파서부터 KDS 자동 하중 생성, OpenSeesPy 3D 해석, Design Check까지의 3D Editor를 구현했고, Midas Gen 벤치마크 112개 지표 중 100개 일치, 약 89퍼센트 정확도로 해석 엔진을 검증한 상태입니다.")
    add_script(doc, "다만 세 가지 한계가 남아 있습니다. 첫째, Design Check 결과에 KDS 조항 본문이 표시되지 않아 보고서로 직결되기 어렵습니다. 둘째, NG가 나면 사용자가 단면 변경과 재해석을 직접 반복해야 합니다. 셋째, KDS 본문이 PDF로 파편화되어 비전공자 검색 비용이 큽니다.")
    add_script(doc, "한마디로, 검증된 3D Editor 위에 LLM Layer가 빠져 있는 상태입니다.")
    add_cue(doc, "\"89퍼센트\" 발음 시 톤 살짝 올려서 강조 · 마지막 한 줄에서 잠시 멈춤")

    # ──────────────── SLIDE 4 ────────────────
    add_slide_header(doc, 4, "01 Background · 1.2 NG 이후의 페인포인트", "약 25초")
    add_script(doc, "그중에서도 가장 큰 페인포인트는 NG 이후 구간입니다. 현재 흐름은 NG 발견부터 재해석까지 수 분에서 수 시간이 걸리지만, LLM Layer를 도입하면 자동 진단부터 1-click 적용까지 수십 초로 단축됩니다.")
    add_script(doc, "즉, 본 과제는 NG 이후의 시행착오를 LLM이 자동화하는 것이 목표입니다.")
    add_cue(doc, "\"수 시간 → 수십 초\" 대비 강조 · 좌우 다이어그램을 손으로 가리키며 전환")

    # ──────────────── SLIDE 5 ────────────────
    add_slide_header(doc, 5, "02 Proposal · 2.1 개념과 목표", "약 30초")
    add_script(doc, "다음은 제안입니다. 프로젝트명은 3D Editor + KDS-RAG 설계 에이전트, IFC 입력부터 NG 부재 진단, 제안, 재해석까지를 한 흐름으로 묶습니다.")
    add_script(doc, "구조는 세 축으로 구성됩니다. 검증된 V2 Editor가 substrate, 엔지니어가 결정 주체, AI Layer가 KDS 인용 기반 제안을 담당합니다. 범위는 Steel, Linear, KDS 41 12, 17, 31에 한정해 기말까지 완성도를 확보합니다.")
    add_script(doc, "핵심은, 엔지니어가 결정하고 AI가 근거를 댄다는 것입니다.")
    add_cue(doc, "마지막 한 줄 \"엔지니어가 결정하고 AI가 근거를 댄다\"는 천천히, 또박또박")

    # ──────────────── SLIDE 6 ────────────────
    add_slide_header(doc, 6, "02 Proposal · 2.2 4개 핵심 기능", "약 45초")
    add_script(doc, "핵심 기능은 네 가지입니다.")
    add_script(doc, "F1, KDS-RAG 챗봇은 임베딩과 BM25, RRF 하이브리드 검색으로 KDS 조항과 출처를 함께 응답합니다.")
    add_script(doc, "F2, Design Check 자동 인용은 검토 결과의 조항을 RAG로 찾아 Citations 형태로 보고서에 첨부합니다.")
    add_script(doc, "F3, 설계 제안 에이전트는 가장 핵심 기능으로, NG 부재를 진단해 단면 · 재료 후보를 찾고, MCP Tool로 모델을 수정한 뒤 재해석까지 자동 반복합니다.")
    add_script(doc, "F4, LLM-as-Judge Verifier는 F3 매 반복의 제안을 과설계, 환각, 기준 누락 관점에서 평가합니다.")
    add_script(doc, "네 기능을 한 단어로 정리하면 Diagnose, Cite, Suggest, Verify 입니다.")
    add_cue(doc, "\"F3는 가장 핵심 기능\"에서 잠깐 멈췄다가 진행 · 마지막 4개 영단어는 손으로 짚어가며")

    # ──────────────── SLIDE 7 ────────────────
    add_slide_header(doc, 7, "03 Method · 3.1 시스템 아키텍처", "약 45초", "★ 가장 중요한 슬라이드 — 화살표 따라 시선 유도")
    add_script(doc, "시스템 아키텍처입니다.")
    add_script(doc, "입력은 IFC가 메인, 자연어가 보조이고, 두 경로 모두 BuildingModel IR로 수렴합니다. 이후 KDS 하중 생성, OpenSeesPy 3D 해석, Design Check까지가 현재 V2 파이프라인입니다.")
    add_script(doc, "그 위에 AI Layer가 얹힙니다. Claude API를 통해 MCP Prompts로 5단계 프롬프트, Tool Use로 V2의 14개 MCP 도구, RAG로 KDS DB와 단면 DB 738건, Citations로 조항 출처를 결합합니다.")
    add_script(doc, "모든 AI 출력은 V2 Editor에 반영되고, 재해석 루프를 통해 OpenSeesPy 결과로 ground truth가 확보됩니다. 즉, 같은 V2 루프 위에 AI Layer만 접합하는 구조입니다.")
    add_cue(doc, "화살표 한 단계씩 짚으며 설명 · \"AI Layer가 얹힙니다\"에서 그림 위쪽 영역 손으로 지목")

    # ──────────────── SLIDE 8 ────────────────
    add_slide_header(doc, 8, "03 Method · 3.2 W1~W7 매핑", "약 35초", "★ 평가 가중치 25% 직격 슬라이드")
    add_script(doc, "수업 내용과의 연계입니다.")
    add_script(doc, "W3 LLM-as-Judge는 F4 Verifier에, W4 Tool Use는 단면 변경과 재해석 호출에, W5 RAG 파이프라인은 KDS 임베딩과 하이브리드 검색에, W6 Citations와 Prompt Caching은 출처 표시와 비용 절감에, W7 FastMCP는 14개 도구와 5단계 프롬프트 노출에 사용합니다.")
    add_script(doc, "결과적으로 7주차 중 5주차 — RAG, Features, MCP를 모두 능동 활용하는 구조입니다.")
    add_cue(doc, "\"7주차 중 5주차\" 강조 · 표의 음영 행을 손으로 가리키며")

    # ──────────────── SLIDE 9 ────────────────
    add_slide_header(doc, 9, "03 Method · 3.3 에이전트 5단계 프롬프트", "약 20초")
    add_script(doc, "방금 언급한 5단계 프롬프트입니다. Solver는 NG 원인 진단, Self-Improvement는 후보 다양화, Verifier는 KDS 적합성 평가, Correction은 모델 수정 명령, Synthesis는 채택안과 근거 정리를 담당합니다.")
    add_script(doc, "이는 W7 게스트 세미나의 표준 패턴을 그대로 구현한 것입니다.")
    add_cue(doc, "다섯 박스 화살표 순서대로 손으로 짚어가며 빠르게 진행")

    # ──────────────── SLIDE 10 ────────────────
    add_slide_header(doc, 10, "04 Plan · 4.1 데모 시나리오", "약 25초")
    add_script(doc, "마지막으로 데모 시나리오입니다.")
    add_script(doc, "Before에서는 240개 부재 중 116개가 NG, 최대 활용도 비 2.46으로 매우 위험한 상태입니다. 에이전트가 KDS 41 31 00 §H1.1 본문을 인용하며 단면 DB 738건에서 H-350×350을 제안하고, After에서는 240개 부재 전부 OK 판정에 도달합니다.")
    add_script(doc, "최종적으로 HTML Report와 CAD 도면 내보내기까지 완성된 구조물 기반 산출물을 제공할 예정입니다.")
    add_cue(doc, "Before / After 두 화면 비교 — \"240개 전부 OK\"에서 톤 살짝 올림")

    # ──────────────── SLIDE 11 ────────────────
    add_slide_header(doc, 11, "Conclusion", "약 15초", "★ 1초 멈춤 후 인사")
    add_script(doc, "결론입니다. 기존 3D Editor 위에 KDS-RAG 설계 에이전트를 얹어, NG 이후의 시행착오를 KDS 근거 기반으로 자동화하는 것이 본 과제의 핵심입니다.")
    add_script(doc, "위험 요소는 Verifier 검증, 발췌 인용과 출처 표기, F1 · F2 우선 구현으로 각각 대응합니다.")
    add_script(doc, "이상으로 발표를 마치겠습니다. 감사합니다.")
    add_cue(doc, "\"마치겠습니다\" 후 1초 멈춤 → 가벼운 목례 → \"감사합니다\"")

    # ──────────────── BACKUP SLIDE ────────────────
    add_h1(doc, "BACKUP SLIDE — 12 (Q&A 응답용, 발표에 미사용)")
    add_para(
        doc,
        "성능 평가 메트릭 질문 (\"평가 지표가 어떻게 되나요?\")이 나올 경우 슬라이드 12번을 띄우고 응답한다. 본 발표 5분에는 포함되지 않으므로 시간 영향 없음.",
        size=10, color=GRAY, space_after=6,
    )

    add_para(doc, "[ 슬라이드 12 디자인 ]", size=11, bold=True, color=BLUE, space_after=4)

    # 슬라이드 헤더 안내
    add_para(doc, "헤더 (좌상단)", size=10, bold=True, space_after=2)
    add_para(doc, "• 섹션명 \"Appendix\" (28pt, #112D4E)  /  번호 \"A1\" (44pt, #3F72AF)", size=10, indent=0.4, space_after=2)
    add_para(doc, "• 부제 \"성능 평가 메트릭 (Backup)\" (#112D4E)", size=10, indent=0.4, space_after=6)

    # 본문 박스 — 4개 기능 박스 + E2E + 평가셋
    add_para(doc, "본문 — 4개 기능 박스 + E2E 통합 박스 + 평가 셋 박스", size=10, bold=True, space_after=4)

    # F1~F4 박스 표로 표현
    add_table(
        doc,
        headers=["F1  KDS-RAG 챗봇", "F2  Design Check 인용", "F3  설계 제안 에이전트", "F4  LLM-as-Judge"],
        rows=[
            [
                "Top-5 Accuracy ≥ 0.95\nMRR ≥ 0.80\nLatency ≤ 3 s (cached)",
                "Hallucination Rate = 0\nCitation Precision ≥ 0.90\nCitation Recall ≥ 0.85",
                "Convergence Rate ≥ 80%\nTime-to-OK ≤ 60 s\nAvg Iter ≤ 2.5",
                "FPR ≤ 0.10\nCohen's Kappa ≥ 0.70\nSelf-Consistency ≥ 0.90",
            ],
        ],
        widths_cm=[4.0, 4.0, 4.0, 4.0],
    )

    add_para(doc, "", size=4, space_after=4)

    # E2E 박스
    add_para(doc, "E2E 통합 (가로로 긴 다크 네이비 박스, 흰 글씨)", size=10, bold=True, space_after=2)
    add_table(
        doc,
        headers=["IFC-to-OK 시간", "1회 실행 비용", "Cache Hit Rate", "User Click 감소"],
        rows=[
            ["≤ 5 min (10층)", "≤ $0.5 / run", "≥ 80%", "≥ 70%"],
        ],
        widths_cm=[4.0, 4.0, 4.0, 4.0],
    )

    add_para(doc, "", size=4, space_after=4)

    # 평가 셋 박스
    add_para(doc, "평가 셋 (작은 회색 박스, 하단 좌측)", size=10, bold=True, space_after=2)
    add_para(
        doc,
        "• KDS Gold Set 50~100건 (자연어 질의 ↔ 정답 조항)",
        size=10, indent=0.4, space_after=2,
    )
    add_para(doc, "• NG Case Set 10~20건 (압축/휨/상관식/drift NG 시나리오)", size=10, indent=0.4, space_after=2)
    add_para(doc, "• Human Baseline 5~10건 (수동 해결 시간/단면 선택 기록)", size=10, indent=0.4, space_after=2)
    add_para(doc, "• 구축 시기 : W2 ~ W4 (계획서 §7 일정과 병행)", size=10, indent=0.4, space_after=6)

    # 하단 메시지
    add_para(doc, "하단 한 줄 메시지 (20pt, #3F72AF)", size=10, bold=True, space_after=2)
    add_para(
        doc,
        "\"기능별 정량 측정 · Hallucination 0% 시스템 차단 · Human Baseline 대비 검증\"",
        size=11, color=BLUE, bold=True, indent=0.4, space_after=8,
    )

    # ASCII 레이아웃 미리보기
    add_para(doc, "[ 슬라이드 12 레이아웃 미리보기 ]", size=11, bold=True, color=BLUE, space_after=4)
    layout = (
        "┌────────────────────────────────────────────────────────────────────────┐\n"
        "│  Appendix                                                       A1    │\n"
        "│  성능 평가 메트릭 (Backup)                                              │\n"
        "├────────────────────────────────────────────────────────────────────────┤\n"
        "│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐               │\n"
        "│  │ F1 RAG   │  │ F2 Cite  │  │ F3 Agent │  │ F4 Judge │               │\n"
        "│  │          │  │          │  │          │  │          │               │\n"
        "│  │ Top-5≥95%│  │ Halluci  │  │ Conv≥80% │  │ FPR≤10%  │               │\n"
        "│  │ MRR≥0.80 │  │  = 0     │  │ TTO≤60s  │  │ Kappa≥0.7│               │\n"
        "│  │ Lat≤3s   │  │ Prec≥90% │  │ Iter≤2.5 │  │ SC ≥90%  │               │\n"
        "│  └──────────┘  └──────────┘  └──────────┘  └──────────┘               │\n"
        "│                                                                        │\n"
        "│  ╔══════════════════════════════════════════════════════════════════╗  │\n"
        "│  ║  E2E   IFC-to-OK ≤5min  ·  $0.5/run  ·  Cache≥80%  ·  Click-70% ║  │\n"
        "│  ╚══════════════════════════════════════════════════════════════════╝  │\n"
        "│                                                                        │\n"
        "│  [평가 셋]                                                              │\n"
        "│  • KDS Gold Set 50~100   • NG Case 10~20   • Human Baseline 5~10      │\n"
        "│                                                                        │\n"
        "│  기능별 정량 측정 · Hallucination 0% 시스템 차단 · Human Baseline 비교 │\n"
        "└────────────────────────────────────────────────────────────────────────┘\n"
    )
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(layout)
    set_font(run, size=8)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")

    # 디자인 메모
    add_para(doc, "[ 디자인 메모 ]", size=11, bold=True, color=BLUE, space_after=2)
    add_para(doc, "• 박스 4개 (F1~F4) — 둥근 모서리, 채우기 #FFFFFF, 테두리 #3F72AF (1.5pt)", size=10, indent=0.4, space_after=2)
    add_para(doc, "• 박스 제목 (F1 KDS-RAG 등) — 16pt, 굵게, #112D4E", size=10, indent=0.4, space_after=2)
    add_para(doc, "• KPI 값 — 14pt, 일반, #112D4E", size=10, indent=0.4, space_after=2)
    add_para(doc, "• E2E 가로 박스 — 채우기 #112D4E, 흰 글씨 16pt", size=10, indent=0.4, space_after=2)
    add_para(doc, "• 평가 셋 박스 — 채우기 #F2F2F2, 12pt #555555", size=10, indent=0.4, space_after=2)
    add_para(doc, "• 하단 메시지 — 20pt, #3F72AF, 굵게 (다른 슬라이드와 동일)", size=10, indent=0.4, space_after=8)

    # 대본
    add_para(doc, "[ Q&A 응답 대본 (약 50초) ]", size=11, bold=True, color=BLUE, space_after=4)
    add_para(
        doc,
        "교수님 또는 청중이 \"성능 평가 지표는 어떻게 되나요?\" 류 질문을 하면, 슬라이드 12번으로 이동한 뒤 다음을 읽는다.",
        size=9, color=GRAY, space_after=4,
    )
    add_script(
        doc,
        "성능 평가는 4개 기능별 정량 지표와 E2E 통합 지표 두 축으로 구성됩니다.",
    )
    add_script(
        doc,
        "F1 RAG 챗봇은 Top-5 정확도 95퍼센트 이상과 MRR 0.80 이상을 목표로, BM25와 임베딩 단독 대비 RRF 결합 우위를 정량 검증합니다.",
    )
    add_script(
        doc,
        "F2 Design Check 자동 인용은 Hallucination Rate를 0퍼센트로 시스템적으로 차단합니다 — retrieve된 청크 ID 외 조항은 인용 자체가 불가능한 구조입니다.",
    )
    add_script(
        doc,
        "F3 설계 제안 에이전트는 수렴률 80퍼센트 이상, NG에서 OK까지 시간 60초 이내, 평균 반복 수 2.5회 이하를 KPI로 둡니다. 본 발표에서 \"수 시간을 수십 초로\"라는 주장의 정량 근거입니다.",
    )
    add_script(
        doc,
        "F4 LLM-as-Judge Verifier는 사람 평가 대비 False Positive Rate 10퍼센트 이하, Cohen's Kappa 0.70 이상으로 안전성 직결 지표를 관리합니다.",
    )
    add_script(
        doc,
        "통합 지표로는 IFC 업로드부터 모든 부재 OK까지 5분 이내, 1회 실행 API 비용 0.5달러 이하, Prompt Caching hit rate 80퍼센트 이상을 목표로 합니다.",
    )
    add_script(
        doc,
        "평가 셋은 KDS Gold Set 50~100건, NG Case Set 10~20건, Human Baseline 5~10건으로 구축하며, 일정상 W2부터 W4까지 RAG 파이프라인 구축과 병행해 확보할 계획입니다.",
    )
    add_cue(doc, "전체 발화 약 50초. 청중 질문이 더 구체적이면 (예: \"Kappa는 왜 0.70?\") 해당 항목만 추가 설명")

    # 짧은 버전
    add_para(doc, "[ 짧은 버전 (20초, 시간 제한 시) ]", size=11, bold=True, color=BLUE, space_after=4)
    add_script(
        doc,
        "F1은 Top-K 정확도와 MRR, F2는 환각률 0퍼센트 시스템 차단, F3는 수렴률 80퍼센트와 NG-OK 시간 60초, F4는 사람 평가 대비 FPR 10퍼센트 이하를 목표로 합니다. 평가 셋은 KDS Gold Set 50건과 NG Case Set 10~20건을 W2~W4에 구축할 예정입니다.",
    )
    add_cue(doc, "Q&A 시간이 짧을 때 또는 메인 슬라이드 그대로 두고 답할 때")

    # ──────────────── 시간 점검표 ────────────────
    add_h1(doc, "시간 점검표 (누적)")
    add_table(
        doc,
        headers=["슬라이드", "내용", "구간", "누적", "비고"],
        rows=[
            ["1", "표지 · 인사", "10s", "0:10", ""],
            ["2", "목차", "5s", "0:15", "빠르게"],
            ["3", "V2 현황 · 한계 3가지", "35s", "0:50", ""],
            ["4", "NG 이후 페인포인트", "25s", "1:15", ""],
            ["5", "개념 · 목표", "30s", "1:45", ""],
            ["6", "4개 핵심 기능", "45s", "2:30", ""],
            ["7", "시스템 아키텍처", "45s", "3:15", "★ 핵심"],
            ["8", "W1~W7 매핑", "35s", "3:50", "★ 평가 25%"],
            ["9", "5단계 프롬프트", "20s", "4:10", ""],
            ["10", "데모 시나리오", "25s", "4:35", ""],
            ["11", "결론 · 마무리", "15s", "4:50", ""],
        ],
        widths_cm=[1.6, 6.0, 1.8, 1.8, 3.6],
        highlight_row=[6, 7],
    )
    add_para(doc, "예상 총합 : 약 4분 50초 + 슬라이드 전환 pause 10초 ≈ 5분", size=10, bold=True, color=BLUE)

    # ──────────────── 발표 팁 ────────────────
    add_h1(doc, "발표 직전 체크리스트")
    add_para(doc, "[ 빠르게 넘길 슬라이드 ]", size=10, bold=True, color=BLUE, space_after=2)
    add_para(doc, "• S2 목차 (5초) · S9 5단계 (20초) — 시각자료가 명확해 짧게 가능", size=10, indent=0.4, space_after=2)

    add_para(doc, "[ 천천히 강조할 슬라이드 ]", size=10, bold=True, color=BLUE, space_after=2)
    add_para(doc, "• S7 아키텍처 — 화살표 따라 시선 유도, 가장 복잡한 슬라이드", size=10, indent=0.4, space_after=2)
    add_para(doc, "• S8 W매핑 — 평가 25% 핵심 포인트, \"5주차 능동 활용\" 명확히 발음", size=10, indent=0.4, space_after=2)

    add_para(doc, "[ 숫자 강조 포인트 ]", size=10, bold=True, color=BLUE, space_after=2)
    add_para(doc, "• \"89퍼센트\" (S3 검증)  ·  \"수 시간 → 수십 초\" (S4 비교)", size=10, indent=0.4, space_after=2)
    add_para(doc, "• \"240개 중 116개 NG\" / \"활용도 비 2.46\" / \"H-350×350\" (S10 데모)", size=10, indent=0.4, space_after=2)
    add_para(doc, "• \"7주차 중 5주차\" (S8 매핑)", size=10, indent=0.4, space_after=4)

    add_para(doc, "[ 마무리 인사 ]", size=10, bold=True, color=BLUE, space_after=2)
    add_para(doc, "• \"마치겠습니다\" 후 1초 멈춤 → 가벼운 목례 → \"감사합니다\"", size=10, indent=0.4, space_after=2)

    add_para(doc, "[ Q&A 대비 ]", size=10, bold=True, color=BLUE, space_after=2)
    add_para(doc, "• 별도 문서 \"LLM-AE-AI_중간발표_QA_손명국.docx\" 참고", size=10, indent=0.4, space_after=2)
    add_para(doc, "• 가장 받기 쉬운 질문 TOP 3 : 일정 / Verifier 환각 누적 / 5단계 본인 기여", size=10, indent=0.4)

    doc.save(OUT_PATH)
    print(f"[OK] saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
