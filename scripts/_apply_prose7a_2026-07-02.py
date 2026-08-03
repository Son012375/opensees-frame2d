"""Session 2026-07-02 (MAJOR REVISION): prose batch 7a — LLM future-work + citations + numerals + P167.

  P165  run_replace: cite where the primitives are exercised (R6 partial-misread: primitives
        ARE tested -> cite tests/test_analysis_gaps.py + tests/test_pdelta_validation.py, both
        confirmed present) and APPEND one LLM/NL/RAG future-work sentence built on the validated
        deterministic substrate (author-approved; keeps LLM out of contributions/validation).
  P124  run_replace: unify V = 239.8 -> 239.84 kN and Cs = 0.126 -> 0.1265 (appendix precision;
        239.84 = 0.1265 * 1896; R3 reproduced Cs=0.1265, V~239.84).
  T8R9  cell: "239.8 kN (Cs = 0.126, ...)" -> "239.84 kN (Cs = 0.1265, ...)".
  P240  run_replace: "thirty-six-combination" -> "36-combination" (numeral consistency).
  P167  DELETE empty conclusion-opening paragraph (runs=0, no sectPr/bookmark) -> paragraphs 241->240.

Table 8 = d.tables[8]. Backup -> ..._v2.pre_prose7a_2026-07-02.docx. Dry-run default; --apply.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_prose7a_2026-07-02.docx")

# P165 primitives citation (run_replace)
P165_CITE_OLD = "were exercised in the verification suite"
P165_CITE_NEW = ("were exercised in the verification suite (for example, "
                 "tests/test_analysis_gaps.py and tests/test_pdelta_validation.py)")
# P165 LLM future-work (append)
P165_LLM_APPEND = (
    " Separately, the natural-language and retrieval-augmented layers deliberately excluded "
    "from this study (Section 2.3) — a natural-language front-end for model and load "
    "specification, KDS retrieval-augmented explanation, and automated narration of the "
    "structural calculation report, all built on the validated deterministic substrate reported "
    "here — are reserved for a dedicated study in which their outputs can be verified "
    "independently of the benchmark evidence."
)

# P124 numerals (run_replace) — "kN" is in the next run (citation split), so match
# within run[15]: "V = 239.8 " (trailing space) and "Cs = 0.126,".
P124_V_OLD = "base shear V = 239.8 "
P124_V_NEW = "base shear V = 239.84 "
P124_CS_OLD = "response coefficient Cs = 0.126,"
P124_CS_NEW = "response coefficient Cs = 0.1265,"

# T8R9 cell — split across runs: "239.8 " | "kN" | " (Cs = 0.126, Ta = 0.42 " | "s" | ")"
T8R9_LOCATE = "239.8 kN (Cs = 0.126, Ta = 0.42 s)"
T8R9_SUBS = [("239.8 ", "239.84 "), ("Cs = 0.126,", "Cs = 0.1265,")]

# P240 numeral (run_replace)
P240_OLD = "thirty-six-combination"
P240_NEW = "36-combination"

LOC_P165 = "Several directions for future work follow from the limitations"
LOC_P124 = "an approximate fundamental period Ta = 0.42 s"
LOC_P240 = "Across all 50 individually checked quantities"
SENTINEL = "reserved for a dedicated study in which their outputs can be verified"


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:50]!r}")
    return hits[0]


def run_replace(par, old, new, tag):
    for r in par.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return
    raise SystemExit(f"ABORT [{tag}]: substring not in a single run: {old[:40]!r}")


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs
    if any(SENTINEL in p.text for p in paras):
        raise SystemExit("ABORT: already applied (P165 LLM sentinel present).")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    print(f"BEFORE: paragraphs={np0} tables={nt0} inline_shapes={ni0} fldChar={fld0} instrText={ins0}")

    # P165 cite + LLM append
    p165 = find_para(paras, LOC_P165, "P165")
    run_replace(p165, P165_CITE_OLD, P165_CITE_NEW, "P165cite")
    p165.add_run(P165_LLM_APPEND)

    # P124 numerals
    p124 = find_para(paras, LOC_P124, "P124")
    run_replace(p124, P124_V_OLD, P124_V_NEW, "P124V")
    run_replace(p124, P124_CS_OLD, P124_CS_NEW, "P124Cs")

    # T8R9 cell (value split across runs)
    t8 = d.tables[8]
    hits = [(ri, c) for ri, row in enumerate(t8.rows) for c in [row.cells[2]] if T8R9_LOCATE in c.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [T8R9]: matched {len(hits)} (want 1)")
    ri, cell = hits[0]
    if "fldChar" in cell._tc.xml or "instrText" in cell._tc.xml:
        raise SystemExit("ABORT [T8R9]: cell carries a field")
    for old, new in T8R9_SUBS:
        done = False
        for cpar in cell.paragraphs:
            for r in cpar.runs:
                if old in r.text:
                    r.text = r.text.replace(old, new)
                    done = True
                    break
            if done:
                break
        if not done:
            raise SystemExit(f"ABORT [T8R9]: '{old}' not in a single run")

    # P240 numeral
    p240 = find_para(paras, LOC_P240, "P240")
    run_replace(p240, P240_OLD, P240_NEW, "P240")

    # P167 delete (empty conclusion-opening paragraph)
    p166 = find_para(paras, "7. CONCLUSION", "P166")
    idx = paras.index(p166)
    p167 = paras[idx + 1]
    if p167.text.strip() != "" or p167.runs:
        raise SystemExit(f"ABORT [P167]: not empty: runs={len(p167.runs)} txt={p167.text[:40]!r}")
    if "sectPr" in p167._p.xml or "bookmark" in p167._p.xml:
        raise SystemExit("ABORT [P167]: carries sectPr/bookmark")
    p167._p.getparent().remove(p167._p)

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    print(f"AFTER : paragraphs={np1} tables={nt1} inline_shapes={ni1} fldChar={fld1} instrText={ins1}")
    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0} -> {fld1}/{ins1}"
    assert (np1, nt1, ni1) == (np0 - 1, nt0, ni0), f"structure count unexpected: {(np1, nt1, ni1)} vs {(np0-1, nt0, ni0)}"

    print(f"\n--- P165 tail ---\n{find_para(d.paragraphs, LOC_P165, 'x').text[-520:]}")
    print(f"\n--- P124 tail ---\n{find_para(d.paragraphs, LOC_P124, 'x').text[-260:]}")
    print(f"\n--- T8R9 cell ---\n{t8.rows[ri].cells[2].text}")
    print(f"\n--- P240 tail ---\n{find_para(d.paragraphs, LOC_P240, 'x').text[-200:]}")
    print("\n--- P166/next after P167 delete ---")
    plist = d.paragraphs
    ci = next(i for i, p in enumerate(plist) if "7. CONCLUSION" in p.text)
    print(f"  P{ci}: {plist[ci].text!r}\n  P{ci+1}: {plist[ci+1].text[:60]!r}")
    p240now = next(p for p in plist if "Across all 50 individually checked" in p.text)
    print(f"  P240 has '36-combination': {'36-combination' in p240now.text}; "
          f"'thirty-six' gone: {'thirty-six' not in p240now.text}")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
