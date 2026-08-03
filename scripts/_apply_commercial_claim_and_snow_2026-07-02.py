"""Session 2026-07-02b: upgrade the commercial-validation framing + snow non-governing note.

Three plain-run edits (no field/SEQ/citation touched):

ITEM 1 (Abstract, P009): the IFC/L-shape application-example sentence still under-claims
  ("returns results consistent with engineering expectations"). Append the ETABS
  cross-check outcome (§5.2 regular example: 31/31 global-response metrics within 1%;
  each example independently cross-checked against ETABS).

ITEM 2 (Conclusion, P171): same under-claim in the conclusion. Append the ETABS
  cross-check (31 metrics within 1%), stating it extends the commercial cross-validation
  to the IFC-to-analysis pipeline itself (closes reviewer gap G3).

ITEM 3 (§5.2, P132): add one sentence noting the generated roof snow load (0.5 kN/m2)
  enters the governing combinations only as a 0.2 companion action and does not control
  the drift / member-strength results, which are governed by the Y-direction earthquake.

P027 (§2.2), P031 (§2.4), Table 1 (This-work row: "Midas Gen + ETABS"), P134 (§5.2 ETABS
  31/31 paragraph), P145/P151 (§5.3 L-shape ETABS) and P154 (§6.1 "cross-checked against
  ETABS to within 1%") already carry the two-program framing and are left untouched.

Numbers verified against validation/ifc_example_etabs_compare.json (31/31 OK; DL corner
  0.94%; drifts/base shears <=0.01%; envelope <=0.04%) and os_envelope gov combos
  (1.2924DL+1.0LL+1.0EQ*+0.2S -> Y-earthquake dominant, snow a 0.2 companion).

Backup -> ..._v2.pre_commclaim_2026-07-02.docx. Dry-run default; --apply saves.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_commclaim_2026-07-02.docx")

# ---- ITEM 1: Abstract P009 ----
LOC_P009 = "node-element L-shaped irregular-plan demonstration, exercises the full pipeline"
OLD_P009 = ("exercises the full pipeline end-to-end and returns results consistent with "
            "engineering expectations under KDS-generated loading.")
NEW_P009 = ("exercises the full pipeline end-to-end and returns results consistent with "
            "engineering expectations under KDS-generated loading; the global response of "
            "each is independently cross-checked against ETABS, the regular example agreeing "
            "on all 31 compared metrics within 1%.")

# ---- ITEM 2: Conclusion P171 ----
LOC_P171 = ("Across five controlled benchmark cases spanning two-dimensional frames, a "
            "three-dimensional moment frame, and a supplementary geometrically nonlinear case")
OLD_P171 = ("returned drift, displacement, and member-strength results consistent with "
            "engineering expectations under KDS-generated loading.")
NEW_P171 = ("returned drift, displacement, and member-strength results consistent with "
            "engineering expectations under KDS-generated loading; its global response was "
            "independently cross-checked against ETABS, with all 31 compared metrics agreeing "
            "within 1%, extending the commercial cross-validation to the IFC-to-analysis "
            "pipeline itself.")

# ---- ITEM 3: §5.2 P132 snow non-governing note ----
LOC_P132 = "The preliminary member strength check"
SNOW_SENTENCE = (" The generated roof snow load (0.5 kN/m2) enters the governing seismic "
                 "combinations only as a 0.2 companion action and does not control any drift "
                 "or member-strength result; the reported demands are governed by the "
                 "Y-direction earthquake as the dominant action.")

SENTINEL = "extending the commercial cross-validation to the IFC-to-analysis pipeline itself"


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:50]!r}")
    return hits[0]


def repl(par, old, new, tag):
    for r in par.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return
    raise SystemExit(f"ABORT [{tag}]: substring not found in a single run: {old[:50]!r}")


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    if any(SENTINEL in p.text for p in paras):
        raise SystemExit("ABORT: already applied (conclusion ETABS sentinel present).")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    print(f"BEFORE: paragraphs={np0} tables={nt0} inline_shapes={ni0} fldChar={fld0} instrText={ins0}")

    p009 = find_para(paras, LOC_P009, "P009")
    repl(p009, OLD_P009, NEW_P009, "P009")

    p171 = find_para(paras, LOC_P171, "P171")
    repl(p171, OLD_P171, NEW_P171, "P171")

    p132 = find_para(paras, LOC_P132, "P132")
    p132.add_run(SNOW_SENTENCE)

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    print(f"AFTER : paragraphs={np1} tables={nt1} inline_shapes={ni1} fldChar={fld1} instrText={ins1}")

    # guards
    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0} -> {fld1}/{ins1}"
    assert (np1, nt1, ni1) == (np0, nt0, ni0), "structure count changed"

    print("\n--- P009 now ---\n", p009.text)
    print("\n--- P171 now ---\n", p171.text)
    print("\n--- P132 tail ---\n", p132.text[-260:])

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
