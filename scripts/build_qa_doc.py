# -*- coding: utf-8 -*-
"""LLM-AE-AI 중간발표 예상 Q&A + 성능 평가 메트릭 문서 생성"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(r"C:\Users\youm\Desktop\백장운_수업")
OUT_PATH = OUT_DIR / "LLM-AE-AI_중간발표_QA_손명국.docx"
FONT = "맑은 고딕"

NAVY = RGBColor(0x11, 0x2D, 0x4E)
BLUE = RGBColor(0x3F, 0x72, 0xAF)
GRAY = RGBColor(0x55, 0x55, 0x55)


def set_font(run, size=10, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for k in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(k), FONT)


def add_para(doc, text="", size=10, bold=False, color=None, space_after=4, indent=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.3
    if indent:
        pf.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=15, bold=True, color=NAVY)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=BLUE)


def add_q(doc, qno, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Q{qno}. ")
    set_font(run, size=11, bold=True, color=NAVY)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=NAVY)


def add_a(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.4)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.35
    run = p.add_run("A. ")
    set_font(run, size=10, bold=True, color=BLUE)
    run = p.add_run(text)
    set_font(run, size=10)


def add_bullet(doc, text, indent=0.6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run("• " + text)
    set_font(run, size=10)


def add_table(doc, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=9, bold=True)
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "DCE6F1")
        tc_pr.append(shd)
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.2
            run = p.add_run(str(val))
            set_font(run, size=9)
    if widths_cm:
        for row in table.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
    return table


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # ── 표제 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LLM-AE-AI 중간발표 — 예상 Q&A 및 성능 평가 메트릭")
    set_font(run, size=17, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("3D Editor + KDS-RAG 설계 에이전트")
    set_font(run, size=11, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("손명국 · 2026310208 · 경희대학교 건축공학과")
    set_font(run, size=10, color=GRAY)

    # ──────────────── 1. 표면적 질문 ────────────────
    add_h1(doc, "1. 표면적 질문 (Surface) — 동기 · 범위 · 차별성")
    add_para(doc, "프로젝트의 동기·차별성·완성 가능성에 대한 일반 질문. 가장 받기 쉬운 유형.", size=9, color=GRAY, space_after=4)

    qa_surface = [
        (
            "왜 이 주제를 선택했나요?",
            "기존 OpenSees 기반 3D Editor를 직접 구현하면서, NG가 발생했을 때 \"왜 NG인지\"와 \"어떻게 고칠지\"를 사용자가 매번 KDS PDF를 뒤져가며 시행착오로 해결해야 한다는 점이 가장 큰 페인포인트였습니다. 이 구간을 LLM이 자동화할 수 있다고 보았고, 마침 LLM-AE-AI 수업의 RAG · Tool Use · MCP가 정확히 이 문제에 맞아 본 과제로 정했습니다.",
        ),
        (
            "Midas Gen 같은 상용 SW가 있는데 왜 직접 만드나요?",
            "상용 SW는 해석 엔진은 강력하지만 KDS 조항을 결과에 자동 인용하거나, NG 부재에 대해 LLM 기반 변경안을 제시하는 기능은 없습니다. 본 과제는 해석 엔진 자체를 대체하는 것이 아니라, 해석 결과 위에 KDS 근거 기반 AI 레이어를 얹는 것입니다. Midas와는 보완 관계로 보고 있습니다.",
        ),
        (
            "왜 KDS인가요? Eurocode나 ASCE 7이 아닌?",
            "국내 실무에서 가장 직접적인 가치를 가지고, 한국어 LLM 환경에서 RAG 효과가 가장 클 것으로 판단했습니다. 또한 본 시스템의 RAG 코퍼스만 교체하면 Eurocode · ASCE 7로 확장 가능한 구조로 설계하기 때문에, KDS는 검증된 첫 번째 사례로 두는 것입니다.",
        ),
        (
            "기말까지 정말 완성 가능한가요?",
            "V2 3D Editor 본체는 이미 Midas 89% 일치 수준으로 검증 완료된 상태이고, 본 과제는 그 위에 RAG · 에이전트 레이어를 얹는 작업으로 범위를 한정했습니다. F1(RAG 챗봇)과 F2(자동 인용)를 우선 확보하고, F3 · F4는 단계적으로 적용하는 전략입니다. F1 · F2까지만으로도 동작하는 데모는 확보 가능합니다.",
        ),
        (
            "발표에서 가장 강조하고 싶은 한 가지는?",
            "검증된 substrate 위에 AI Layer를 접합한다는 점입니다. LLM의 출력 자체를 결과로 신뢰하는 것이 아니라, 모든 제안이 OpenSeesPy 재해석을 거쳐 ground truth가 확보되는 구조라는 것이 핵심 차별점입니다.",
        ),
        (
            "개인 발표인데 작업량이 너무 크지 않나요?",
            "V2 본체가 완성 상태이고 KDS DB 712건도 이미 적재된 상태입니다. 본 과제에서 추가하는 것은 RAG 파이프라인과 5단계 프롬프트, V2 UI에 패널 통합 정도로, 6주 분량으로 산정 가능한 범위입니다.",
        ),
    ]
    for i, (q, a) in enumerate(qa_surface, start=1):
        add_q(doc, i, q)
        add_a(doc, a)

    # ──────────────── 2. 구체적 질문 ────────────────
    add_h1(doc, "2. 구체적 질문 (Specific) — 기술 · 구현 · 검증")
    add_para(doc, "RAG · 에이전트 · 평가 등 기술 내부 구조에 대한 질문. 교수 또는 LLM 학습 중인 동료가 던질 가능성이 큰 유형.", size=9, color=GRAY, space_after=4)

    qa_specific = [
        (
            "RAG 청킹 전략은 어떻게 설계하나요?",
            "KDS는 조항 단위 계층 구조(절 · 항 · 호)가 명확하므로 structure-aware chunking을 채택할 예정입니다. 본문은 조항 단위로, 표는 별도 청크로 분리하고 메타데이터에 조항 번호를 명시해 Citations에 직접 활용합니다. 청크 크기는 우선 500~800 토큰 기준으로 잡고, 임베딩 품질을 보면서 조정하겠습니다.",
        ),
        (
            "BM25와 임베딩의 가중치는 어떻게 정하나요?",
            "초기에는 RRF (Reciprocal Rank Fusion)로 두 결과를 결합해 가중치 튜닝 부담을 줄이고, k=60 표준값으로 시작합니다. 평가 셋(KDS 자체 예제 + 직접 작성한 질의 30~50개) 위에서 Top-K 정확도를 측정해 가중치 또는 RRF k값을 조정할 계획입니다.",
        ),
        (
            "Verifier 자체가 LLM이면 환각이 누적되지 않나요?",
            "정확한 지적입니다. 그래서 두 단계 안전장치를 둡니다. 첫째, 모든 수치 결과는 OpenSeesPy 재해석 ground truth로 확정되어 LLM 출력이 결과로 사용되지 않습니다. 둘째, Verifier는 수치 판정이 아니라 제안의 KDS 적합성만 평가하며, 판정이 모호하면 사용자 승인 게이트로 escalation됩니다. 즉 LLM은 진단 · 제안 · 근거 인용까지만, 최종 검증은 solver + 사람입니다.",
        ),
        (
            "F3가 H-350을 제안했는데 단면 DB에 없는 사이즈라면?",
            "후보는 KS D 3502/3568 기반 738건 단면 DB 안에서만 검색하도록 제한합니다. LLM이 임의로 단면을 생성하는 것이 아니라, RAG retrieval 결과로 유효한 단면 후보군만 LLM에게 제공하는 grounded generation 방식입니다. DB에 없는 단면은 제안 자체가 불가능합니다.",
        ),
        (
            "max_iter=5로 116개 NG를 모두 OK로 만들 수 있나요?",
            "모든 NG를 한 번에 수렴시키는 것이 아니라, 부재별 또는 그룹별로 루프를 돌립니다. 동일 단면을 쓰는 기둥 그룹 단위로 처리하면 5회 반복 안에 활용도 비를 1.0 이하로 낮추는 것이 일반적으로 충분합니다. 수렴 못 하는 경우는 사용자에게 escalation하고 그룹을 재분할하도록 설계할 예정입니다.",
        ),
        (
            "Prompt Caching으로 KDS 본문을 캐싱하면 1MB 이상일 텐데 비용은?",
            "Anthropic Prompt Caching은 5분 TTL 기준 cache hit 시 비용이 약 1/10로 떨어지고, 1MB 수준의 KDS 컨텍스트는 한 번 캐싱하면 동일 세션의 다수 질의에서 재활용되어 비용 효율이 매우 높습니다. 실측 후 캐시 단위(전체 KDS vs 챕터별)를 조정할 계획입니다.",
        ),
        (
            "LLM이 환각으로 KDS에 없는 조항을 인용하면?",
            "Citations는 RAG로 retrieve된 청크 ID를 직접 인용 출처로 사용하므로, retrieval 결과 밖의 조항은 인용 자체가 불가능한 구조입니다. 추가로 응답 후처리에서 조항 번호 정규표현식 매칭으로 retrieval 결과와 일치 여부를 검증하는 가드를 둘 수 있습니다.",
        ),
        (
            "IFC가 Revit · Tekla 등 BIM 도구마다 export 결과가 달라지는데?",
            "V2 IFC 파서가 벽 기반과 기둥 기반 두 가지 building type을 모두 지원하고, 4차시도 10층 골조 모델로 ALL PASS 검증한 상태입니다. 단면 매핑은 현재 SHS까지 검증 완료, CHS · L · 채널은 Revit IFC 추가 입수 후 확장 예정으로 현재 보완 작업 중입니다.",
        ),
        (
            "5단계 프롬프트가 W7 게스트 세미나 패턴 그대로라면 본인 기여는?",
            "패턴 구조는 동일하지만 본 과제에서의 기여는 두 가지입니다. 첫째, 각 단계에서 호출되는 도메인 도구(MCP Tools)를 구조해석 영역으로 매핑한 것 — change_section, re_analyze, KDS retrieval 등. 둘째, Verifier 단계의 평가 프롬프트를 KDS 코드 적합성 · 과설계 · 환각 3축으로 정의한 것입니다. 즉 패턴은 차용, 도메인 적응이 기여입니다.",
        ),
        (
            "KDS 자체가 개정되면 RAG DB 업데이트는 어떻게?",
            "KDS 개정은 보통 수년 주기이므로 버전 메타데이터(code_version)를 모든 청크에 포함시키고, 개정판 입수 시 해당 조항만 재청킹 · 재임베딩하는 방식으로 운영합니다. Citations에도 버전이 함께 표시되어 사용자가 적용 버전을 명확히 확인 가능합니다.",
        ),
    ]
    for i, (q, a) in enumerate(qa_specific, start=7):
        add_q(doc, i, q)
        add_a(doc, a)

    # ──────────────── 3. Fallback ────────────────
    add_h1(doc, "3. Fallback — 답하기 곤란할 때")
    add_q(doc, "F1", "실제 시연 가능한가요?")
    add_a(doc, "MCP Inspector 시연 영상은 기말 발표 시 제출 예정이며, 오늘은 계획(Proposal) 발표라서 실제 동작 데모는 포함하지 않았습니다. 다만 V2 3D Editor 본체는 https://opensees-v2.onrender.com/ 에서 현재도 접근 가능합니다.")

    add_q(doc, "F2", "성능 평가 메트릭은? (간단 답변)")
    add_a(doc, "F1 retrieval은 Top-K accuracy와 MRR, F3 에이전트는 수렴 성공률과 평균 iter 수, F4 Verifier는 사람 평가 대비 precision/recall로 측정 예정입니다. 자세한 평가 체계는 다음 섹션 4에 정리되어 있습니다.")

    # ──────────────── 4. 성능 평가 메트릭 (확장) ────────────────
    add_h1(doc, "4. 성능 평가 메트릭 (확장 — 발표 보완)")
    add_para(
        doc,
        "발표 슬라이드에는 분량 한계로 포함되지 않았으나, Q&A 또는 기말 보고서에서 반드시 필요한 정량 평가 체계를 본 섹션에 정리한다. 각 핵심 기능(F1~F4)마다 측정 지표 · 평가 셋 · baseline을 명시하고, 마지막에 발표용 요약 슬라이드 안을 제안한다.",
    )

    # 4.1
    add_h2(doc, "4.1 평가 체계 개요")
    add_para(doc, "본 시스템의 평가는 두 축으로 구성된다.")
    add_bullet(doc, "기능별 평가 (Component-level) — F1~F4 각각 독립적 지표")
    add_bullet(doc, "통합 평가 (End-to-End) — IFC 업로드부터 OK 수렴까지의 wall-clock 시간 및 비용")
    add_para(doc, "평가 셋은 다음 세 종류를 구축한다.")
    add_bullet(doc, "KDS Gold Set : 50~100개 자연어 질의 ↔ 정답 조항 매핑 (F1, F2 평가용)")
    add_bullet(doc, "NG Case Set : 10~20개 NG 시나리오 (압축 NG, 휨 NG, 상관식 NG, 층간변위 NG 등) — F3 평가용")
    add_bullet(doc, "Human Baseline : 본인 + 1~2명 외부 평가자가 동일 NG를 수동으로 해결한 시간 및 단면 선택 (F3, F4 비교용)")

    # 4.2 F1
    add_h2(doc, "4.2 F1 — KDS RAG 챗봇 평가")
    add_table(
        doc,
        headers=["지표", "정의", "목표값", "측정 방법"],
        rows=[
            ["Top-1 Accuracy", "1순위 retrieval이 정답 조항인 비율", "≥ 0.75", "Gold Set 50건"],
            ["Top-5 Accuracy (Recall@5)", "정답 조항이 Top-5 안에 포함된 비율", "≥ 0.95", "Gold Set 50건"],
            ["MRR (Mean Reciprocal Rank)", "정답 조항 순위의 역수 평균", "≥ 0.80", "Gold Set 50건"],
            ["nDCG@10", "순위 가중 누적 적합도", "≥ 0.85", "Gold Set 50건"],
            ["응답 지연 시간 (latency)", "질의 입력 → 응답 완료 시간", "≤ 3 s (cached)", "100회 반복 평균"],
        ],
        widths_cm=[3.6, 5.0, 2.4, 4.0],
    )
    add_para(doc, "Baseline 비교 : BM25 단독 · 임베딩 단독 · BM25+임베딩 RRF — 세 조합을 동일 셋으로 측정해 RRF 우위를 정량 확인한다.", size=9)

    # 4.3 F2
    add_h2(doc, "4.3 F2 — Design Check 자동 인용 평가")
    add_table(
        doc,
        headers=["지표", "정의", "목표값", "측정 방법"],
        rows=[
            ["Citation Precision", "인용된 조항이 해당 검토에 실제 적용되는 비율", "≥ 0.90", "NG Case Set + 전문가 검수"],
            ["Citation Recall", "적용되어야 할 조항이 모두 인용된 비율", "≥ 0.85", "NG Case Set + 전문가 검수"],
            ["Hallucination Rate", "KDS에 존재하지 않는 조항을 인용한 비율", "= 0", "응답 사후 KDS DB 매칭 검증"],
            ["인용 완전성", "조항 번호 · 본문 · 버전 메타데이터 누락률", "≤ 5 %", "응답 포맷 검사"],
        ],
        widths_cm=[3.6, 5.5, 2.0, 3.9],
    )
    add_para(doc, "Hallucination Rate 0%는 retrieval 결과 청크 ID 외의 조항 인용을 시스템적으로 차단하는 가드 기반으로 달성한다.", size=9)

    # 4.4 F3
    add_h2(doc, "4.4 F3 — 설계 제안 에이전트 평가")
    add_table(
        doc,
        headers=["지표", "정의", "목표값", "측정 방법"],
        rows=[
            ["Convergence Rate", "max_iter(=5) 내 모든 부재 ratio ≤ 1.0 도달 비율", "≥ 0.80", "NG Case Set 10~20건"],
            ["Avg. Iterations to Converge", "수렴 시까지 평균 루프 반복 수", "≤ 2.5회", "NG Case Set"],
            ["First-Suggestion Acceptance", "1차 제안이 채택되는 비율", "≥ 0.60", "NG Case Set"],
            ["Over-design Rate", "최종 ratio < 0.5인 부재 비율 (필요 이상 큰 단면)", "≤ 0.15", "NG Case Set 평균"],
            ["Time-to-OK", "NG 검출 → 모든 부재 OK 도달까지 wall-clock", "≤ 60 s (10층 모델)", "5회 반복 평균"],
            ["Manual vs Agent 시간 감소율", "수동 해결 시간 대비 절감 비율", "≥ 80 %", "Human Baseline 비교"],
        ],
        widths_cm=[3.8, 5.4, 2.4, 3.4],
    )
    add_para(doc, "Time-to-OK는 본 과제의 핵심 KPI. 5분 발표에서 \"수 시간 → 수십 초\"라 주장한 부분의 정량 근거.", size=9)

    # 4.5 F4
    add_h2(doc, "4.5 F4 — LLM-as-Judge Verifier 평가")
    add_table(
        doc,
        headers=["지표", "정의", "목표값", "측정 방법"],
        rows=[
            ["Precision (vs Human)", "Verifier OK 판정 중 사람도 OK한 비율", "≥ 0.85", "50건 사람-LLM 페어 평가"],
            ["Recall (vs Human)", "사람 OK 판정 중 Verifier도 OK한 비율", "≥ 0.80", "50건 페어 평가"],
            ["False Positive Rate", "잘못된 제안을 통과시킨 비율 (안전 직결)", "≤ 0.10", "50건 페어 평가"],
            ["Cohen's Kappa", "사람-Verifier inter-rater agreement", "≥ 0.70", "50건 페어 평가"],
            ["Self-Consistency", "동일 입력 5회 반복 시 판정 일치율", "≥ 0.90", "10건 × 5회 반복"],
        ],
        widths_cm=[3.6, 5.6, 2.2, 3.6],
    )
    add_para(doc, "False Positive Rate는 안전성 직결 지표 — 0.10 초과 시 사용자 승인 게이트 강화로 보완한다.", size=9)

    # 4.6 E2E
    add_h2(doc, "4.6 End-to-End 통합 평가")
    add_table(
        doc,
        headers=["지표", "정의", "목표값"],
        rows=[
            ["IFC-to-OK Total Time", "IFC 업로드부터 모든 부재 OK까지 wall-clock", "≤ 5 분 (10층 모델)"],
            ["User Click Reduction", "기존 수동 워크플로우 대비 클릭/조작 수 감소", "≥ 70 %"],
            ["Cost per Run (API)", "1회 E2E 실행 시 Anthropic API 비용", "≤ $0.5 (Prompt Caching 적용)"],
            ["Cache Hit Rate", "Prompt Caching 활용률", "≥ 0.80"],
        ],
        widths_cm=[4.0, 6.4, 4.8],
    )

    # 4.7 평가 셋 구축 계획
    add_h2(doc, "4.7 평가 셋 구축 및 일정")
    add_para(doc, "평가 셋 구축은 RAG 파이프라인 구축과 병행한다 (계획서 §7 일정의 W2~W3 구간).")
    add_table(
        doc,
        headers=["평가 셋", "규모", "구축 방법", "주차"],
        rows=[
            ["KDS Gold Set", "50~100건", "KDS 본문에서 자연어 질의 작성 + 정답 조항 라벨링", "W2~"],
            ["NG Case Set", "10~20건", "기존 4차시도 IFC + 인위적 NG 케이스 (단면 축소 등)", "W3~"],
            ["Human Baseline", "5~10건", "본인 + 외부 평가자가 동일 NG 수동 해결 (시간/단면 기록)", "W4~"],
            ["Cost & Latency", "100회", "각 기능별 100회 반복 측정", "W5~"],
        ],
        widths_cm=[3.4, 2.4, 6.8, 2.6],
    )

    # 4.8 발표 보완 제안
    add_h2(doc, "4.8 발표 슬라이드 보완 제안")
    add_para(doc, "현재 11장 슬라이드에는 평가 지표가 명시되어 있지 않다. 분량을 늘리지 않으면서 보완하는 두 가지 안을 제시한다.")
    add_para(doc, "Option A — Backup 슬라이드 (질문 대응용, 본 발표에는 미사용)", bold=True, color=BLUE, space_after=2)
    add_bullet(doc, "표 형태로 F1~F4 + E2E의 핵심 KPI만 압축 (Top-1 Acc, Conv Rate, Time-to-OK, FPR 등 5~6개)")
    add_bullet(doc, "Q&A에서 \"성능 평가 메트릭은?\" 질문 시 띄워 보여주는 용도")
    add_bullet(doc, "발표 시간에 포함되지 않으므로 5분 제약 영향 없음")

    add_para(doc, "Option B — 슬라이드 10번 (Demo) 우측 하단에 작은 박스 추가", bold=True, color=BLUE, space_after=2)
    add_bullet(doc, "\"평가 지표\" 박스 한 줄 : Conv Rate ≥ 80% · Time-to-OK ≤ 60s · Hallucination 0% · FPR ≤ 10%")
    add_bullet(doc, "본문에서 별도 설명하지 않고, 데모 시나리오 옆에 정량 목표만 노출")
    add_bullet(doc, "발표 시간 추가 없이 시각적으로 평가 의식이 있다는 신호 제공")

    add_para(doc, "권장 : Option A + Option B 병행. 슬라이드 10에 한 줄, Backup 한 장 별도 준비.", color=NAVY, bold=True)

    # ──────────────── 5. 준비 우선순위 ────────────────
    add_h1(doc, "5. 준비 우선순위 (TOP 3)")
    add_para(doc, "발표 후 가장 받기 쉬운 질문 — 답변을 거의 외워서 대답할 수 있게 준비할 것.")
    add_bullet(doc, "Q4 (기말까지 완성 가능?) — 일정 · 범위 질문은 거의 확정적으로 출제")
    add_bullet(doc, "Q9 (Verifier 환각 누적?) — 교수님이 가장 좋아하는 critical thinking 질문")
    add_bullet(doc, "Q15 (5단계 본인 기여?) — W7 게스트 세미나 패턴 차용 부분, 반드시 검증 들어옴")
    add_bullet(doc, "F2 (성능 평가 메트릭?) — 본 문서 §4 표를 외워두면 좋음 — Conv Rate, Time-to-OK, FPR 정도")

    doc.save(OUT_PATH)
    print(f"[OK] saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
