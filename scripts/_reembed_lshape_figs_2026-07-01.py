"""Task C: re-embed the regenerated L-shape figures (in-image "Figure 6/7" prefixes
removed) into the canonical v2 DOCX.

Manuscript Figure 7 (zone decomposition) = inline shape index 6 = fig6_lshape_zones.png
Manuscript Figure 8 (L-shape 3D model)   = inline shape index 7 = fig7_lshape_3d.png
(The generator/final file names are offset from the manuscript figure numbers; target by
shape index + an old-bytes guard.) Keeps display width, recomputes height. Refreshes
final/fig6.png, final/fig7.png. Backup -> ..._v2.pre_lshapefigs_2026-07-01.docx.
"""
from __future__ import annotations

import shutil
import struct
import sys
from pathlib import Path

import docx
from docx.shared import Emu

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs" / "paper1_open_source_alternative"
DOC = BASE / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_lshapefigs_2026-07-01.docx")

# (shape index, source png, final copy, expected old bytes)
JOBS = [
    (6, BASE / "figures" / "fig6_lshape_zones.png", BASE / "figures" / "final" / "fig6.png", 82010),
    (7, BASE / "figures" / "fig7_lshape_3d.png", BASE / "figures" / "final" / "fig7.png", 230763),
]


def png_size(b):
    assert b[12:16] == b"IHDR"
    return struct.unpack(">II", b[16:24])


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    if len(d.inline_shapes) != 8:
        raise SystemExit(f"ABORT: expected 8 shapes, got {len(d.inline_shapes)}")
    for idx, src, _final, oldlen in JOBS:
        sh = d.inline_shapes[idx]
        part = d.part.related_parts[sh._inline.graphic.graphicData.pic.blipFill.blip.embed]
        if len(part._blob) != oldlen:
            raise SystemExit(f"ABORT: shape[{idx}] bytes {len(part._blob)} != {oldlen}")
        nb = src.read_bytes()
        nw, nh = png_size(nb)
        ow = int(sh.width)
        print(f"shape[{idx}] <- {src.name} {nw}x{nh} {len(nb)}B; width kept {ow/914400:.2f}in")
    if not apply:
        print("\nDRY-RUN. Re-run with --apply.")
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"backup -> {BACKUP.name}")
    for idx, src, final, _oldlen in JOBS:
        sh = d.inline_shapes[idx]
        part = d.part.related_parts[sh._inline.graphic.graphicData.pic.blipFill.blip.embed]
        nb = src.read_bytes()
        nw, nh = png_size(nb)
        ow = int(sh.width)
        part._blob = nb
        sh.height = Emu(int(ow * nh / nw))
        shutil.copy2(src, final)
        print(f"  embedded shape[{idx}], refreshed {final.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")

    d2 = docx.Document(str(DOC))
    media = [p.partname for p in d2.part.package.iter_parts() if "media" in str(p.partname)]
    print(f"[verify] shapes={len(d2.inline_shapes)} media={len(media)} "
          f"paras={len(d2.paragraphs)} tables={len(d2.tables)}")
    for idx, _s, _f, _o in JOBS:
        b = d2.part.related_parts[d2.inline_shapes[idx]._inline.graphic.graphicData.pic.blipFill.blip.embed]._blob
        print(f"  shape[{idx}] now {len(b)} bytes")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
