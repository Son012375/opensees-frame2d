"""§2.3.3 Zone-Based Decomposition + §4.3 L-Shape Application Example 삽입.

`_patch_paper_draft.py` → `_add_case1_table5a.py` 실행 후, `_add_appendix_a2.py`
전에 실행한다.

이 스크립트는 v2 docx에 두 개의 신규 section을 삽입한다:
- §2.3.3 Zone-Based Decomposition for Orthogonal Irregular Plans
  - 위치: §2.3.2 본문 (para 42) 직후, §2.4 heading (para 44) 앞
  - 내용: subheading + 3 body paragraphs
- §4.3 Irregular-Plan Application Example (L-Shape)
  - 위치: §4 마지막 (Figure 5 caption, para 120) 직후, §5 heading (para 121) 앞
  - 내용: subheading + 4 body paragraphs + Table 8 + Figure 6 + Figure 7

운영 순서:
    1. python scripts/_patch_paper_draft.py        # 원본 → v2 (text rewrites)
    2. python scripts/_add_case1_table5a.py        # v2 → v2 (Table 4a 삽입)
    3. python scripts/_add_irregular_sections.py   # v2 → v2 (§2.3.3 + §4.3 삽입)
    4. python scripts/_add_appendix_a2.py          # v2 → v2 (Appendix A 추가)

본 스크립트는 idempotent하지 않다 — 재실행 전 처음부터 다시 실행해야 한다.
"""
from __future__ import annotations

import json
from pathlib import Path

import docx
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
FIG_DIR = ROOT / "docs" / "paper1_open_source_alternative" / "figures"
METRICS_JSON = ROOT / "docs" / "paper1_open_source_alternative" / "validation" / "lshape_demonstration.json"


# ──────────────────────────────────────────────────────────────
# Anchors (verified against v2 docx state after _add_case1_table5a)
# ──────────────────────────────────────────────────────────────
# After _add_case1_table5a, all paragraphs >= 84 shift by +1.
# §2.3.2 body ends at source para 42 → v2 para 42 (before Table 4a anchor 83/84)
# §4 ends around source para 120 (Figure 5) → v2 para 121 (+1 shift)
# We'll use text-search anchoring for robustness.

ANCHOR_TEXT_2_3_2_END = "diagnostic layer is central to the open-source"
ANCHOR_TEXT_FIG5 = "Figure 5. Summary of preliminary review outputs"


# ──────────────────────────────────────────────────────────────
# §2.3.3 content
# ──────────────────────────────────────────────────────────────

SEC_2_3_3_HEADING = "2.3.3. Zone-Based Decomposition for Orthogonal Irregular Plans"

SEC_2_3_3_BODY = [
    (
        "The node-element representation produced by Sections 2.3.1–2.3.2 "
        "supports both regular orthogonal grids and orthogonal irregular plan "
        "configurations. For irregular plans, the workflow decomposes the "
        "floor footprint into one or more orthogonal zones (such as the two "
        "rectangular wings of an L-shaped plan or the upper and lower "
        "stories of a setback), each described by its bay spacings, plan "
        "origin, and active-story range. Nodes on shared zone boundaries are "
        "merged by a tolerance-based coordinate search so that the resulting "
        "graph is a single connected analysis model rather than a disjoint "
        "set of subgrids."
    ),
    (
        "When the input is an IFC file, zones are detected automatically "
        "from the column-occupancy pattern. For each story, column "
        "endpoints are projected onto the inferred grid intersections; "
        "intersections without an occupying column define the unoccupied "
        "region, and connected unoccupied regions are grouped into "
        "rectangular zones via a row-profile clustering procedure. The "
        "result is refined story-by-story so that setbacks — stories with "
        "reduced occupancy relative to lower stories — are represented by "
        "zones with shorter active-story ranges. Topmost stories with no "
        "columns are interpreted as roof levels and removed."
    ),
    (
        "Zone-based decomposition supports L-shape, T-shape, and setback "
        "configurations on orthogonal grids. Non-orthogonal grids, rotated "
        "zones, curved members, and internal column removal within a "
        "rectangular zone are outside the present implementation. KDS load "
        "generation accounts for zone-wise floor area through an "
        "inclusion–exclusion calculation, so that gravity loads on setbacks "
        "reflect the reduced active footprint at each story."
    ),
]


# ──────────────────────────────────────────────────────────────
# §4.3 content
# ──────────────────────────────────────────────────────────────

SEC_4_3_HEADING = "4.3 Irregular-Plan Application Example (L-Shape)"

SEC_4_3_BODY_PRE_TABLE = [
    (
        "To demonstrate the workflow's capability on an orthogonal irregular "
        "plan configuration, an L-shaped steel frame with a five-story left "
        "wing and a three-story setback right wing is assembled directly as "
        "a node-element graph consistent with the zone-based representation "
        "described in Section 2.3.3. The example exercises the analysis "
        "pipeline on the resulting irregular model; it is not presented as "
        "an IFC auto-detection benchmark."
    ),
    (
        "The plan comprises a 12 m × 8 m left wing (five stories) and a "
        "12 m × 4 m right wing (three stories), with a uniform 3.5 m story "
        "height. The two wings share a common boundary at x = 12 m on which "
        "the zone-boundary nodes are merged, producing a single connected "
        "node-element graph that the OpenSeesPy frame module treats "
        "identically to a regular orthogonal grid. The assembled model "
        "contains 70 nodes and 135 elements (57 columns and 78 beams) "
        "supported on 13 base nodes corresponding to the combined column "
        "footprints of the two wings. The plan footprint per story is 12 × "
        "8 m for the full five-story left wing and 12 × 4 m for the "
        "three-story right wing, yielding nominal floor areas of 96 m² "
        "(stories 1–5 left) and 48 m² (stories 1–3 right). KDS-style dead "
        "and live loads (DL = 5.1 kN/m², LL = 2.5 kN/m² for office occupancy "
        "on all floors) are applied as floor-area loads; the open-source "
        "load generator distributes each story-level pressure to the beams "
        "of that story using a tributary-width / panel-based scheme."
    ),
    (
        "Under the gravity case, the total vertical base reaction is "
        "3,458 kN for dead load and 1,695 kN for live load — a ratio of "
        "2.04, matching the applied DL/LL intensity ratio (5.1 / 2.5 = "
        "2.04), which indicates that the two gravity load cases are "
        "distributed consistently by the same procedure. Horizontal base "
        "reactions remain below 0.01 kN, confirming static equilibrium of "
        "the assembled model. These results indicate that the same "
        "workflow pipeline — node-element model assembly, KDS load "
        "distribution, and OpenSeesPy analysis — runs end-to-end on the "
        "L-shaped configuration without geometry-specific manual "
        "intervention."
    ),
]

SEC_4_3_BODY_POST_TABLE = [
    (
        "Distinct in scope from the benchmark validation reported in "
        "Section 3, this example illustrates the workflow's irregular-plan "
        "capability through end-to-end execution and gravity-equilibrium "
        "verification rather than through numerical comparison against a "
        "commercial reference. A direct Midas Gen benchmark for L-shape, "
        "T-shape, and setback configurations — covering zone-boundary node "
        "merging, rigid-diaphragm extension across zones, and tributary-area "
        "treatment in inclusion–exclusion form — is reserved for the "
        "irregular-plan validation campaign identified in Section 5.3."
    ),
]

TABLE_8_HEADING = "Table 8. L-shaped application example: model size and gravity-load equilibrium."

TABLE_8_HEADERS = ["Quantity", "Value", "Source / Cross-check"]
TABLE_8_ROWS = [
    ["Plan configuration", "L-shape: 12 × 8 m (5F) + 12 × 4 m (3F)", "Section 2.3.3, Figure 6"],
    ["Story height", "3.5 m", "Geometric input"],
    ["Zone structure", "2 rectangular zones (left wing 5F, right wing 3F)", "Consistent with the zone representation described in §2.3.3"],
    ["Nodes", "70", "After zone-boundary merging"],
    ["Elements", "135 (57 columns + 78 beams)", "Node-element graph"],
    ["Base supports (fixed)", "13", "Combined wing footprints"],
    ["Nominal floor area, left wing", "96 m² × 5 stories", "Per-story plan footprint"],
    ["Nominal floor area, right wing", "48 m² × 3 stories", "Per-story plan footprint"],
    ["DL total vertical base reaction", "3,458 kN", "Sum of nodal RZ; ratio with LL (2.04) matches DL/LL input ratio"],
    ["LL total vertical base reaction", "1,695 kN", "Sum of nodal RZ"],
    ["Horizontal base reactions (DL, LL)", "< 0.01 kN", "Static-equilibrium check"],
]
TABLE_8_COL_CM = [4.6, 5.8, 5.8]

FIG_6_CAPTION = (
    "Figure 6. Zone-based decomposition of the L-shaped plan example, "
    "showing the two rectangular zones and the merged boundary nodes "
    "at x = 12 m."
)

FIG_7_CAPTION = (
    "Figure 7. L-shaped node-element model assembled from the zone "
    "decomposition (five-story left wing and three-story right wing)."
)


# ──────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────

def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _populate_table(t, headers, rows, col_widths_cm, font_size=9) -> None:
    t.autofit = False
    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[ci].width = Cm(w)

    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
        _set_cell_borders(cell)

    for ri, row in enumerate(rows, start=1):
        for ci, v in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(v)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
            _set_cell_borders(cell)


def _find_anchor_index(doc, anchor_text: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if anchor_text in p.text:
            return i
    raise RuntimeError(f"Anchor text not found: {anchor_text!r}")


def _make_subheading(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


def _make_body(doc, text: str):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def _make_caption(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    return p


def _insert_after(anchor_para, elements):
    """Insert a list of (paragraph or table) XML elements after anchor.

    Elements are inserted in given order: anchor → elements[0] → elements[1] → ...
    addnext puts the new element immediately after anchor; to preserve order,
    insert in reverse so each successive call pushes prior insertions further.
    """
    for elem in reversed(elements):
        anchor_para._element.addnext(elem)


# ──────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────

def main() -> int:
    print(f"[*] Opening: {TARGET.relative_to(ROOT)}")
    d = docx.Document(str(TARGET))

    n_paras_before = len(d.paragraphs)
    n_tables_before = len(d.tables)

    # ──────────────────────────────────────────────────────────
    # 1) §4.3 insertion (LOWER in document — insert first so §2.3.3
    #    indices stay stable)
    # ──────────────────────────────────────────────────────────
    anchor_4_3_idx = _find_anchor_index(d, ANCHOR_TEXT_FIG5)
    anchor_4_3 = d.paragraphs[anchor_4_3_idx]
    print(f"[*] §4.3 anchor (para {anchor_4_3_idx}): {anchor_4_3.text[:80]!r}")

    # Create paragraphs/table at end-of-document, then move
    new_elements_4_3 = []
    sub = _make_subheading(d, SEC_4_3_HEADING)
    new_elements_4_3.append(sub._element)

    for body_text in SEC_4_3_BODY_PRE_TABLE:
        p = _make_body(d, body_text)
        new_elements_4_3.append(p._element)

    cap_t8 = _make_caption(d, TABLE_8_HEADING)
    new_elements_4_3.append(cap_t8._element)

    tbl_8 = d.add_table(rows=1 + len(TABLE_8_ROWS), cols=len(TABLE_8_HEADERS))
    _populate_table(tbl_8, TABLE_8_HEADERS, TABLE_8_ROWS, TABLE_8_COL_CM)
    new_elements_4_3.append(tbl_8._tbl)

    # Figure 6: zones top-view
    fig6_para = d.add_paragraph()
    fig6_para.add_run().add_picture(str(FIG_DIR / "fig6_lshape_zones.png"), width=Inches(5.5))
    new_elements_4_3.append(fig6_para._element)

    cap_f6 = _make_caption(d, FIG_6_CAPTION)
    new_elements_4_3.append(cap_f6._element)

    # Figure 7: 3D wireframe
    fig7_para = d.add_paragraph()
    fig7_para.add_run().add_picture(str(FIG_DIR / "fig7_lshape_3d.png"), width=Inches(5.5))
    new_elements_4_3.append(fig7_para._element)

    cap_f7 = _make_caption(d, FIG_7_CAPTION)
    new_elements_4_3.append(cap_f7._element)

    for body_text in SEC_4_3_BODY_POST_TABLE:
        p = _make_body(d, body_text)
        new_elements_4_3.append(p._element)

    _insert_after(anchor_4_3, new_elements_4_3)
    print(f"[OK] §4.3 inserted: {len(new_elements_4_3)} elements after para {anchor_4_3_idx}")

    # ──────────────────────────────────────────────────────────
    # 2) §2.3.3 insertion (UPPER — insert last to keep upper anchor stable)
    # ──────────────────────────────────────────────────────────
    # Reload (positions shifted by §4.3 insertion below, but anchor we search
    # for (§2.3.2 end) is above §4.3 insertion so its index is unchanged).
    # However, doc.paragraphs is rebuilt automatically — search again.
    anchor_2_3_3_idx = _find_anchor_index(d, ANCHOR_TEXT_2_3_2_END)
    anchor_2_3_3 = d.paragraphs[anchor_2_3_3_idx]
    print(f"[*] §2.3.3 anchor (para {anchor_2_3_3_idx}): {anchor_2_3_3.text[:80]!r}")

    new_elements_2_3_3 = []
    sub3 = _make_subheading(d, SEC_2_3_3_HEADING)
    new_elements_2_3_3.append(sub3._element)

    for body_text in SEC_2_3_3_BODY:
        p = _make_body(d, body_text)
        new_elements_2_3_3.append(p._element)

    _insert_after(anchor_2_3_3, new_elements_2_3_3)
    print(f"[OK] §2.3.3 inserted: {len(new_elements_2_3_3)} elements after para {anchor_2_3_3_idx}")

    # Save
    d.save(str(TARGET))

    # Verify
    d2 = docx.Document(str(TARGET))
    n_paras_after = len(d2.paragraphs)
    n_tables_after = len(d2.tables)
    print(f"[OK] Final: {n_paras_after} paragraphs (+{n_paras_after - n_paras_before}), "
          f"{n_tables_after} tables (+{n_tables_after - n_tables_before})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
