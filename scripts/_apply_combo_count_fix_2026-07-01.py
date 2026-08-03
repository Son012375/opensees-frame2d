"""Reconcile the §5 example's load-case / load-combination description with the
actual analysis (which the regenerated example_section4_results.json and the G3
ETABS writeup confirm: SEVEN cases incl. snow, THIRTY-SIX combinations incl. the
orthogonal 100%/30% seismic set). The manuscript's "six load cases (…) / 18
combinations" predates the orthogonal-seismic addition and is now contradicted in
the paper itself by Fig 5's governing-combo label (…+0.3EQX+0.2S).

  P126 : "The six load cases (DL, LL, EQX, EQY, WX, WY) … 18 load combinations …"
         -> add snow (0.5 kN/m2), "seven load cases (DL, LL, S, …)", "36 …",
            note the ± directional and orthogonal 100%/30% seismic combinations.
  P133 : "under the 18 load combinations evaluated" -> "36".
  Table 8 : "18 (KDS standard)" -> "36 (KDS standard)"; add a Snow | Roof | 0.5 kN/m2
            row (deep-copied from the Dead-Load row so the superscript m2 + cell
            formatting are preserved), inserted after the Wind row.

Snow value: os_response S-case base ΣFz ≈ 60 kN over the 120 m2 roof = 0.5 kN/m2.
Backup -> ..._v2.pre_combofix_2026-07-01.docx. Dry-run by default; --apply saves.
"""
from __future__ import annotations

import copy
import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_combofix_2026-07-01.docx")

P126_OLD = (" in the Y-direction. The six load cases (DL, LL, EQX, EQY, WX, WY) are "
            "combined into 18 load combinations following the KDS standard provisions")
P126_NEW = (" in the Y-direction, and a roof snow load of 0.5 kN/m2 is generated for the "
            "Seoul site. The seven load cases (DL, LL, S, EQX, EQY, WX, WY) are combined "
            "into 36 load combinations following the KDS standard provisions, including "
            "the positive and negative directional cases and the orthogonal 100%/30% "
            "seismic combinations")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs

    if any("seven load cases (DL, LL, S," in p.text for p in paras):
        raise SystemExit("ABORT: combo-count fix already applied.")

    # ---- P126 (run21 plain text) --------------------------------------------
    p126 = next(p for p in paras if "The six load cases (DL, LL, EQX, EQY, WX, WY)" in p.text)
    hit = [r for r in p126.runs if P126_OLD in r.text]
    if len(hit) != 1:
        raise SystemExit(f"ABORT: P126 target run matched {len(hit)} (want 1).")
    hit[0].text = hit[0].text.replace(P126_OLD, P126_NEW)

    # ---- P133 (18 -> 36) -----------------------------------------------------
    p133 = next(p for p in paras if "under the 18 load combinations evaluated" in p.text)
    r133 = next(r for r in p133.runs if "18 load combinations" in r.text)
    r133.text = r133.text.replace("18 load combinations", "36 load combinations")

    # ---- Table 8 -------------------------------------------------------------
    t8 = next(t for t in d.tables
              if [c.text for c in t.rows[0].cells][:3] == ["Category", "Parameter", "Value"]
              and len(t.rows) == 12)
    # 18 -> 36 in the Combinations cell (content-anchored)
    combo_cell = next(c for row in t8.rows for c in row.cells if "18 (KDS standard)" in c.text)
    cr = next(r for r in combo_cell.paragraphs[0].runs if "18 (KDS standard" in r.text)
    cr.text = cr.text.replace("18 (KDS standard", "36 (KDS standard")

    # add Snow row: deep-copy the Dead-Load row (row 7) to preserve superscript m2
    dead_idx = next(i for i, row in enumerate(t8.rows) if row.cells[0].text == "Dead Load")
    wind_idx = next(i for i, row in enumerate(t8.rows) if row.cells[0].text == "Wind")
    new_tr = copy.deepcopy(t8.rows[dead_idx]._tr)
    t8.rows[wind_idx]._tr.addnext(new_tr)
    snow = t8.rows[wind_idx + 1]  # freshly inserted row
    assert snow.cells[0].text == "Dead Load", f"copied row unexpected: {snow.cells[0].text!r}"
    snow.cells[0].paragraphs[0].runs[0].text = "Snow"
    snow.cells[1].paragraphs[0].runs[0].text = "Roof"
    snow.cells[2].paragraphs[0].runs[0].text = "0.5 "   # keeps 'kN' '/m' superscript '2'

    print("Planned edits:")
    print("  [P126] six->seven cases (+S), 18->36 combos, +snow 0.5 kN/m2, +orthogonal note")
    print("  [P133] 18 -> 36 load combinations")
    print("  [T8]   Combinations 18 -> 36; inserted Snow | Roof | 0.5 kN/m2 row")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply.")
        print("P126 tail:", p126.text[-230:])
        print("T8 snow row:", [snow.cells[i].text for i in range(3)])
        print("T8 rows now:", len(t8.rows))
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
