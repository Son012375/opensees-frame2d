"""Session 2026-07-02 (MAJOR REVISION): secondary — disclose the distributed-load path.

Reviewer secondary issue: the beam distributed-load application path (floor-area pressure ->
consistent member/nodal loads) was corrected on 2026-07-01 (end-node bias removed; see
memory g3_e2_gravity_fix); the five Midas Gen benchmark cases use point/nodal loads and never
exercise this path, though both Section 5 examples do. Add a one-sentence disclosure to the
§6.2 first limitation (P159). The current §5 numbers already reflect the corrected path
(example_section4_results.json regenerated 2026-07-01), so no manuscript numbers are stale.

Single append to P159 (no field). Backup -> ..._v2.pre_distload_2026-07-02.docx.
Dry-run default; --apply.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_distload_2026-07-02.docx")

LOC_P159 = "the benchmark-validated structural scope is restricted to regular orthogonal steel frames"
P159_APPEND = (
    " Separately, the beam distributed-load application path — which converts floor-area "
    "pressures to consistent member and nodal loads, and was corrected during this study to "
    "remove an end-node bias — is exercised by both Section 5 application examples but not by "
    "the five Midas Gen benchmark cases, which apply point and nodal loads; the Section 5 "
    "results reflect the corrected load path, and a dedicated distributed-load benchmark case "
    "against a commercial reference is left to future work."
)

SENTINEL = "the beam distributed-load application path"


def find_para(paras, needle, tag):
    hits = [p for p in paras if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ABORT [{tag}]: matched {len(hits)} (want 1): {needle[:50]!r}")
    return hits[0]


def count_fields(d):
    xml = d.element.xml
    return xml.count("fldChar"), xml.count("instrText")


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs
    if any(SENTINEL in p.text for p in paras):
        raise SystemExit("ABORT: already applied (distload sentinel present).")

    fld0, ins0 = count_fields(d)
    np0, nt0, ni0 = len(paras), len(d.tables), len(d.inline_shapes)
    print(f"BEFORE: paragraphs={np0} tables={nt0} inline_shapes={ni0} fldChar={fld0} instrText={ins0}")

    p159 = find_para(paras, LOC_P159, "P159")
    if "fldChar" in p159._p.xml or "instrText" in p159._p.xml:
        raise SystemExit("ABORT [P159]: paragraph carries a field")
    p159.add_run(P159_APPEND)

    fld1, ins1 = count_fields(d)
    np1, nt1, ni1 = len(d.paragraphs), len(d.tables), len(d.inline_shapes)
    print(f"AFTER : paragraphs={np1} tables={nt1} inline_shapes={ni1} fldChar={fld1} instrText={ins1}")
    assert (fld1, ins1) == (fld0, ins0), f"FIELD COUNT CHANGED {fld0}/{ins0} -> {fld1}/{ins1}"
    assert (np1, nt1, ni1) == (np0, nt0, ni0), "structure count changed"

    print(f"\n--- P159 now ({len(p159.text.split())} words) ---\n{p159.text}")

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
