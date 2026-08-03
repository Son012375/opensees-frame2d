# -*- coding: utf-8 -*-
"""Fig 8 caption: note it is a real 3D-viewer render (now a tool screenshot).

Field-safe single-run text edit; content anchor; backup; field-invariant verify.
Usage: python scripts/_apply_fig8_caption_2026-07-09.py [--apply]
"""
from __future__ import annotations

import re
import shutil
import sys
import zipfile
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_fig8caption_2026-07-09.docx")

OLD = "(five-story left wing and three-story right wing)."
NEW = "(five-story left wing and three-story right wing), rendered in the open-source 3D viewer."


def field_counts(p):
    fld = instr = 0
    with zipfile.ZipFile(p) as z:
        for n in z.namelist():
            if n.endswith(".xml"):
                x = z.read(n).decode("utf-8", "replace")
                fld += len(re.findall(r"<w:fldChar\b", x)); instr += len(re.findall(r"<w:instrText\b", x))
    return fld, instr


def main(apply):
    d = docx.Document(str(DOC))
    hits = [p for p in d.paragraphs if "assembled from the zone decomposition" in p.text and p.text.startswith("Figure 8")]
    if len(hits) != 1:
        raise SystemExit(f"ABORT: Fig 8 caption anchor matched {len(hits)} (want 1)")
    p = hits[0]
    if "rendered in the open-source 3D viewer" in p.text:
        raise SystemExit("ABORT: already applied")
    if p.text.count(OLD) != 1:
        raise SystemExit(f"ABORT: OLD occurs {p.text.count(OLD)}x")
    runs = [i for i, r in enumerate(p.runs) if OLD in r.text]
    if len(runs) != 1:
        raise SystemExit(f"ABORT: OLD spans runs ({len(runs)})")
    print("Fig8 caption ->", NEW)
    if not apply:
        print("DRY-RUN"); return 0
    fb, ib = field_counts(DOC)
    shutil.copy2(DOC, BACKUP)
    p.runs[runs[0]].text = p.runs[runs[0]].text.replace(OLD, NEW)
    d.save(str(DOC))
    fa, ia = field_counts(DOC)
    d2 = docx.Document(str(DOC))
    print(f"SAVED. paragraphs={len(d2.paragraphs)} tables={len(d2.tables)} shapes={len(d2.inline_shapes)} "
          f"fields {fb}/{ib}->{fa}/{ia} [{'OK' if (fb,ib)==(fa,ia) else 'CHANGED!!'}]")
    return 0


if __name__ == "__main__":
    raise SystemExit(main("--apply" in sys.argv))
