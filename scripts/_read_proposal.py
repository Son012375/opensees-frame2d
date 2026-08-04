# -*- coding: utf-8 -*-
import sys, os, glob
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

# Find the latest docx in Downloads matching pattern
candidates = glob.glob(r"C:\Users\youm\Downloads\*.docx")
target = None
for c in candidates:
    if "LLM-AE-AI" in c or "계획서" in c:
        target = c
        break
if target is None:
    target = max(candidates, key=os.path.getmtime)

print(f"FILE: {target}")
print()
d = Document(target)
print("=== PARAGRAPHS ===")
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t:
        print(f"[P{i}] {t}")
print()
print("=== TABLES ===")
for ti, t in enumerate(d.tables):
    print(f"--- Table {ti} ({len(t.rows)} rows x {len(t.columns)} cols) ---")
    for r in t.rows:
        cells = [c.text.strip().replace(chr(10), " / ") for c in r.cells]
        print(" | ".join(cells))
    print()
