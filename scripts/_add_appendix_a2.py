"""Appendix A KDS Load Generation Hand-Check를 v2 docx에 추가.

`_patch_paper_draft.py` 실행 후에 본 스크립트를 실행하면, References 마지막
이후에 Appendix A (본문 + 7 tables)이 추가된다.

운영 순서:
    1. python scripts/_patch_paper_draft.py      # 원본 → v2 (rewrites)
    2. python scripts/_add_appendix_a2.py        # v2 → v2 (appendix 추가)

본 스크립트는 idempotent하지 않다 — 매 실행마다 appendix를 *추가*하므로
재실행 전 _patch_paper_draft.py로 v2를 원본 기반 재생성해야 한다.

내용 출처: docs/paper1_open_source_alternative/rewrites/appendix_a2_draft.md
"""
from __future__ import annotations

from pathlib import Path

import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"


# ============================================================
# Helpers
# ============================================================

def add_top_heading(doc, text: str) -> None:
    """Top-level Appendix heading (bold, 14pt)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def add_sub_heading(doc, text: str) -> None:
    """A.2.X sub-heading (bold, 12pt)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def add_body(doc, text: str) -> None:
    """Regular body paragraph."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)


def add_table_caption(doc, text: str) -> None:
    """Table caption (bold italic, 10pt)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)


def _set_cell_borders(cell) -> None:
    """Add thin black borders to a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")  # 0.5pt
        border.set(qn("w:color"), "000000")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_data_table(doc, headers: list[str], rows: list[list[str]],
                   col_widths_cm: list[float] | None = None,
                   font_size: int = 9) -> None:
    """Add a data table with bold header + bordered cells."""
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.autofit = False

    # Set column widths if provided
    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[ci].width = Cm(w)

    # Header row
    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
        _set_cell_borders(cell)

    # Data rows
    for ri, row in enumerate(rows, start=1):
        for ci, v in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(v)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
            _set_cell_borders(cell)

    # Spacer after table
    doc.add_paragraph()


# ============================================================
# Table data
# ============================================================

TABLE_A2_1_HEADERS = ["#", "Quantity", "KDS clause", "Hand calc", "Auto-gen", "Δ", "Status"]
TABLE_A2_1_ROWS = [
    ["A1", "RC unit weight γ_c", "KDS 24 12 21 표 4.2-1", "24.0 kN/m³", "24.0", "0.0%", "✓"],
    ["A2", "Slab self-weight", "KDS 41 12 00 §2.1 (γ_c × t_slab)", "3.60 kN/m²", "3.60", "0.0%", "✓"],
    ["A3", "Finish, 1F (retail)", "Engineering input", "1.50 kN/m²", "1.50", "0.0%", "✓"],
    ["A4", "Finish, 2F–3F (office)", "Engineering input", "1.00 kN/m²", "1.00", "0.0%", "✓"],
    ["A5", "MEP allowance", "Workflow default", "0.50 kN/m²", "0.50", "—", "⚠"],
    ["A6", "DL total, 1F", "A2 + A3 + A5", "5.60 kN/m²", "5.60", "0.0%", "✓"],
    ["A7", "DL total, 2F–3F", "A2 + A4 + A5", "5.10 kN/m²", "5.10", "0.0%", "✓"],
]

TABLE_A2_2_HEADERS = TABLE_A2_1_HEADERS
TABLE_A2_2_ROWS = [
    ["B1", "LL, retail (판매장)", "KDS 41 12 00 표 3.1-1", "5.0 kN/m²", "5.00", "0.0%", "✓"],
    ["B2", "LL, office (사무실)", "KDS 41 12 00 표 3.1-1", "2.5 kN/m²", "2.50", "0.0%", "✓"],
]

TABLE_A2_3_HEADERS = TABLE_A2_1_HEADERS
TABLE_A2_3_ROWS = [
    ["C1", "Seismic zone factor z", "KDS 41 17 00 §3.1 (Seoul, Zone 1)", "0.11 g", "0.11", "0.0%", "✓"],
    ["C2", "Site coeff F_a", "KDS 17 10 00 표 4.2-8, S3, linear interp at z = 0.11", "1.68", "1.68", "0.0%", "✓"],
    ["C3", "Site coeff F_v", "KDS 17 10 00 표 4.2-8, S3, linear interp at z = 0.11", "1.69", "1.69", "0.0%", "✓"],
    ["C4", "Importance factor I_E", "KDS 41 17 00 §3.1.4", "1.0", "1.0", "0.0%", "✓"],
    ["C5", "S (단주기)", "KDS 17 10 00 §4.2.1, z × F_a", "0.1848 g", "0.1848", "0.0%", "✓"],
    ["C6", "S₁ (1초 주기)", "KDS 17 10 00 §4.2.1, z × F_v", "0.1859 g", "0.1859", "0.0%", "✓"],
    ["C7", "S_DS", "KDS 17 10 00 §4.2.1, 2.5 × S", "0.4620 g", "0.4620", "0.0%", "✓"],
    ["C8", "S_D1", "KDS 17 10 00 §4.2.1", "0.1859 g", "0.1859", "0.0%", "✓"],
    ["C9", "Response modification R", "KDS 41 17 00 표 6.2-1 (ordinary steel MF, 3-c)", "3.5", "3.5", "0.0%", "✓"],
    ["C10", "Overstrength Ω₀", "KDS 41 17 00 표 6.2-1", "3.0", "3.0", "0.0%", "✓"],
    ["C11", "Deflection amplification C_d", "KDS 41 17 00 표 6.2-1", "3.0", "3.0", "0.0%", "✓"],
    ["C12", "Period coefficient C_T", "KDS 41 17 00 §7.2.4 (1), steel MF", "0.0724", "0.0724", "0.0%", "✓"],
    ["C13", "Period exponent x", "KDS 41 17 00 §7.2.4 (1)", "0.80", "0.80", "0.0%", "✓"],
    ["C14", "Building height h_n", "Geometry, Σ story heights", "9.0 m", "9.0", "0.0%", "✓"],
    ["C15", "Approximate period T_a", "KDS 41 17 00 §7.2.4 (1), C_T × h_n^x", "0.4199 s", "0.4199", "0.0%", "✓"],
    ["C16", "Period upper limit C_u", "KDS 41 17 00 표 7.2-1, linear interp at S_D1 = 0.1859", "1.5282", "1.5282", "0.0%", "✓"],
    ["C17", "Design period T", "min(T_a, C_u·T_a); T = T_a", "0.4199 s", "0.4199", "0.0%", "✓"],
    ["C18", "C_s candidate (식 7.2-2)", "KDS 41 17 00 §7.2.2, S_DS·I_E/R", "0.1320", "0.1320", "0.0%", "✓"],
    ["C19", "C_s lower bound (식 7.2-5)", "KDS 41 17 00 §7.2.2, max(0.044·S_DS·I_E, 0.01)", "0.02033", "0.02033", "0.0%", "✓"],
    ["C20", "C_s upper bound (식 7.2-3)", "KDS 41 17 00 §7.2.2, S_D1·I_E/(T·R)", "0.1265", "0.1265", "0.0%", "✓"],
    ["C21", "C_s (final)", "C18 bounded by C19, C20; C20 governs", "0.1265", "0.1265", "0.0%", "✓"],
    ["C22", "Effective weight W_1 (1F)", "KDS 41 17 00 §7.2.1, DL_1F × A", "672.0 kN", "672.0", "0.0%", "✓"],
    ["C23", "Effective weight W_2 (2F)", "KDS 41 17 00 §7.2.1", "612.0 kN", "612.0", "0.0%", "✓"],
    ["C24", "Effective weight W_3 (3F)", "KDS 41 17 00 §7.2.1", "612.0 kN", "612.0", "0.0%", "✓"],
    ["C25", "Total weight W", "Σ W_i", "1,896.0 kN", "1,896.0", "0.0%", "✓"],
    ["C26", "Base shear V (식 7.2-1)", "KDS 41 17 00 §7.2.1, C_s · W", "239.84 kN", "239.84", "0.0%", "✓"],
    ["C27", "Distribution exponent k", "KDS 41 17 00 §7.2.5 (T ≤ 0.5 → k = 1.0)", "1.0", "1.0", "0.0%", "✓"],
]

TABLE_A2_4_HEADERS = ["Story", "w_i (kN)", "h_i (m)", "w_i · h_i", "C_vx", "F_x (kN)", "Status"]
TABLE_A2_4_ROWS = [
    ["1", "672.0", "3.0", "2,016", "0.1801", "43.19", "✓"],
    ["2", "612.0", "6.0", "3,672", "0.3280", "78.66", "✓"],
    ["3", "612.0", "9.0", "5,508", "0.4920", "117.99", "✓"],
    ["Σ", "1,896.0", "—", "11,196", "1.0000", "239.84", "✓"],
]

TABLE_A2_5_HEADERS = TABLE_A2_1_HEADERS
TABLE_A2_5_ROWS = [
    ["D1", "Basic wind speed V₀", "KDS 41 12 00 §5 (Seoul)", "30 m/s", "30", "0.0%", "✓"],
    ["D2", "Topographic factor K_zt", "KDS 41 12 00 §5.2.5 (flat terrain)", "1.0", "1.0", "—", "⚠"],
    ["D3", "Air density ρ", "Physical constant", "1.225 kg/m³", "1.225", "0.0%", "✓"],
    ["D4", "Exposure category", "Engineering input", "B", "B", "—", "✓"],
    ["D5", "Gradient height z_g (B)", "KDS 41 12 00 표 5.2-4", "365 m", "365", "0.0%", "✓"],
    ["D6", "Power-law exponent α (B)", "KDS 41 12 00 표 5.2-4", "9.5", "9.5", "0.0%", "✓"],
    ["D7", "Gust factor G_f (B)", "Workflow simplification", "2.20", "2.20", "—", "⚠"],
    ["D8", "Total pressure coefficient C_p", "Workflow simplification (0.8 + 0.5)", "1.30", "1.30", "—", "⚠"],
]

TABLE_A2_6_HEADERS = ["Story", "z_eff (m)", "K_z", "q_z (kN/m²)", "p (kN/m²)", "F_x (kN)", "F_y (kN)", "Status"]
TABLE_A2_6_ROWS = [
    ["1", "5.0", "0.8145", "0.3657", "1.0460", "31.38", "37.66", "✓"],
    ["2", "5.0", "0.8145", "0.3657", "1.0460", "31.38", "37.66", "✓"],
    ["3", "7.5", "0.8871", "0.4338", "1.2408", "37.22", "44.67", "✓"],
    ["Σ", "—", "—", "—", "—", "99.98", "119.99", "✓"],
]

TABLE_A2_7_HEADERS = ["Group", "Count", "Combinations"]
TABLE_A2_7_ROWS = [
    ["Gravity only", "2", "1.4 DL; 1.2 DL + 1.6 LL"],
    ["Seismic X (±)", "4", "1.2 DL + 1.0 LL ± 1.0 EQX; 0.9 DL ± 1.0 EQX"],
    ["Seismic Y (±)", "4", "1.2 DL + 1.0 LL ± 1.0 EQY; 0.9 DL ± 1.0 EQY"],
    ["Wind X (±)", "4", "1.2 DL + 1.0 LL ± 1.0 WX; 0.9 DL ± 1.0 WX"],
    ["Wind Y (±)", "4", "1.2 DL + 1.0 LL ± 1.0 WY; 0.9 DL ± 1.0 WY"],
    ["Total", "18", "matches the count reported in Section 4.1"],
]

# Column widths (cm) for each table — tuned to fit page width
W_NARROW_7 = [1.0, 3.5, 4.5, 2.4, 2.4, 1.4, 1.2]  # tables A.2-1/2/3/5
W_STORY_DIST = [1.4, 2.0, 1.8, 1.8, 1.6, 1.8, 1.6]  # A.2-4
W_WIND_STORY = [1.2, 1.6, 1.6, 1.8, 1.8, 1.6, 1.6, 1.2]  # A.2-6
W_COMBO = [3.0, 1.5, 11.0]  # A.2-7


# ============================================================
# Main: build appendix into v2 docx
# ============================================================

def main() -> int:
    print(f"[*] Opening: {TARGET.relative_to(ROOT)}")
    d = docx.Document(str(TARGET))
    n_paras_before = len(d.paragraphs)
    n_tables_before = len(d.tables)

    # Page break before appendix
    d.add_page_break()

    # ─── Appendix A top heading ────────────────────────────
    add_top_heading(d, "Appendix A. KDS Load Generation Hand-Check")
    d.add_paragraph()  # spacer

    # ─── A.1 Scope and Reference Inputs ────────────────────
    add_sub_heading(d, "A.1 Scope and Reference Inputs")
    add_body(d,
        "This appendix records the clause-by-clause verification of the KDS-based "
        "load values generated by the open-source pipeline for the three-story "
        "regular steel frame example described in Section 4.1. For each quantity, "
        "the value retrieved from the workflow (the auto-generated value) is "
        "compared against the value obtained by direct evaluation of the "
        "corresponding Korean Design Standards clause (the hand-calculated value). "
        "Forty-nine quantities are tabulated across dead, live, equivalent lateral "
        "seismic, wind, and combination categories."
    )
    add_body(d,
        "The reference building has three stories of 3.0 m uniform height arranged "
        "on a 3 × 2 bay plan (X-direction span 4.0 m, Y-direction span 5.0 m), "
        "with 0.15 m reinforced-concrete slabs, H-300×300 columns, and H-400×200 "
        "beams of SS275 steel. The first story is occupied as retail and the "
        "second and third stories as office space, with the site located in Seoul "
        "(z = 0.11 g), Site Class S3, Importance Class II (I_E = 1.0), assigned "
        "to an ordinary steel moment frame, and exposed to wind exposure category "
        "B with a basic wind speed V₀ = 30 m/s."
    )
    add_body(d,
        "Status symbols used throughout the tables: ✓ indicates agreement within "
        "rounding (Δ ≤ 0.1%); ⚠ indicates a workflow-introduced assumption or "
        "simplification disclosed in Section 2.4."
    )

    # ─── A.2 Dead Load ──────────────────────────────────────
    add_sub_heading(d, "A.2 Dead Load (KDS 41 12 00 §2, KDS 24 12 21 §4.2)")
    add_table_caption(d, "Table A-1. Dead-load hand-check.")
    add_data_table(d, TABLE_A2_1_HEADERS, TABLE_A2_1_ROWS, col_widths_cm=W_NARROW_7)
    add_body(d,
        "The MEP allowance (A5) is a workflow constant of 0.5 kN/m² applied in "
        "lieu of fixture-specific computation; KDS 41 12 00 §2 leaves equipment "
        "loads to engineering judgment. This is the first of the four workflow "
        "simplifications disclosed in Section 2.4."
    )

    # ─── A.3 Live Load ──────────────────────────────────────
    add_sub_heading(d, "A.3 Live Load (KDS 41 12 00 §3 표 3.1-1)")
    add_table_caption(d, "Table A-2. Live-load hand-check.")
    add_data_table(d, TABLE_A2_2_HEADERS, TABLE_A2_2_ROWS, col_widths_cm=W_NARROW_7)

    # ─── A.4 Seismic ────────────────────────────────────────
    add_sub_heading(d,
        "A.4 Equivalent Lateral Seismic Load "
        "(KDS 41 17 00 §6.2, §7.2; KDS 17 10 00 §4.2.1)"
    )
    add_table_caption(d, "Table A-3. Seismic equivalent-lateral-force hand-check.")
    add_data_table(d, TABLE_A2_3_HEADERS, TABLE_A2_3_ROWS, col_widths_cm=W_NARROW_7)
    add_table_caption(d,
        "Table A-4. Story-wise seismic-force distribution (식 7.2-8, 7.2-9)."
    )
    add_data_table(d, TABLE_A2_4_HEADERS, TABLE_A2_4_ROWS, col_widths_cm=W_STORY_DIST)
    add_body(d,
        "Two adjustments to the implementation were made during the hand-check "
        "exercise. First, the F_a and F_v site coefficients were originally "
        "retrieved from KDS 17 10 00 표 4.2-8 by bucketing the seismic zone "
        "factor z into the three table anchor rows (z ≤ 0.1, z = 0.2, z = 0.3); "
        "the same table specifies linear interpolation between the anchor "
        "values, and the lookup was revised to compute F_a = 1.68 and F_v = 1.69 "
        "at z = 0.11 by interpolation between the S ≤ 0.1 and S = 0.2 rows of "
        "the S3 site class. Second, the period upper-limit coefficient C_u was "
        "originally retrieved from KDS 41 17 00 표 7.2-1 by bucketing the design "
        "spectral acceleration S_D1; the same table likewise specifies linear "
        "interpolation, and the lookup was revised to compute C_u = 1.5282 at "
        "S_D1 = 0.1859 by interpolation between the table rows at S_D1 = 0.15 "
        "and S_D1 = 0.20. The C_u adjustment does not propagate to the base "
        "shear because the implementation uses the approximate period T = T_a "
        "directly without modal computation, and the condition T_a < C_u · T_a "
        "is satisfied for both values. Both adjustments are reflected in the "
        "values reported above."
    )

    # ─── A.5 Wind ───────────────────────────────────────────
    add_sub_heading(d, "A.5 Wind Load (KDS 41 12 00 §5)")
    add_table_caption(d, "Table A-5. Wind parameter hand-check.")
    add_data_table(d, TABLE_A2_5_HEADERS, TABLE_A2_5_ROWS, col_widths_cm=W_NARROW_7)
    add_table_caption(d, "Table A-6. Story-wise wind force.")
    add_data_table(d, TABLE_A2_6_HEADERS, TABLE_A2_6_ROWS, col_widths_cm=W_WIND_STORY)
    add_body(d,
        "Rows D2, D7, and D8 are the three wind-related items among the four "
        "workflow simplifications disclosed in Section 2.4. Specifically: the "
        "topographic factor K_zt is held at 1.0 corresponding to flat terrain "
        "(K_zt would otherwise be computed per KDS §5.2.5 from the site "
        "topography); the gust factor G_f is taken as a fixed value per exposure "
        "category (2.20 for category B) rather than computed through the full "
        "KDS §5 expression that depends on building geometry, period, and "
        "damping; and the total wind pressure coefficient C_p is set to 1.3 "
        "representing 0.8 (windward) + 0.5 (leeward) without aspect-ratio "
        "refinement from the KDS pressure-coefficient table. The effective "
        "velocity-pressure height z_eff is also clamped to a minimum of 5 m, "
        "which causes Stories 1 and 2 to share the same K_z because both story "
        "centroids fall below the 5 m threshold."
    )

    # ─── A.6 Load Combinations ──────────────────────────────
    add_sub_heading(d, "A.6 Load Combinations (KDS 41 12 00 §1.7)")
    add_table_caption(d, "Table A-7. Load combination set.")
    add_data_table(d, TABLE_A2_7_HEADERS, TABLE_A2_7_ROWS, col_widths_cm=W_COMBO)
    add_body(d,
        "The KDS 41 12 00 §1.7 식 3 combinations involving roof live load (L_r), "
        "snow (S), or rain (R) are not generated for this example because those "
        "load cases are not produced by the workflow for the Seoul site."
    )

    # ─── A.7 Summary ────────────────────────────────────────
    add_sub_heading(d, "A.7 Summary")
    add_body(d,
        "Across all 49 individually checked quantities, the auto-generated values "
        "agree with the corresponding hand calculation to within rounding "
        "(Δ ≤ 0.1%), with the exception of four entries (A5, D2, D7, D8) that "
        "are explicit workflow simplifications disclosed in Section 2.4 rather "
        "than values derived from a KDS clause. No material discrepancy was "
        "identified, and the eighteen-combination set matches the KDS 41 12 00 "
        "§1.7 load-combination requirement for the available load cases. The "
        "hand-check exercise itself prompted two corrections to the "
        "implementation — the linear interpolation of F_a and F_v in KDS 17 10 "
        "00 표 4.2-8 and of C_u in KDS 41 17 00 표 7.2-1 — which are noted in "
        "Section A.4 and reflected in the tabulated values."
    )

    d.save(str(TARGET))

    n_paras_after = len(d.paragraphs)
    n_tables_after = len(d.tables)
    print(f"[OK] Appendix A added to {TARGET.relative_to(ROOT)}")
    print(f"     Paragraphs: {n_paras_before} → {n_paras_after} (+{n_paras_after - n_paras_before})")
    print(f"     Tables:     {n_tables_before} → {n_tables_after} (+{n_tables_after - n_tables_before})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
