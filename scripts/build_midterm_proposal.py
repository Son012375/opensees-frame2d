# -*- coding: utf-8 -*-
"""
LLM-AE-AI 중간발표 계획서 생성 스크립트
출력: C:\\Users\\youm\\Desktop\\백장운_수업\\LLM-AE-AI_중간발표_계획서_손명국.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = Path(r"C:\Users\youm\Desktop\백장운_수업")
OUT_PATH = OUT_DIR / "LLM-AE-AI_중간발표_계획서_손명국.docx"

FONT = "맑은 고딕"


def set_korean_font(run, size_pt=10, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), FONT)
    rfonts.set(qn("w:cs"), FONT)


def add_para(doc, text, size=10, bold=False, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.3
    run = p.add_run(text)
    set_korean_font(run, size_pt=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 14, 2: 11}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_korean_font(run, size_pt=sizes.get(level, 11), bold=True)
    return p


def add_bullet(doc, text, size=10, indent_cm=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run("• " + text)
    set_korean_font(run, size_pt=size)
    return p


def add_table(doc, headers, rows, col_widths_cm=None, header_bold=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_korean_font(run, size_pt=9, bold=header_bold)
        # 헤더 음영
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "DCE6F1")
        tc_pr.append(shd)
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.2
            run = p.add_run(str(val))
            set_korean_font(run, size_pt=9)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return table


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # 기본 스타일
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # ──────────────── 표제 ────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LLM-AE-AI 중간발표 계획서")
    set_korean_font(run, size_pt=18, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    run = sub.add_run("V2 + KDS 설계 에이전트: IFC 기반 모델링과 KDS 근거 기반 자동 설계 보조")
    set_korean_font(run, size_pt=11, bold=False)

    # 발표자 정보 표
    info = doc.add_table(rows=1, cols=4)
    info.style = "Table Grid"
    cells = info.rows[0].cells
    info_pairs = [("이름", "손명국"), ("학번", "2026310208")]
    idx = 0
    for label, val in info_pairs:
        cells[idx].text = ""
        p = cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_korean_font(run, size_pt=9, bold=True)
        tc_pr = cells[idx]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "DCE6F1")
        tc_pr.append(shd)
        cells[idx + 1].text = ""
        p = cells[idx + 1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        set_korean_font(run, size_pt=9)
        idx += 2

    info2 = doc.add_table(rows=1, cols=4)
    info2.style = "Table Grid"
    cells = info2.rows[0].cells
    info_pairs2 = [("지도교수", "백장운 교수님"), ("강의명", "LLM-AE-AI")]
    idx = 0
    for label, val in info_pairs2:
        cells[idx].text = ""
        p = cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_korean_font(run, size_pt=9, bold=True)
        tc_pr = cells[idx]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "DCE6F1")
        tc_pr.append(shd)
        cells[idx + 1].text = ""
        p = cells[idx + 1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        set_korean_font(run, size_pt=9)
        idx += 2

    add_para(doc, "", size=2, space_after=2)

    # ──────────────── §1 ────────────────
    add_heading(doc, "1. 프로젝트 제목 및 한 줄 요약")
    add_para(
        doc,
        "프로젝트명: V2 + KDS 설계 에이전트 — IFC 기반 모델링과 KDS 근거 기반 자동 설계 보조 시스템.",
    )
    add_para(
        doc,
        "한 줄 요약: IFC 파일을 파싱해 3D 구조모델을 자동 생성·해석하고, NG 부재가 발생하면 LLM이 KDS 조항을 인용해 단면·재료 변경안을 제안·재해석까지 수행하는 LLM-네이티브 구조해석 어시스턴트.",
    )

    # ──────────────── §2 ────────────────
    add_heading(doc, "2. 문제 정의 및 동기")
    add_para(
        doc,
        "발표자가 진행 중인 OpenSees-MCP 프로젝트(이하 V2)는 IFC 파일 기반의 3D 노드-요소 에디터, KDS 자동 하중 생성, OpenSeesPy 해석, KDS·AISC 기반 Design Check를 갖춘 구조해석 파이프라인이다. Midas Gen과의 5개 벤치마크 케이스 비교(112 metrics)를 통해 해석 엔진의 정확도는 검증되었다. 그러나 실무적 활용 관점에서 다음 세 가지 한계가 남아 있다.",
    )
    add_bullet(
        doc,
        "한계 1 — 근거 부재: Design Check가 NG/OK 판정과 ratio 값을 제공하지만, 적용된 KDS 조항의 본문·출처·해설이 결과 옆에 함께 표시되지 않는다. 설계 보고서로 직접 활용하기 어렵다.",
    )
    add_bullet(
        doc,
        "한계 2 — NG 이후의 시행착오: 부재가 NG로 판정되면 사용자가 단면을 직접 변경하고 재해석을 반복해야 한다. 어떤 단면 후보가 KDS 기준을 만족시킬지에 대한 안내가 없다.",
    )
    add_bullet(
        doc,
        "한계 3 — KDS 본문 접근성: KDS 41 12 00 / 41 17 00 / 41 31 00 등 본문은 PDF로 파편화되어 있어, 비-구조 전공자(건축 설계자, 학생)가 적용 조항을 찾는 비용이 크다.",
    )
    add_para(
        doc,
        "본 과제는 위 세 한계를 LLM 기반으로 해결하는 것을 목표로 한다. 핵심 차별점은 단순 챗봇이 아니라 진단·근거·수정안·재해석을 한 루프 안에서 수행하는 능동 설계 에이전트라는 점이며, 이는 W7 게스트 세미나 예시(Solver→Self-Improvement→Verifier→Correction→Synthesis 5단계 프롬프트)와 정확히 부합한다.",
    )

    # ──────────────── §3 ────────────────
    add_heading(doc, "3. 목표 사용자 및 사용 시나리오")
    add_para(doc, "두 종류의 페르소나를 설정한다.", bold=False)
    add_bullet(
        doc,
        "페르소나 A — 구조 초년 설계자: 모델링과 해석은 가능하나 KDS 조항을 빠르게 찾기 어렵고, NG 부재의 적정 단면 선정에 시행착오가 많다. 본 시스템은 NG 부재에 대해 KDS 근거를 제시하고 후보 단면을 자동 제안하여 학습 도구이자 실무 보조 도구로 동작한다.",
    )
    add_bullet(
        doc,
        "페르소나 B — 건축 설계자(비-구조 전공): IFC 파일을 던져 초기 안의 구조 안정성을 빠르게 확인하고, 자연어로 KDS 본문을 질의하길 원한다. 예: “사무실의 활하중 KDS 기준은?”.",
    )
    add_para(doc, "대표 사용 흐름은 다음 5단계이다.")
    add_bullet(doc, "① IFC 파일 업로드 → V2가 노드/요소/단면/슬래브를 자동 파싱하여 3D 모델 생성.")
    add_bullet(doc, "② KDS 자동 하중 생성(DL/LL/EQ/Wind, 18개 조합) 및 OpenSeesPy 해석 수행.")
    add_bullet(
        doc,
        "③ Design Check 결과 패널에 KDS 조항이 자동 인용된다(F2). 사용자는 자연어로 KDS 본문을 자유 질의할 수 있다(F1).",
    )
    add_bullet(
        doc,
        "④ NG 부재가 검출되면 설계 에이전트(F3)가 KDS 근거 + 단면/재료 변경안을 제시하고, 사용자가 승인하면 모델을 자동 수정·재해석한다.",
    )
    add_bullet(
        doc,
        "⑤ Verifier(F4)가 매 반복마다 제안의 적정성을 평가해 환각·과설계를 차단하며, 수렴 또는 max_iter 도달 시 종료한다.",
    )

    # ──────────────── §4 ────────────────
    add_heading(doc, "4. 핵심 기능 (입력·처리·출력 명세)")
    add_para(
        doc,
        "본 과제는 4개의 핵심 기능으로 구성된다. 각 기능의 입력·처리·출력과 활용 학습 주차를 명시한다.",
    )
    add_table(
        doc,
        headers=["#", "기능명", "입력 → 처리 → 출력", "주요 활용 주차"],
        rows=[
            [
                "F1",
                "KDS RAG 챗봇",
                "자연어 질의 → 하이브리드 검색(임베딩+BM25+RRF) → KDS 조항 본문 + 출처가 포함된 응답.",
                "W2, W5, W6",
            ],
            [
                "F2",
                "Design Check 자동 인용",
                "Design Check 결과(부재 ID, ratio, 적용기준) → 관련 조항 RAG 검색 → Citations 형태로 보고서에 첨부.",
                "W5, W6",
            ],
            [
                "F3",
                "KDS 기반 설계 제안 에이전트",
                "NG 부재 + 해석 결과 → 원인 진단(Solver) → 단면·재료 후보(KDS·단면 DB 738건 검색) → MCP Tool로 모델 수정 → 재해석 → 수렴까지 반복.",
                "W4, W5, W7",
            ],
            [
                "F4",
                "LLM-as-Judge Verifier",
                "F3의 매 반복 제안 → 평가 프롬프트(과설계/환각/기준 누락 검사) → 채택/기각 + 사유 출력.",
                "W3",
            ],
        ],
        col_widths_cm=[1.0, 3.6, 8.0, 2.6],
    )
    add_para(
        doc,
        "F3의 5단계 프롬프트 구조: ① Solver(NG 원인 진단) → ② Self-Improvement(후보 다양화) → ③ Verifier(KDS 적합성 평가) → ④ Correction(모델 수정 명령 생성) → ⑤ Synthesis(채택안 + 근거 정리). 이 5단계는 W7 MCP Prompts로 노출하여 Claude Desktop / MCP Inspector에서 직접 호출 가능하도록 한다.",
    )

    # ──────────────── §5 ────────────────
    add_heading(doc, "5. 사용 기술 스택 (학습 주차 매핑)")
    add_table(
        doc,
        headers=["주차", "학습 기법", "본 프로젝트 활용", "관련 기능"],
        rows=[
            ["W1", "프롬프트 엔지니어링 6기법", "진단·제안용 도메인 프롬프트 설계(Few-shot, CoT, Role).", "F3"],
            ["W2", "Claude API + JSON 추출", "BuildingIntent, 제안 후보 구조화 추출.", "F1, F3"],
            ["W3", "LLM-as-Judge", "Verifier 단계 — 제안 적정성 자동 평가.", "F4"],
            ["W4", "Tool Use (MCP 도구)", "change_section / change_material / re-analyze 루프 도구 호출.", "F3"],
            ["W5", "RAG 풀파이프라인", "KDS PDF 청킹 + Voyage 임베딩 + BM25 + RRF 하이브리드.", "F1, F2, F3"],
            ["W6", "Citations / PDF / Prompt Caching", "KDS 조항 출처 표시, 대용량 PDF 직접 입력, 캐싱으로 비용 절감.", "F1, F2"],
            ["W7", "FastMCP — Tools + Prompts", "기존 V2 MCP 14 tools + 5단계 설계 프롬프트를 표준 프로토콜로 노출.", "F3"],
        ],
        col_widths_cm=[1.2, 3.6, 7.2, 3.0],
    )

    # ──────────────── §6 ────────────────
    add_heading(doc, "6. 시스템 아키텍처")
    add_para(
        doc,
        "전체 데이터 흐름은 다음과 같다. IFC 파일이 메인 입력이며, 자연어는 보조 질의·수정 인터페이스로 동작한다.",
    )

    arch = (
        "[IFC 파일] ──> V2 IFC Parser ──> BuildingModel(IR) ──> KDS Load Generator ──> OpenSeesPy 해석\n"
        "                                                                                   │\n"
        "                                                                                   ▼\n"
        "                                                                           Design Check (KDS+AISC)\n"
        "                                                                                   │\n"
        "          ┌────────────────────────────────────────────────────────────────────────┘\n"
        "          ▼\n"
        "   ┌──────────────────────────────────────────────────────────────┐\n"
        "   │  Claude API  (Extended Thinking · JSON 추출 · Streaming)      │\n"
        "   │   ├─ MCP Prompts (W7) ── Solver → Self-Imp → Verifier         │\n"
        "   │   │                       → Correction → Synthesis            │\n"
        "   │   ├─ Tool Use (W4)  ──> V2 MCP Server (analyze_building,      │\n"
        "   │   │                       design_check, change_section …)    │\n"
        "   │   ├─ RAG (W5)       ──> KDS Vector DB + 단면 DB 738건         │\n"
        "   │   │                       (Voyage 임베딩 + BM25 + RRF)         │\n"
        "   │   └─ Citations (W6) ──> 조항 출처 표시 + Prompt Caching        │\n"
        "   └──────────────┬───────────────────────────────────────────────┘\n"
        "                  ▼\n"
        "        V2 3D Editor (UI)  ◄── NL 보조 질의/수정 입력\n"
        "                  │\n"
        "                  └── (NG 시 재해석 루프) ──> 위 OpenSeesPy 해석으로 재진입\n"
    )

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(4)
    run = p.add_run(arch)
    set_korean_font(run, size_pt=8)
    # 등폭 글꼴 적용
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")

    add_para(
        doc,
        "핵심 설계 결정: ① IFC가 메인 입력이므로 사용자는 BIM 도구에서 만든 모델을 그대로 가져올 수 있다. ② 자연어는 RAG 질의(F1)와 부분 모델 수정 명령에 한정한다. ③ 설계 에이전트의 모든 제안은 OpenSeesPy 재해석을 거쳐 ground truth를 확보하므로, LLM 출력 자체가 결과로 사용되지 않는다.",
    )

    # ──────────────── §7 ────────────────
    add_heading(doc, "7. 개발 일정 (기말발표까지 마일스톤)")
    add_para(
        doc,
        "총 기간 미정(추후 학사일정에 따라 확정). 발표일 2026-05-12를 W0로 두고 6주 작업을 가정한 상대 일정이며, 각 주의 종료 시점 기준 산출물을 명시한다.",
    )
    add_table(
        doc,
        headers=["주차(상대)", "주요 작업", "산출물"],
        rows=[
            ["W1~", "KDS 대상 문서 선정·수집, 청킹 전략 결정, Voyage 임베딩 파이프라인 구축.", "KDS 벡터 DB 1차"],
            ["W2~", "F1 RAG 챗봇 구현, BM25/RRF 하이브리드 검색, Citations 포맷 통합.", "KDS 자유 질의 동작"],
            ["W3~", "F3 에이전트의 Solver + Correction 단계 구현, MCP Tool 연동.", "NG → 1차 제안 동작"],
            ["W4~", "F4 Verifier + Synthesis 통합, max_iter·승인 게이트 구현.", "수렴 루프 동작"],
            ["W5~", "F2 Design Check 자동 인용 + V2 UI에 RAG/제안 패널 통합.", "End-to-End 데모"],
            ["W6~", "MCP Prompts 5단계 노출, MCP Inspector 시연 영상 제작, 기능-주차 매핑서 정리.", "기말 제출본"],
        ],
        col_widths_cm=[2.4, 9.0, 3.6],
    )

    # ──────────────── §8 ────────────────
    add_heading(doc, "8. 예상 데모 화면 및 산출물 형태")
    add_para(
        doc,
        "최종 데모는 V2 웹 에디터 위에 RAG·제안 패널을 통합한 단일 화면으로 시연한다. 좌측은 기존 3D 뷰어와 부재 테이블, 우측은 KDS 인용 + 제안 카드가 위치한다.",
    )
    add_para(doc, "[Before / After 시연 시나리오]", bold=True, space_after=2)
    add_bullet(
        doc,
        "Before: 10층 IFC 모델 해석 후 3층 기둥(H-300x300) 압축-휨 상관식 ratio 1.05 → 빨간색 NG 표시.",
    )
    add_bullet(
        doc,
        "Agent: KDS 41 31 00 §H1.1 본문 인용(P_r/φP_n + 8/9·M_r/φM_n ≤ 1.0) + 단면 DB에서 H-350×350 후보 제시.",
    )
    add_bullet(
        doc,
        "After: 사용자 승인 후 모델 자동 수정 + OpenSeesPy 재해석 → ratio 0.78 OK(초록), 변경 이력은 사이드 패널에 누적.",
    )
    add_para(doc, "기말 제출 산출물:", bold=True, space_after=2)
    add_bullet(doc, "GitHub 저장소 (Son012375/opensees-frame2d) — CLAUDE.md, README, 의존성 명세 포함.")
    add_bullet(doc, "MCP Inspector / Claude Desktop 시연 영상 또는 스크린샷 모음.")
    add_bullet(doc, "기능-학습 주차 매핑서(이 계획서의 §5 표를 구현 결과로 갱신한 버전).")

    # ──────────────── §9 ────────────────
    add_heading(doc, "9. 위험 요소 및 완화 방안")
    add_table(
        doc,
        headers=["위험", "영향", "완화 방안"],
        rows=[
            [
                "환각·과설계 (LLM 제안이 실제 KDS와 불일치)",
                "잘못된 단면 채택 시 안전성 저하",
                "F4 Verifier(W3 LLM-as-Judge) 검증 + 모든 결과는 OpenSeesPy 재해석 ground truth로 확정.",
            ],
            [
                "KDS 본문 저작권",
                "원문 무단 게재 시 저작권 침해 우려",
                "발췌 인용 + Citations로 출처 표기, 본문 전체 게재 금지(과제 안내문 §10 준수).",
            ],
            [
                "에이전트 무한 루프",
                "재해석 비용·시간 증가",
                "max_iter=5, 수렴 임계(ratio 변화 < 1%) 설정, 사용자 승인 게이트 필수화.",
            ],
            [
                "API 비용",
                "KDS PDF 컨텍스트 크기로 비용 누적",
                "W6 Prompt Caching으로 KDS 컨텍스트 캐시 재사용, 검색 결과만 매 호출 갱신.",
            ],
            [
                "일정 지연",
                "기말 데모 미완성 가능성",
                "V2 본체는 이미 완성 → RAG·에이전트 레이어로 범위 한정. F1·F2 우선, F3·F4는 단계별 적용.",
            ],
            [
                "API 키 노출",
                ".env 키가 GitHub에 커밋될 위험",
                ".gitignore 등록 + pre-commit 점검, 발표자료에는 키 비포함.",
            ],
        ],
        col_widths_cm=[3.6, 4.0, 7.4],
    )

    # ──────────────── §10 ────────────────
    add_heading(doc, "10. 참고자료")
    add_bullet(doc, "한국건설기준 KDS 41 12 00 (건축구조기준 설계하중), KDS 41 17 00 (내진설계기준), KDS 41 31 00 (강구조설계기준).")
    add_bullet(doc, "AISC 360 Specification for Structural Steel Buildings (압축·휨·상관식).")
    add_bullet(doc, "KS D 3502 / KS D 3568 — 단면 DB 738건 기준.")
    add_bullet(doc, "OpenSeesPy 공식 문서 (https://openseespydoc.readthedocs.io).")
    add_bullet(doc, "ifcopenshell 0.8.x 공식 문서.")
    add_bullet(doc, "Anthropic Claude API · Tool Use · Prompt Caching · Citations 공식 문서.")
    add_bullet(doc, "FastMCP / Model Context Protocol 공식 문서 (https://modelcontextprotocol.io).")
    add_bullet(doc, "VoyageAI Embeddings 공식 문서.")
    add_bullet(doc, "본인 GitHub 저장소: https://github.com/Son012375/opensees-frame2d.")

    doc.save(OUT_PATH)
    print(f"[OK] saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
