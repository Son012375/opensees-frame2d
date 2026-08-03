"""§5 manuscript revision after the V1/V2 gravity fix + case6 resolution (2026-07-01).

Applies A1-A4 + B1 from
docs/paper1_open_source_alternative/validation/manuscript_revision_checklist_2026-07-01.md
to the canonical v2 DOCX, editing PLAIN-TEXT runs only (no EndNote/Zotero/SEQ fields).

  A1  P131  roof disp 5.4/12.0 -> 5.3/11.9 mm ; design drift 0.22/0.47 -> 0.21/0.46 %
  A2  P132  max interaction 0.507 (2F col, H1-1b) -> 0.467 (1F col, H1-1a) ; beam 0.322 -> 0.324
  A3  T9    same numbers in the summary table (Table 9)
  A4  P144  §5.3 ETABS cross-check paragraph -> resolved-agreement prose (gap -> ~0.7%,
            setback moment exact); keeps the unchanged 36.5/14.24/3,458 examples + §6.3 reservation
  B1  P130  add one sentence stating SS275 nominal E = 205,000 MPa for the example

Backup -> ..._v2.pre_5rev_2026-07-01.docx. Dry-run by default; pass --apply to save.
Idempotency-guarded: aborts if any expected OLD string is missing or a NEW marker
is already present.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_5rev_2026-07-01.docx")

A4_NEW = (
    "As an independent cross-check of the irregular-plan model, the same L-shaped "
    "configuration was also analyzed in ETABS under the four load cases, with both "
    "programs assigned identical member sections, supports, and the same "
    "tributary-distributed gravity loads. The global response agrees closely: the "
    "Zone A roof displacements, the maximum inter-story drifts, and the inter-zone "
    "torsional displacement difference match to within 0.1% under both the X- and "
    "Y-direction equivalent-static cases (for example, 36.5 mm versus 36.5 mm for the "
    "Zone A fifth-floor X-displacement and 14.24 mm versus 14.24 mm for the X-direction "
    "torsional difference), and the total vertical base reaction matches to 0.0% "
    "(3,458 kN dead load in both programs). The per-support gravity reactions also agree "
    "to within about 0.7%, including at the re-entrant boundary: the Zone A corner "
    "support carries 156.3 versus 155.2 kN, the shared-boundary support 248.5 versus "
    "246.8 kN, and the far Zone B corner 90.3 versus 90.3 kN, the small residual "
    "reflecting the section-catalog cross-sectional area difference rather than the load "
    "path. Representative member end moments agree as well: the setback column's "
    "double-curvature strong-axis moment matches at every output station (for example, "
    "-55.2 versus -55.2 kN·m at the upper end), with the two programs differing only "
    "in the local-axis sign convention for the weak-axis and corner-column moments. A "
    "full per-member force comparison together with a Midas Gen cross-check is reserved "
    "for the dedicated irregular-plan validation campaign (Section 6.3)."
)

B1_SENTENCE = (
    " The steel members use SS275 with a nominal Young's modulus of E = 205,000 MPa, "
    "the KS value for this grade (the controlled benchmark cases of Section 4 use "
    "210,000 MPa)."
)


def _replace_in_run(run, old, new, log, tag):
    if old not in run.text:
        raise SystemExit(f"ABORT [{tag}]: expected old string not found: {old!r}")
    run.text = run.text.replace(old, new)
    log.append(f"  [{tag}] {old!r} -> {new!r}")


def _replace_cell(cell, old, new, log, tag):
    par = cell.paragraphs[0]
    if len(par.runs) != 1:
        raise SystemExit(f"ABORT [{tag}]: cell has {len(par.runs)} runs (want 1): {par.text!r}")
    _replace_in_run(par.runs[0], old, new, log, tag)


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs
    log: list[str] = []

    # ---- idempotency guards -------------------------------------------------
    if any("agree to within about 0.7%" in p.text for p in paras):
        raise SystemExit("ABORT: A4 already applied (0.7% prose present).")
    if any("E = 205,000 MPa" in p.text for p in paras):
        raise SystemExit("ABORT: B1 already applied (205,000 MPa present).")

    # ---- A1: P131 ------------------------------------------------------------
    p131 = paras[131]
    assert "maximum roof displacement is 5.4 mm" in p131.text, "P131 anchor moved"
    _replace_in_run(p131.runs[0], "5.4 mm", "5.3 mm", log, "A1 roof-X")
    _replace_in_run(p131.runs[0], "12.0 mm", "11.9 mm", log, "A1 roof-Y")
    _replace_in_run(p131.runs[2], "0.22%", "0.21%", log, "A1 drift-X")
    _replace_in_run(p131.runs[2], "0.47%", "0.46%", log, "A1 drift-Y")

    # ---- A2: P132 ------------------------------------------------------------
    p132 = paras[132]
    r = p132.runs[3]
    _replace_in_run(
        r,
        "0.507 at a second-story column under the H1-1b equation (governed by "
        "1.2DL + 1.0LL + 1.0EQY)",
        "0.467 at a first-story column under the H1-1a equation (governed by a seismic "
        "load combination with the Y-direction earthquake as the dominant lateral "
        "component)",
        log, "A2 interaction")
    _replace_in_run(r, "beam interaction ratio is 0.322",
                    "beam interaction ratio is 0.324", log, "A2 beam")

    # ---- A3: Table 9 ---------------------------------------------------------
    t9 = None
    for t in d.tables:
        hdr = [c.text for c in t.rows[0].cells]
        if hdr[:3] == ["Category", "Parameter", "Value"] and len(t.rows) == 11:
            t9 = t
            break
    if t9 is None:
        raise SystemExit("ABORT: Table 9 not located.")
    _replace_cell(t9.rows[4].cells[2], "5.4 / 12.0 mm", "5.3 / 11.9 mm", log, "A3 disp")
    _replace_cell(t9.rows[5].cells[2], "0.22% (ratio 0.108)", "0.21% (ratio 0.107)", log, "A3 drift-X")
    _replace_cell(t9.rows[6].cells[2], "0.47% (ratio 0.233)", "0.46% (ratio 0.232)", log, "A3 drift-Y")
    _replace_cell(t9.rows[8].cells[2], "0.507 (column, Story 2, H1-1b)",
                  "0.467 (column, Story 1, H1-1a)", log, "A3 interaction")
    _replace_cell(t9.rows[9].cells[2], "0.322 (Story 1)", "0.324 (Story 1)", log, "A3 beam")

    # ---- A4: P144 (full paragraph replace) -----------------------------------
    p144 = paras[144]
    if "carry 20.6% and 8.5% more" not in p144.text:
        raise SystemExit("ABORT: P144 does not contain the old 20.6%/8.5% gap text.")
    p144.runs[0].text = A4_NEW
    for extra in p144.runs[1:]:
        extra.text = ""
    log.append(f"  [A4] P144 replaced ({len(A4_NEW)} chars), {len(p144.runs)-1} extra run(s) cleared")

    # ---- B1: P130 append -----------------------------------------------------
    p130 = paras[130]
    assert p130.text.rstrip().endswith("Section 3.4."), f"P130 tail moved: {p130.text[-30:]!r}"
    p130.add_run(B1_SENTENCE)
    log.append("  [B1] P130 appended SS275 E=205,000 MPa sentence")

    print("Planned edits:")
    print("\n".join(log))

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply to write.")
        print("P131 now:", paras[131].text[:130])
        print("P132 now:", paras[132].text[:170])
        print("P144 now:", paras[144].text[:130])
        return 0

    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
