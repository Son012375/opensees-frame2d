"""Insert a new Section 2 'Related Work' (prose + comparison table) into the v2
DOCX, renumbering subsequent sections (2..6 -> 3..7) and tables (1..N -> 2..N+1).

Safe because all in-text 'Section N' / 'Table N' cross-references and all heading
leading numbers were verified to live within a single run. Atomic +1 increments
(read original, add one) avoid double-application. Backup written first.
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"

HEAD_RE = re.compile(r'^[2-6](\.\d+){0,2}\.?\s+[A-Za-z]')


def incr_section(t):
    return re.sub(r'Section\s+(\d+)((?:\.\d+)*)',
                  lambda m: f"Section {int(m.group(1)) + 1 if int(m.group(1)) >= 2 else int(m.group(1))}{m.group(2)}", t)


def incr_table(t):
    return re.sub(r'Table\s+(\d+)', lambda m: f"Table {int(m.group(1)) + 1}", t)


def bump_lead(p):
    if not p.runs:
        return
    m = re.match(r'^(\s*)(\d+)((?:\.\d+){0,2})(.*)$', p.runs[0].text, re.S)
    if not m:
        return
    n = int(m.group(2))
    if 2 <= n <= 6:
        p.runs[0].text = f"{m.group(1)}{n + 1}{m.group(3)}{m.group(4)}"


# ---- Section 2 content ----
P_INTRO = ("Three streams of prior work bear on the proposed pipeline: extraction of "
           "analysis-ready structural models from BIM/IFC data, open-source BIM-to-"
           "finite-element workflows built on OpenSees, and the emerging use of generative "
           "AI and large language models for structural analysis. The present work is "
           "positioned within the first two streams and deliberately defers the third.")
P_21 = ("Converting a coordination-oriented BIM into an analysis-ready model has been "
        "studied for over a decade. Ramaji and Memari (2018) interpret structural "
        "analytical models from the coordination view of building information models, and "
        "Hasan et al. (2019) generate geometrically accurate structural-analysis models "
        "within BIM-centered software. Fernández-Mora et al. (2022) review the broader "
        "integration of the structural project into the BIM paradigm and document the "
        "persistent gap between authoring models and analysis-ready models. More recently, "
        "Singh et al. (2024) automate the generation of a structural analytical model "
        "directly from an architectural IFC model to improve open-BIM interoperability, "
        "and Rudenko and Petryna (2025) generate finite-element models of varying "
        "complexity and dimensionality from BIM. These works establish the extraction "
        "methodology that the present pipeline adopts. The distinguishing feature here is "
        "an explicit, inspectable node-element analysis graph coupled with an automated "
        "connectivity-repair and validation-diagnostics layer (Section 3.3.2), rather than "
        "a single-pass translation whose intermediate representation is not surfaced for "
        "inspection.")
P_22 = ("OpenSees (McKenna, 2011) and its Python interface OpenSeesPy (Zhu et al., 2018) "
        "provide an open computational substrate that a growing number of workflows target. "
        "Leonardi et al. (2024) present a scalable, open BIM-based workflow for the "
        "structural analysis of masonry building aggregates, and Llanos and Delgadillo "
        "(2025) perform linear seismic analysis and optimization of reinforced-concrete "
        "frames in OpenSeesPy. An IFC-based framework has also been proposed to integrate "
        "structural-analysis results back into BIM for code-compliance support (Buildings, "
        "2026, 16(4):746). These confirm OpenSeesPy as a credible analysis engine for "
        "BIM-driven pipelines. The present work adds Korean Design Standard (KDS) "
        "clause-traced load automation and, distinctively, a metric-level cross-validation "
        "against two independent commercial programs (Midas Gen and ETABS); most BIM-to-FEM "
        "studies report agreement at the feasibility or visualization level rather than at "
        "the response-quantity level.")
P_23 = ("A rapidly growing 2024-2026 body of work applies generative AI and large language "
        "models (LLMs) to structural modeling. Liao et al. (2024) and Xie et al. (2025) "
        "survey generative-AI and AI-driven structural design automation, and Wang et al. "
        "(2025) apply BIM to the seismic reliability analysis of reinforced-concrete "
        "structures. Closest to an analysis pipeline, Liang et al. (2025) use an LLM to "
        "translate natural-language structural-analysis word problems into executable "
        "OpenSeesPy scripts, reporting full accuracy with GPT-4o on a 20-problem benchmark; "
        "and a Model Context Protocol approach (Buildings, 2025, 15(17):3190) couples GPT-4o "
        "to OpenSeesPy and ETABS, reporting agreement within roughly one percent on storey "
        "drift, nodal displacement, and vibrational period for reinforced-concrete frames. "
        "This cluster is the most active and the most architecturally adjacent to a "
        "natural-language front-end. The present paper deliberately excludes the LLM, "
        "retrieval-augmented, and chatbot layers from its contributions and validates only "
        "the underlying BIM-to-analysis substrate, so that the quantitative commercial-"
        "baseline evidence is not entangled with the reliability of LLM-generated output; "
        "the natural-language and KDS-retrieval-augmented layers are reserved for a separate "
        "study.")
P_24 = ("Table 1 positions the present work against representative recent systems across "
        "input modality, intermediate representation, connectivity repair and diagnostics, "
        "solver, code system, and the depth of commercial cross-validation. None of the "
        "individual components — IFC parsing, OpenSeesPy analysis, or KDS load lookup — is "
        "claimed as novel in isolation. The contribution is their joint realization as a "
        "transparent, inspectable, clause-traced toolchain for the Korean regulatory "
        "context, with a surfaced node-element representation and connectivity diagnostics, "
        "validated at the response-quantity level against two commercial references.")

TABLE1 = [
    ("System", "Primary input", "Representation", "Conn. repair / diagnostics", "Solver", "Code system", "Commercial cross-validation"),
    ("Ramaji & Memari (2018)", "IFC (coordination view)", "analytical model", "not surfaced", "external", "—", "none"),
    ("Hasan et al. (2019)", "BIM", "geometrically accurate FE", "partial", "BIM-centered", "—", "qualitative"),
    ("Leonardi et al. (2024)", "openBIM / IFC", "FE (masonry)", "not surfaced", "OpenSees-based", "Eurocode context", "feasibility-level"),
    ("Singh et al. (2024)", "architectural IFC", "structural analytical model", "not surfaced", "export to FE", "—", "none"),
    ("Rudenko & Petryna (2025)", "BIM", "multi-complexity FE", "not surfaced", "generic FE", "—", "none"),
    ("Llanos & Delgadillo (2025)", "parametric", "OpenSeesPy RC frame", "not surfaced", "OpenSeesPy", "—", "none"),
    ("Liang et al. (2025)", "natural language", "LLM-generated script", "not surfaced", "OpenSeesPy", "—", "20 analytical problems"),
    ("MCP approach (Buildings, 2025)", "natural language", "LLM + MCP commands", "not surfaced", "OpenSeesPy / ETABS", "—", "ETABS (~1%)"),
    ("This work", "IFC + NL + manual", "node-element analysis graph", "yes (merge / snap / split + diagnostics)", "OpenSeesPy", "KDS (clause-traced)", "Midas Gen + ETABS, 112 metrics"),
]
TBL_CAP = "Table 1. Positioning of the present work against representative recent systems."


def main() -> int:
    d = docx.Document(str(DOC))
    paras = d.paragraphs
    refs_idx = next(i for i, p in enumerate(paras) if p.text.strip().upper().startswith("REFERENCES"))

    # safety: confirm 'Table N' in-text refs are single-run
    split = 0
    for p in paras[:refs_idx]:
        for m in re.finditer(r'Table\s+\d+', p.text):
            if not any(m.group(0) in r.text for r in p.runs):
                split += 1
    if split:
        print(f"ABORT: {split} 'Table N' tokens split across runs")
        return 1

    shutil.copy2(DOC, DOC.with_name("open_source_alternative_review_draft_v2.pre_g5_backup.docx"))

    # ---- renumber existing sections + tables (paragraphs before References) ----
    for p in paras[:refs_idx]:
        for r in p.runs:
            if r.text:
                r.text = incr_section(incr_table(r.text))
        if HEAD_RE.match(p.text.strip()) and len(p.text.strip()) < 95:
            bump_lead(p)

    # ---- roadmap: add a clause for the new Section 2 ----
    for p in paras[:refs_idx]:
        if "remainder of this paper is organized as follows" in p.text:
            for r in p.runs:
                if "Section 3 presents the proposed" in r.text:
                    r.text = r.text.replace(
                        "Section 3 presents the proposed",
                        "Section 2 reviews related work and positions the present "
                        "contribution against recent systems; Section 3 presents the proposed")
            break

    # ---- anchor: the (now renumbered) '3. PROPOSED...' heading ----
    anchor = next(p for p in d.paragraphs
                  if p.text.strip().startswith("3.") and "PROPOSED" in p.text.upper())

    # build new section content (appended, then relocated before anchor)
    def mk_par(text, style):
        return d.add_paragraph(text, style=style)

    h1 = mk_par("2. RELATED WORK", "Heading 1")
    items = [
        mk_par(P_INTRO, "Normal"),
        mk_par("2.1 IFC/BIM-to-Analytical-Model Extraction", "Heading 2"),
        mk_par(P_21, "Normal"),
        mk_par("2.2 Open-Source BIM-to-FEM and OpenSeesPy Workflows", "Heading 2"),
        mk_par(P_22, "Normal"),
        mk_par("2.3 Generative AI and Large Language Models for Structural Analysis (Deferred)", "Heading 2"),
        mk_par(P_23, "Normal"),
        mk_par("2.4 Positioning", "Heading 2"),
        mk_par(P_24, "Normal"),
        mk_par(TBL_CAP, "Normal"),
    ]
    tbl = d.add_table(rows=len(TABLE1), cols=7)
    tbl.style = "Table Grid"
    for ri, row in enumerate(TABLE1):
        for ci, val in enumerate(row):
            c = tbl.cell(ri, ci)
            c.text = val
            for par in c.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(8)
                    if ri == 0 or ri == len(TABLE1) - 1:
                        run.font.bold = True

    cur = anchor._p
    cur.addprevious(h1._p)
    cur = h1._p
    for it in items:
        cur.addnext(it._p)
        cur = it._p
    cur.addnext(tbl._tbl)

    d.save(str(DOC))

    # ---- verify ----
    d2 = docx.Document(str(DOC))
    print("=== Heading 1 sections (in order) ===")
    for p in d2.paragraphs:
        if p.style.name == "Heading 1" and re.match(r'^\d+\.', p.text.strip()):
            print("  ", p.text.strip()[:55])
    caps = [p.text.strip()[:45] for p in d2.paragraphs if re.match(r'^Table \d+\.', p.text.strip())]
    print(f"=== table captions ({len(caps)}) ===")
    for c in caps:
        print("  ", c)
    print(f"tables={len(d2.tables)} inline_shapes={len(d2.inline_shapes)} paragraphs={len(d2.paragraphs)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
