"""Apply the G2 (Case 4 ablation) prose edits to the canonical v2 DOCX.

QW/G2 text swaps only (the Fig 8 image and Table 9 are handed off separately for
manual placement):
  [103]  Replace the qualitative "plausible explanation ... 3D element formulation"
         hand-wave with the quantified rigid-zone attribution (refs Table 9, Fig 8).
  [150]  Section 5.3 future-work item 3 — the Case 4 ablation is now DONE, so
         replace "should be studied through a targeted ablation ..." with a
         forward item (optional panel-zone modeling feature).

A backup (..._v2.pre_g2_backup.docx) is written first.
"""
from __future__ import annotations

import shutil
from pathlib import Path
import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"


def replace_in_para(p, old, new) -> bool:
    if old not in p.text:
        return False
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    full = p.text.replace(old, new)
    for i, r in enumerate(p.runs):
        r.text = full if i == 0 else ""
    return True


P103_OLD = ("A plausible explanation is that the discrepancy arises from differences "
            "in three-dimensional element formulation or response recovery between the "
            "two programs. Small differences in how 3D frame stiffness, local-axis "
            "treatment, or end-force recovery are handled can produce systematic offsets "
            "in displacement and moment quantities, which is consistent with the "
            "directional pattern observed in Section 3.3.")

P103_NEW = ("To attribute this discrepancy rather than leave it qualitative, the Case 4 "
            "frame was re-analyzed in OpenSeesPy under a controlled ablation of the "
            "candidate drivers (Table 9, Figure 8). Because the benchmark specification "
            "fixes shear deformation off and the section beta angle at zero in both "
            "programs, and because the OpenSeesPy and ETABS centerline models agree to "
            "0.00% on the affected metrics, the residual is a commercial-side stiffening "
            "effect rather than a solver error. Enabling shear deformation increases the "
            "Story 1 drift by 5.4%, the opposite direction, which rules it out as the "
            "cause. Introducing rigid end zones at the beam-column joints instead reduces "
            "the drift monotonically, from 4.1% above the Midas Gen value at the "
            "centerline limit (no rigid zone) to 4.0% below it at the full geometric "
            "panel-zone limit; a rigid-zone factor of 0.51, essentially the default "
            "panel-zone setting of Midas Gen, reproduces the Midas Story 1 drift to "
            "within 0.1% and the story and roof displacements to within 0.1%. "
            "Equivalently, the gap corresponds to a uniform 6.6% increase in effective "
            "column stiffness. The Case 4 differences are therefore consistent with "
            "beam-column joint rigid-zone (panel-zone) modeling that the centerline "
            "benchmark models deliberately omit, and they do not indicate an error in "
            "the open-source analysis chain.")

P150_OLD = ("Third, the element-formulation-sensitive response metrics observed in the "
            "primary three-dimensional benchmark case (Case 4) should be studied through "
            "a targeted ablation covering local-axis convention, beta angle, shear "
            "deformation (Euler-Bernoulli versus Timoshenko), rigid-offset treatment, "
            "and the rigid-diaphragm constraint.")

P150_NEW = ("Third, the beam-column joint rigid-zone effect identified for Case 4 in "
            "Section 3.3.2 could be offered as an optional modeling feature, so that the "
            "open-source workflow can reproduce the commercial panel-zone convention in "
            "addition to its default centerline model.")


def main() -> int:
    shutil.copy2(DOC, DOC.with_name("open_source_alternative_review_draft_v2.pre_g2_backup.docx"))
    d = docx.Document(str(DOC))
    log = []
    ok103 = any(replace_in_para(p, P103_OLD, P103_NEW) for p in d.paragraphs)
    log.append(f"[103] Case 4 attribution: {'APPLIED' if ok103 else 'NOT FOUND'}")
    ok150 = any(replace_in_para(p, P150_OLD, P150_NEW) for p in d.paragraphs)
    log.append(f"[150] sec5.3 item-3: {'APPLIED' if ok150 else 'NOT FOUND'}")
    if ok103 and ok150:
        d.save(str(DOC))
        log.append(f"SAVED -> {DOC.name}")
    else:
        log.append("NOT SAVED (an edit failed)")
    print("\n".join(log))
    return 0 if (ok103 and ok150) else 1


if __name__ == "__main__":
    raise SystemExit(main())
