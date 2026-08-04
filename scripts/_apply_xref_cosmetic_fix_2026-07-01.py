"""Safe, unambiguous cross-reference + spelling fixes surfaced by the full-manuscript
audit (all are PLAIN-TEXT single-run table cells — no fields, no SEQ, no numbers of
substance). Higher-risk / judgment items (Appendix A 18->36, P114/P115/Table6-7 claims,
Fig 7/8 in-image titles, P167/P091 field-adjacent refs) are handled separately.

  Table 10 (T10): "Section 2.3.3, Figure 6" -> "Section 3.3.3, Figure 7" (T10R1C2);
                  "...described in §2.3.3" -> "...§3.3.3" (T10R3C2)
                  (zone-based decomposition is §3.3.3 (P054/P142); the L-shape zone
                   figure is Figure 7 (P148); Figure 6 is the IFC-example screening.)
  Table 4  (T3): "See Section 2.3.1" -> "See Section 3.3.1" (T3R6C2)
                  (IFC endpoint extraction / section mapping is §3.3.1 (P046).)
  Table 2  (T1): British "storey" -> American "story" in the 4 descriptive cells
                  (the rest of the paper uses "story"); the IfcBuildingStorey entity
                  name is NOT touched.

Backup -> ..._v2.pre_xrefcosmetic_2026-07-01.docx. Dry-run by default; --apply saves.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "paper1_open_source_alternative" / "drafts" / "open_source_alternative_review_draft_v2.docx"
BACKUP = DOC.with_name("open_source_alternative_review_draft_v2.pre_xrefcosmetic_2026-07-01.docx")

# (table-locator substring, cell old text, cell new text)
CELL_EDITS = [
    ("Section 2.3.3, Figure 6", "Section 2.3.3, Figure 6", "Section 3.3.3, Figure 7"),
    ("described in §2.3.3", "Consistent with the zone representation described in §2.3.3",
     "Consistent with the zone representation described in §3.3.3"),
    ("See Section 2.3.1", "See Section 2.3.1", "See Section 3.3.1"),
    ("Global endpoints, section, material, storey relation",
     "Global endpoints, section, material, storey relation",
     "Global endpoints, section, material, story relation"),  # appears twice (T1R1,T1R2)
    ("Elevation and storey name", "Elevation and storey name", "Elevation and story name"),
    ("Node story assignment and storey height calculation",
     "Node story assignment and storey height calculation",
     "Node story assignment and story height calculation"),
]


def main(apply: bool) -> int:
    d = docx.Document(str(DOC))
    log: list[str] = []
    for _loc, old, new in CELL_EDITS:
        n = 0
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        for r in par.runs:
                            if old in r.text:
                                r.text = r.text.replace(old, new)
                                n += 1
        if n == 0:
            raise SystemExit(f"ABORT: no cell run matched {old[:40]!r}")
        log.append(f"  [{n}x] {old[:45]!r} -> {new[:45]!r}")

    print("Planned cell edits:")
    print("\n".join(log))

    if not apply:
        print("\nDRY-RUN (no save). Re-run with --apply.")
        return 0
    shutil.copy2(DOC, BACKUP)
    print(f"\nbackup -> {BACKUP.name}")
    d.save(str(DOC))
    print(f"SAVED -> {DOC.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(apply="--apply" in sys.argv))
